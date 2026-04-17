"""DOCX_RENDER node - convert markdown to .docx."""
import json
import logging
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..state import ReportState, WORKFLOW_RUNS_DIR

logger = logging.getLogger(__name__)


class QAHardBlockError(Exception):
    """Raised when QA gate fails."""
    pass


def markdown_to_docx(md_text: str, output_path: str) -> None:
    """Convert markdown to DOCX using python-docx."""
    doc = Document()
    
    lines = md_text.split("\n")
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            doc.add_paragraph()
            i += 1
            continue
        
        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # Bullet lists
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        # Numbered lists
        elif re.match(r"^\d+\.\s", line):
            match = re.match(r"^(\d+)\.\s(.*)", line)
            if match:
                doc.add_paragraph(match.group(2), style="List Number")
        # Tables
        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not re.match(r"^\|[\s\-\|]+\|$", lines[i].strip()):
                    table_lines.append(lines[i].strip())
                i += 1
            if table_lines:
                num_cols = len(table_lines[0].split("|")) - 2
                table = doc.add_table(rows=len(table_lines), cols=num_cols)
                for row_idx, tline in enumerate(table_lines):
                    cells = [c.strip() for c in tline.split("|")[1:-1]]
                    for col_idx, cell in enumerate(cells):
                        if col_idx < num_cols:
                            table.rows[row_idx].cells[col_idx].text = cell
            continue
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("*") and part.endswith("*"):
                    run = para.add_run(part[1:-1])
                    run.italic = True
                else:
                    para.add_run(part)
        
        i += 1
    
    doc.save(output_path)


def run_docx_render(state: ReportState) -> ReportState:
    """T15: DOCX_RENDER - convert markdown to .docx."""
    qa_decision = state.qa.get("qa_decision")
    
    if qa_decision and qa_decision != "pass":
        raise QAHardBlockError(f"QA gate failed: {qa_decision}")
    
    cited_md_path = state.drafts.get("merged_draft_cited_md")
    if not cited_md_path or not Path(cited_md_path).exists():
        cited_md_path = state.drafts.get("merged_draft_md")
    
    if not cited_md_path or not Path(cited_md_path).exists():
        raise QAHardBlockError("No merged draft found")
    
    with open(cited_md_path) as f:
        md_content = f.read()
    
    run_dir = WORKFLOW_RUNS_DIR / state.job_id
    output_dir = Path(state.output.get("output_dir", run_dir))
    output_dir.mkdir(parents=True, exist_ok=True)
    
    final_docx_path = output_dir / "final.docx"
    
    try:
        markdown_to_docx(md_content, str(final_docx_path))
    except Exception as exc:
        logger.exception("[DOCX_RENDER] markdown_to_docx failed, falling back to stub document")
        try:
            doc = Document()
            doc.add_heading("Report", level=1)
            doc.add_paragraph(md_content[:1000])
            doc.save(str(final_docx_path))
        except Exception as save_exc:
            logger.exception("[DOCX_RENDER] fallback doc.save also failed")
            state.runtime["error"] = f"DOCX_RENDER failed: {type(exc).__name__}: {exc}; fallback save failed: {save_exc}"
            raise QAHardBlockError(f"DOCX render failed and fallback save also failed: {save_exc}") from save_exc
    
    state.output["final_docx_path"] = str(final_docx_path)
    return state
