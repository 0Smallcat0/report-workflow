"""Semi-structured parser for PDF, DOCX, TXT files."""
from pathlib import Path
from typing import Any, Optional


def parse_pdf(file_path: str) -> dict:
    """Parse PDF using pdfplumber."""
    try:
        import pdfplumber
        blocks = []
        all_text = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    all_text.append(text)
                    blocks.append({
                        "block_id": f"page_{page_num}",
                        "block_type": "paragraph",
                        "content": text,
                        "page_number": page_num,
                        "table_data": None
                    })
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    blocks.append({
                        "block_id": f"page_{page_num}_table_{t_idx}",
                        "block_type": "table",
                        "content": "",
                        "page_number": page_num,
                        "table_data": table
                    })
        return {
            "blocks": blocks,
            "raw_content": "\n\n".join(all_text),
            "success": True
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_docx(file_path: str) -> dict:
    """Parse DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        blocks = []
        all_text = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)
            block_type = "paragraph"
            if para.style.name.startswith("Heading"):
                block_type = "heading"
            blocks.append({
                "block_id": f"para_{i}",
                "block_type": block_type,
                "content": text,
                "page_number": None,
                "table_data": None
            })
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            blocks.append({
                "block_id": f"table_{t_idx}",
                "block_type": "table",
                "content": "",
                "page_number": None,
                "table_data": rows
            })
        return {
            "blocks": blocks,
            "raw_content": "\n\n".join(all_text),
            "success": True
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_txt(file_path: str) -> dict:
    """Parse TXT file directly."""
    try:
        with open(file_path, encoding="utf-8") as f:
            content = f.read()
        blocks = [{
            "block_id": "main",
            "block_type": "paragraph",
            "content": content,
            "page_number": None,
            "table_data": None
        }]
        return {
            "blocks": blocks,
            "raw_content": content,
            "success": True
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_semi_structured(file_path: str, file_type: str) -> dict:
    """Parse semi-structured file and return content blocks."""
    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    elif file_type == "txt":
        return parse_txt(file_path)
    else:
        return {"blocks": [], "error": f"Unsupported file type: {file_type}", "success": False}
