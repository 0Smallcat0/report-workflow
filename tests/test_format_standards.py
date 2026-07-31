"""Format-standards regressions: page numbers, TOC placement/language, math, CJK front matter."""
import re
import shutil
import tempfile
import unittest
import zipfile
from pathlib import Path

from report_workflow.errors import QAHardBlockError
from report_workflow.nodes.docx_render import (
    _REFERENCE_DOC,
    _find_pandoc,
    _inject_toc,
    _render_via_pandoc,
    _resolve_reference_doc,
    _toc_openxml_block,
    reference_docx_error,
)
from report_workflow.nodes.front_matter_build import (
    _parse_affiliation_from_user_prompt,
    _parse_author_from_user_prompt,
    _parse_title_from_user_prompt,
)


class TocInjectionTest(unittest.TestCase):
    def test_zh_document_gets_zh_toc_after_front_matter(self):
        md = (
            "# 管線監控導入報告\n\n作者:王小明\n\n---\n\n# 緒論\n\n"
            "本研究說明監控管線的導入流程與觀察結果,並整理後續維運建議。"
            "全文以中文撰寫,涵蓋方法、結果與討論。"
        )
        out = _inject_toc(md, has_front_matter=True)
        self.assertIn("目錄", out)
        self.assertIn('TOC \\o "1-3"', out)
        # Title page stays first; body follows the TOC.
        self.assertLess(out.index("作者"), out.index("目錄"))
        self.assertLess(out.index("目錄"), out.index("緒論"))

    def test_en_document_gets_en_toc(self):
        md = "Title page\n\n---\n\n# Introduction\n\nEnglish body text for detection."
        out = _inject_toc(md, has_front_matter=True)
        self.assertIn("Table of Contents", out)
        self.assertNotIn("目錄", out)

    def test_without_front_matter_toc_leads_document(self):
        md = "# Introduction\n\nBody."
        out = _inject_toc(md, has_front_matter=False)
        self.assertTrue(out.startswith("```{=openxml}"))
        self.assertIn("# Introduction", out)

    def test_block_uses_toc_heading_style_and_field(self):
        block = _toc_openxml_block("en", page_break_before=True)
        self.assertIn('w:pStyle w:val="TOCHeading"', block)
        self.assertIn('w:fldCharType="begin"', block)
        self.assertIn('w:br w:type="page"', block)


class ReferenceTemplateTest(unittest.TestCase):
    def test_template_footer_has_page_field(self):
        with zipfile.ZipFile(_REFERENCE_DOC) as z:
            footers = [n for n in z.namelist() if re.match(r"word/footer\d*\.xml", n)]
            self.assertTrue(footers, "reference.docx must carry a footer for page numbers")
            self.assertTrue(
                any("PAGE" in z.read(n).decode("utf-8") for n in footers),
                "footer must contain a PAGE field",
            )


class CoverTocPlacementTest(unittest.TestCase):
    def test_cover_promoted_to_title_page_with_toc_after(self):
        md = (
            "# 封面\n\n本報告為機械工程實驗課程之實驗報告,依據講義與量測數據撰寫。\n\n"
            "# 1. 實驗目的\n\n量測懸臂樑自由端撓度並與理論比較,驗證線性關係與誤差來源。"
        )
        out = _inject_toc(md, has_front_matter=False, cover_title="封面")
        # The cover heading is dropped (a title page does not label itself,
        # and without a Heading 1 it stays out of the TOC field).
        self.assertNotIn("# 封面", out)
        self.assertIn('w:jc w:val="center"', out)
        self.assertIn("目錄", out)
        self.assertLess(out.index("本報告為"), out.index("目錄"))
        self.assertLess(out.index("目錄"), out.index("實驗目的"))

    def test_unrelated_first_heading_keeps_toc_on_top(self):
        md = "# Introduction\n\nBody text.\n\n# Methods\n\nMore body."
        out = _inject_toc(md, has_front_matter=False, cover_title="封面")
        self.assertTrue(out.startswith("```{=openxml}"))

    def test_revised_document_title_leads_and_toc_follows(self):
        md = "# My Report Title\n\n# Introduction\n\nBody text for the report."
        out = _inject_toc(md, has_front_matter=False, title_leads=True)
        self.assertLess(out.index("My Report Title"), out.index("Table of Contents"))
        self.assertLess(out.index("Table of Contents"), out.index("Introduction"))


class TableFigureTitleTest(unittest.TestCase):
    def test_table_title_uses_columns(self):
        from report_workflow.nodes.figure_recommend import _human_figure_title

        title = _human_figure_title(
            "table",
            "",
            "",
            {"columns": ["荷重(N)", "實測撓度(mm)", "理論撓度(mm)", "誤差(%)"]},
            "量測數據.csv",
            "",
        )
        self.assertIn("實測撓度", title)
        self.assertIn("依荷重", title)
        self.assertNotIn("view of", title)


class DerivedStatsEvidenceTest(unittest.TestCase):
    def _registry(self, cols):
        import json

        rows = [
            {cols[0]: "5", cols[1]: "1.52", cols[2]: "1.45", cols[3]: "4.8"},
            {cols[0]: "10", cols[1]: "3.01", cols[2]: "2.90", cols[3]: "3.8"},
            {cols[0]: "15", cols[1]: "4.48", cols[2]: "4.35", cols[3]: "3.0"},
        ]
        blocks = [
            {"block_type": "csv_row", "content": json.dumps(r, ensure_ascii=False)}
            for r in rows
        ]
        return [{
            "source_id": "s1",
            "file_name": "m.csv",
            "file_path": "m.csv",
            "file_type": "csv",
            "parsed_content": blocks,
        }]

    def test_en_measurement_columns_produce_citable_stats(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(
            self._registry(
                ["Load (N)", "Measured deflection (mm)", "Theoretical deflection (mm)", "Error (%)"]
            ),
            "2026-07-22T00:00:00+00:00",
        )
        self.assertEqual(len(units), 2)
        regression = units[0]
        self.assertEqual(regression["evidence_grade"], "high")
        self.assertIn("least-squares slope", regression["content"])
        self.assertIn("R²", regression["content"])
        self.assertEqual(regression["derivation"]["method"], "least_squares_fit")
        self.assertIn("mean", units[1]["content"])

    def test_rated_column_is_the_reference_curve(self):
        """A reference column is rarely labelled "theoretical" on a real sheet.

        Both runs that motivated this carried the manufacturer's figure as
        "Rated Effectiveness" and 廠商標稱有效度, so the comparison against it
        never fired in either language — even though the code has carried an
        output template for both all along. Without it the author has to
        compute the ratio by hand, and the gates then correctly refuse to
        publish a number with no evidence behind it.
        """
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(
            self._registry(
                ["Flow (L/min)", "Measured effectiveness", "Rated effectiveness", "Error (%)"]
            ),
            "2026-07-22T00:00:00+00:00",
        )
        self.assertTrue(units)
        regression = units[0]
        self.assertIn("Rated effectiveness", regression["content"])
        self.assertIn("Rated effectiveness", regression["derivation"]["input_columns"])

    def test_zh_nominal_column_is_the_reference_curve(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(
            self._registry(["流量(L/min)", "實測有效度", "廠商標稱有效度", "誤差(%)"]),
            "2026-07-22T00:00:00+00:00",
        )
        self.assertTrue(units)
        self.assertIn("理論斜率", units[0]["content"])
        self.assertIn("廠商標稱有效度", units[0]["derivation"]["input_columns"])

    def test_a_flow_rate_column_is_not_a_rated_value(self):
        """"Rate" is not "rated" — widening the tokens must not turn an
        independent variable into the reference curve."""
        from report_workflow.nodes.evidence_normalize import _THEORETICAL_COL_RE

        self.assertIsNone(_THEORETICAL_COL_RE.search("Flow Rate (L/min)"))
        self.assertIsNone(_THEORETICAL_COL_RE.search("Sampling rate (Hz)"))

    def test_zh_columns_produce_zh_content(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(
            self._registry(["荷重(N)", "實測撓度(mm)", "理論撓度(mm)", "誤差(%)"]),
            "2026-07-22T00:00:00+00:00",
        )
        self.assertEqual(len(units), 2)
        self.assertIn("最小平方法", units[0]["content"])
        self.assertIn("平均", units[1]["content"])

    def test_unrelated_columns_add_nothing(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(
            self._registry(["A", "B", "C", "D"]), "2026-07-22T00:00:00+00:00"
        )
        self.assertEqual(units, [])


class BudgetTotalDerivedStatsTest(unittest.TestCase):
    """A budget is read for its total, and no row states it."""

    def _quote_registry(self):
        import json

        rows = [
            {"品項": "加速規", "單價": "6800", "數量": "2", "小計": "13600"},
            {"品項": "訊號線", "單價": "450", "數量": "2", "小計": "900"},
            {"品項": "耗材", "單價": "300", "數量": "1", "小計": "300"},
        ]
        return [{
            "source_id": "s1",
            "file_name": "quote.csv",
            "file_path": "quote.csv",
            "file_type": "csv",
            "parsed_content": [
                {"block_type": "csv_row", "content": json.dumps(r, ensure_ascii=False)}
                for r in rows
            ],
        }]

    def test_amount_column_total_is_citable(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(self._quote_registry(), "2026-07-25T00:00:00+00:00")
        self.assertEqual(len(units), 1)
        total = units[0]
        self.assertEqual(total["derivation"]["method"], "column_total")
        self.assertEqual(total["derivation"]["input_columns"], ["小計"])
        self.assertEqual(total["evidence_grade"], "high")
        self.assertIn("14,800", total["content"])
        self.assertIn("13,600", total["content"])

    def test_product_column_is_recognized(self):
        from report_workflow.nodes.evidence_normalize import _is_product_column

        numeric = {"unit price": [10.0, 4.0], "qty": [3.0, 5.0], "cost": [30.0, 20.0]}
        self.assertTrue(_is_product_column(numeric, "cost"))
        self.assertFalse(_is_product_column(numeric, "qty"))

    def test_two_row_table_still_gets_a_total(self):
        import json

        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        rows = [{"item": "a", "total": "100"}, {"item": "b", "total": "250"}]
        registry = [{
            "source_id": "s1",
            "file_name": "small.csv",
            "file_path": "small.csv",
            "file_type": "csv",
            "parsed_content": [
                {"block_type": "csv_row", "content": json.dumps(r)} for r in rows
            ],
        }]
        units = _derived_stats_units(registry, "2026-07-25T00:00:00+00:00")
        self.assertEqual(len(units), 1)
        self.assertIn("350", units[0]["content"])


class OptionsComparisonTotalTest(unittest.TestCase):
    """Alternatives are not line items: their costs must not be summed."""

    def _registry(self, label_header):
        import json

        rows = [
            {label_header: "沿用既有設備", "採購成本": "0"},
            {label_header: "開源自組", "採購成本": "8500"},
            {label_header: "商用套裝", "採購成本": "86000"},
        ]
        return [{
            "source_id": "s1",
            "file_name": "compare.csv",
            "file_path": "compare.csv",
            "file_type": "csv",
            "parsed_content": [
                {"block_type": "csv_row", "content": json.dumps(r, ensure_ascii=False)}
                for r in rows
            ],
        }]

    def test_option_table_gets_no_total(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(self._registry("方案"), "2026-07-26T00:00:00+00:00")
        self.assertEqual(units, [])

    def test_line_item_table_still_gets_a_total(self):
        from report_workflow.nodes.evidence_normalize import _derived_stats_units

        units = _derived_stats_units(self._registry("品項"), "2026-07-26T00:00:00+00:00")
        self.assertEqual(len(units), 1)
        self.assertIn("94,500", units[0]["content"])


class UnitSignatureTest(unittest.TestCase):
    """An unrecognized unit is still a unit."""

    def test_cjk_and_unknown_units_are_distinguished(self):
        from report_workflow.nodes.figure_utils import unit_signature

        rate = unit_signature("最高取樣率(kS/s)")
        cost = unit_signature("採購成本(元)")
        bits = unit_signature("解析度(bit)")
        for signature in (rate, cost, bits):
            self.assertNotEqual(signature, "")
        self.assertEqual(len({rate, cost, bits}), 3)

    def test_column_without_parentheses_has_no_unit(self):
        from report_workflow.nodes.figure_utils import unit_signature

        self.assertEqual(unit_signature("輸入通道"), "")

    def test_year_parenthetical_is_not_a_unit(self):
        from report_workflow.nodes.figure_utils import unit_signature

        self.assertEqual(unit_signature("Revenue (2026)"), "")

    def test_known_units_keep_their_canonical_token(self):
        from report_workflow.nodes.figure_utils import unit_signature

        self.assertEqual(unit_signature("載重(kg)"), "kg")
        self.assertEqual(unit_signature("不良率(%)"), "%")


class CjkAbstractLengthTest(unittest.TestCase):
    """Policy bounds are English word counts; CJK is counted per character."""

    def test_chinese_abstract_is_measured_against_scaled_bounds(self):
        from report_workflow.nodes.abstract_check import _word_count_check

        text = "本報告說明變轉速條件下的軸承診斷方法與量化結果。" * 14
        self.assertEqual(_word_count_check(text, "admissions_project_report"), [])

    def test_chinese_abstract_can_still_be_too_short(self):
        from report_workflow.nodes.abstract_check import _word_count_check

        text = "本報告說明變轉速條件下的軸承診斷方法與量化結果。" * 5
        errors = _word_count_check(text, "admissions_project_report")
        self.assertTrue(errors)
        self.assertIn("characters", errors[0])
        self.assertIn("minimum 300", errors[0])

    def test_english_bounds_are_unchanged(self):
        from report_workflow.nodes.abstract_check import _word_count_check

        errors = _word_count_check("word " * 400, "admissions_project_report")
        self.assertTrue(errors)
        self.assertIn("maximum 250", errors[0])


class ReferencesHeadingLevelTest(unittest.TestCase):
    def test_references_sits_at_the_same_level_as_its_siblings(self):
        from report_workflow.nodes.heading_contract_check import _canonical_heading

        self.assertEqual(_canonical_heading("references", "參考文獻", None), "# 參考文獻")
        self.assertEqual(_canonical_heading("abstract", "摘要", None), "# 摘要")
        self.assertEqual(_canonical_heading("introduction", "緒論", 1), "# 1. 緒論")


class MarkdownTableProvenanceTest(unittest.TestCase):
    """The same table must not score lower for living inside a Markdown file."""

    TABLE = (
        "| 項目 | 權重 | | --- | --- | | 計畫書與期中審查 | 30% | "
        "| 實作完成度與數據品質 | 40% | | 期末書面報告 | 20% |"
    )

    def test_markdown_table_scores_like_a_structured_row(self):
        from report_workflow.nodes.evidence_normalize import compute_provenance_score

        score = compute_provenance_score(
            {"file_type": "md"}, {"block_type": "paragraph", "content": self.TABLE}
        )
        self.assertAlmostEqual(score, 0.75, places=6)

    def test_plain_prose_is_not_promoted(self):
        from report_workflow.nodes.evidence_normalize import compute_provenance_score

        score = compute_provenance_score(
            {"file_type": "md"},
            {"block_type": "paragraph", "content": "本組於 2026 年量測 3 次，結果一致。"},
        )
        self.assertAlmostEqual(score, 0.5, places=6)


class ProposalFigureSectionTest(unittest.TestCase):
    """`proposal` has no results-like section; a figure must not target one."""

    def test_falls_back_to_a_section_the_blueprint_defines(self):
        from report_workflow.nodes.figure_recommend import _section_for_recommendation

        class _State:
            plan = {
                "blueprint": {
                    "sections": {
                        "executive_summary": {"section_type": "exec_summary"},
                        "problem_statement": {"section_type": "problem"},
                        "budget_resources": {"section_type": "budget"},
                    }
                }
            }

        self.assertEqual(_section_for_recommendation(_State()), "budget_resources")

    def test_results_still_wins_when_present(self):
        from report_workflow.nodes.figure_recommend import _section_for_recommendation

        class _State:
            plan = {"blueprint": {"sections": {"results": {"section_type": "results"}}}}

        self.assertEqual(_section_for_recommendation(_State()), "results")


class BarLabelColumnTest(unittest.TestCase):
    def test_label_column_prefers_distinct_values(self):
        from report_workflow.nodes.figure_recommend import _distinct_label_index

        rows = [
            ["品項", "規格", "小計"],
            ["測試軸承", "健康件", "720"],
            ["測試軸承", "外環缺陷", "1560"],
            ["測試軸承", "內環缺陷", "1560"],
            ["耗材", "黏著劑", "300"],
        ]
        self.assertEqual(_distinct_label_index(rows, [0, 1]), 1)

    def test_single_categorical_column_is_kept(self):
        from report_workflow.nodes.figure_recommend import _distinct_label_index

        rows = [["name", "value"], ["a", "1"], ["b", "2"]]
        self.assertEqual(_distinct_label_index(rows, [0]), 0)


class LocalArtifactReferenceTest(unittest.TestCase):
    """Every local-file label the citation formatter emits must be excluded."""

    LABELS = ("[Text file]", "[Word document]", "[Dataset]", "[Data file]")

    def test_no_local_artifact_reaches_the_publication_list(self):
        from report_workflow.nodes.reference_verify import (
            _check_reference_curation,
            _is_publication_reference_candidate,
        )

        for label in self.LABELS:
            ref = f"measurements. (n.d.). *measurements* {label}."
            with self.subTest(label=label):
                self.assertFalse(_is_publication_reference_candidate(ref))
                self.assertFalse(_check_reference_curation(ref)[0])

    def test_real_publication_is_still_kept(self):
        from report_workflow.nodes.reference_verify import (
            _is_publication_reference_candidate,
        )

        self.assertTrue(_is_publication_reference_candidate(
            "Kording, K. (2017). Ten simple rules. *PLOS Comput Biol*, 13(9), "
            "e1005619. doi: 10.1371/journal.pcbi.1005619"
        ))


class FigureShortfallHintTest(unittest.TestCase):
    """A figure count on its own is not actionable.

    Three ways to end up with fewer embedded figures than expected, and the
    hint must distinguish them: a placeholder naming an id that was never
    built, an outline referencing one, and a figure built that no section
    mentions at all — the last surfaced in a revision run, where section
    drafts are validated but never become body content.
    """

    def _state(self, tmpdir, built_ids):
        import json as _json

        from report_workflow.state import ReportState

        state = ReportState.new("render report", [], str(Path(tmpdir) / "out"))
        manifest = Path(tmpdir) / "figure_manifest.json"
        manifest.write_text(
            _json.dumps({"figures": [{"figure_id": fid} for fid in built_ids]}),
            encoding="utf-8",
        )
        state.output["figure_manifest_path"] = str(manifest)
        return state

    def test_unresolved_placeholder_is_named(self):
        from report_workflow.nodes.post_render_validate import _figure_shortfall_hint

        with tempfile.TemporaryDirectory() as tmpdir:
            hint = _figure_shortfall_hint(
                self._state(tmpdir, ["1"]), "prose [FIGURE:figrec_1] more", {"figrec_1"}
            )
            self.assertIn("[FIGURE:figrec_1]", hint)
            self.assertIn("Built figure ids: 1", hint)

    def test_built_but_unreferenced_figure_is_explained(self):
        from report_workflow.nodes.post_render_validate import _figure_shortfall_hint

        with tempfile.TemporaryDirectory() as tmpdir:
            hint = _figure_shortfall_hint(self._state(tmpdir, ["1"]), "prose only", set())
            self.assertIn("no section references them", hint)
            self.assertIn("Built figure ids: 1", hint)

    def test_a_templates_own_cover_page_is_reported_as_not_carried_over(self):
        # A course hands out a .docx with the cover it grades. --reference-doc
        # takes styles only, so the fonts came through, the cover did not, and
        # the template was accepted without a word about it.
        from docx import Document as _Document
        from report_workflow.nodes.docx_render import _reference_doc_body_carryover

        with tempfile.TemporaryDirectory() as td:
            template = Path(td) / "course_template.docx"
            doc = _Document()
            doc.add_paragraph("國立成功大學 機械工程學系")
            doc.add_paragraph("姓名：____________  學號：____________")
            doc.save(str(template))

            carried = _reference_doc_body_carryover(template)
            self.assertIn("國立成功大學", carried)
            self.assertIn("姓名：", carried)

            empty = Path(td) / "styles_only.docx"
            _Document().save(str(empty))
            self.assertEqual(_reference_doc_body_carryover(empty), "")
            self.assertEqual(_reference_doc_body_carryover(Path(td) / "nope.docx"), "")

    def test_what_pandoc_could_not_render_reaches_the_issue_list(self):
        # pandoc names what it dropped: a figure replaced by its alt text, a
        # formula printed as raw TeX in a submitted report. That went to
        # logger.info truncated at 300 characters — shorter than the TeX
        # warning itself, so even the log was cut mid-sentence.
        from report_workflow.nodes.docx_render import _pandoc_warnings

        stderr = (
            "[WARNING] Could not fetch resource chart.png: replacing image with description\n"
            "[WARNING] Could not convert TeX math C_D = \\frac{2F}{\\rho U^2 A, rendering as TeX:\n"
            "  frac{2F}{\\rho U^2 A\n"
            "                     ^\n"
            "  unexpected eof\n"
            "  expecting \"\\\\bangle\", \"{\", letter, digit\n"
            "[WARNING] Could not fetch resource chart.png: replacing image with description\n"
        )
        found = _pandoc_warnings(stderr)
        self.assertEqual(len(found), 2)
        self.assertIn("replacing image with description", found[0])
        self.assertIn("rendering as TeX", found[1])
        # The parser's expectation dump is for whoever fixes the markup.
        self.assertFalse([item for item in found if "unexpected eof" in item])
        self.assertEqual(_pandoc_warnings(""), [])

    def test_a_chart_that_cannot_draw_its_own_labels_says_so(self):
        # The font list falls back to DejaVu Sans, which has no CJK glyphs, so
        # on a machine with no Chinese font — an ordinary Linux runner, or the
        # Colab notebook this project advertises — every Chinese axis label
        # became an empty box and the report went out carrying them.
        # matplotlib warns about each character; nobody was listening.
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from report_workflow.nodes.figure_build import _save_figure

        def issues_for(fonts, label):
            previous = matplotlib.rcParams["font.sans-serif"]
            matplotlib.rcParams["font.sans-serif"] = fonts
            figure, axes = plt.subplots()
            try:
                axes.plot([2.0, 3.0], [72.4, 76.1])
                axes.set_xlabel(label)
                with tempfile.TemporaryDirectory() as tmpdir:
                    reported = _save_figure(
                        figure, axes, "fig_1", "line", {},
                        Path(tmpdir) / "fig.png", 100,
                    )
            finally:
                plt.close(figure)
                matplotlib.rcParams["font.sans-serif"] = previous
            return [i for i in reported if i["type"] == "characters_missing_from_font"]

        found = issues_for(["DejaVu Sans"], "流量 (L/min)")
        self.assertTrue(found)
        self.assertIn("empty boxes", found[0]["detail"])
        self.assertIn("font", found[0]["repair_hint"])

        # No false alarm when the characters are drawable.
        self.assertFalse(issues_for(["DejaVu Sans"], "Flow rate (L/min)"))

    def test_the_word_after_figure_is_not_a_figure_number(self):
        # "kept in figure metadata rather than recomputed in prose" hard-blocked
        # a rendered report for citing figure "m" — the first letter of the next
        # word. "the figure shows" would have named figure "s".
        from report_workflow.nodes.post_render_validate import _FIGURE_MENTION_RE

        def mentions(text):
            return sorted(m.group(1).lower() for m in _FIGURE_MENTION_RE.finditer(text))

        self.assertEqual(mentions("kept in figure metadata rather than prose"), [])
        self.assertEqual(mentions("the figure shows a rising trend"), [])
        self.assertEqual(mentions("as Figure 3 shows"), ["3"])
        self.assertEqual(mentions("see Fig. 12 below"), ["12"])
        self.assertEqual(mentions("Figure A: apparatus"), ["a"])
        self.assertEqual(mentions("由圖 2 可見"), ["2"])
        # A trailing \b would have been the obvious guard and would have broken
        # this: CJK are word characters, so there is no boundary after the 3.
        self.assertEqual(mentions("由圖3可見有效度上升"), ["3"])

    def test_no_figures_at_all_stays_silent(self):
        from report_workflow.nodes.post_render_validate import _figure_shortfall_hint

        with tempfile.TemporaryDirectory() as tmpdir:
            self.assertEqual(
                _figure_shortfall_hint(self._state(tmpdir, []), "prose only", set()), ""
            )


class ClaimlessSectionTypeTest(unittest.TestCase):
    """A required cover page cannot cite anything, so it must not be asked to.

    Found writing a real lab report (2026-07-27): `cover` is required in the
    outline, and PLAN_LOCK then hard-blocked it for carrying no claims. Omitting
    it failed the early gate; including it failed the last one, after all the
    drafting was already done. Exemption is a property of the section type, not
    a list of ids to remember per profile.
    """

    def _blueprints(self):
        import glob

        import yaml

        for path in sorted(glob.glob("src/report_workflow/blueprints/*.yaml")):
            with open(path, encoding="utf-8") as f:
                yield path, yaml.safe_load(f)

    def test_front_matter_never_requires_claims_in_any_profile(self):
        from report_workflow.nodes.section_contract import (
            CLAIMLESS_SECTION_TYPES,
            section_requires_claims,
        )

        checked = 0
        for path, blueprint in self._blueprints():
            for section_id, section in (blueprint.get("sections") or {}).items():
                expected = section.get("section_type") not in CLAIMLESS_SECTION_TYPES
                self.assertEqual(
                    section_requires_claims(blueprint, section_id), expected,
                    f"{path}:{section_id} ({section.get('section_type')})",
                )
                checked += 1
        self.assertGreater(checked, 0, "no blueprint sections were checked")

    def test_a_required_section_may_still_be_claimless(self):
        from report_workflow.nodes.section_contract import (
            section_requires_claims,
            validate_required_outline_sections,
        )

        blueprint = {"sections": {
            "cover": {"required": True, "section_type": "front_matter"},
            "results": {"required": True, "section_type": "results"},
        }}
        validate_required_outline_sections(blueprint, {"cover": {}, "results": {}})
        self.assertFalse(section_requires_claims(blueprint, "cover"))
        self.assertTrue(section_requires_claims(blueprint, "results"))


class LocalArtifactTableConsistencyTest(unittest.TestCase):
    """One table drives the label, the citation, and the curation filter.

    Iterating the table is the point: a newly supported source format that is
    not flagged fails here instead of reaching a rendered document.
    """

    def test_every_local_type_is_curated_out_and_left_uncited(self):
        from report_workflow.nodes.citation_bind import (
            LOCAL_ARTIFACT_FILE_TYPES,
            _format_apa_reference_entry,
            _format_in_text_citation,
        )
        from report_workflow.nodes.reference_verify import (
            _check_reference_curation,
            _is_publication_reference_candidate,
        )

        for file_type in sorted(LOCAL_ARTIFACT_FILE_TYPES):
            reference = _format_apa_reference_entry(
                f"measurements.{file_type}", file_type, "s1"
            )
            with self.subTest(file_type=file_type):
                self.assertFalse(_is_publication_reference_candidate(reference))
                self.assertFalse(_check_reference_curation(reference)[0])
                self.assertEqual(
                    _format_in_text_citation({
                        "source_role": "primary_source",
                        "source_file_name": f"measurements.{file_type}",
                        "file_type": file_type,
                    }),
                    "",
                )

    def test_a_pdf_source_is_still_a_publication(self):
        from report_workflow.nodes.citation_bind import (
            LOCAL_ARTIFACT_FILE_TYPES,
            _format_in_text_citation,
        )

        self.assertNotIn("pdf", LOCAL_ARTIFACT_FILE_TYPES)
        self.assertTrue(
            _format_in_text_citation({
                "source_role": "research_document",
                "source_file_name": "kording2017.pdf",
                "file_type": "pdf",
            }).startswith("(")
        )


class LocalArtifactCitationTest(unittest.TestCase):
    """A citation must point at something that appears in the reference list."""

    def test_local_file_sources_emit_no_in_text_citation(self):
        from report_workflow.nodes.citation_bind import _format_in_text_citation

        for file_type in ("csv", "json", "docx", "md", "txt"):
            with self.subTest(file_type=file_type):
                self.assertEqual(
                    _format_in_text_citation({
                        "source_role": "primary_source",
                        "source_file_name": f"quote.{file_type}",
                        "file_type": file_type,
                    }),
                    "",
                )

    def test_the_numeric_notation_suppresses_them_too(self):
        # The rule above was enforced only where the author-year formatter runs.
        # The GB/T numeric branch never called it, so a lab report citing its
        # own .csv put [1] in the prose, curation removed the entry it pointed
        # at, and the document went out with thirteen markers over an empty
        # bibliography.
        from report_workflow.nodes.citation_bind import _is_local_artifact

        for file_type in ("csv", "json", "docx", "md", "txt"):
            with self.subTest(file_type=file_type):
                self.assertTrue(_is_local_artifact({"file_type": file_type}))
        self.assertFalse(_is_local_artifact({"file_type": "pdf"}))

    def test_a_marker_with_no_entry_is_refused_after_render(self):
        # The safety net for the same invariant, asked of the deliverable
        # itself: whatever binding and curation did between them, a number in
        # the prose must be findable in the list.
        from report_workflow.nodes.post_render_validate import (
            _NUMERIC_REFERENCE_ENTRY_RE,
            _cited_numbers,
        )

        self.assertEqual(_cited_numbers("see [2,3] and [4-6] and [1]"), {1, 2, 3, 4, 5, 6})
        # A bracketed date is not a citation.
        self.assertEqual(_cited_numbers("dated [2026-05-13]"), set())

        listed = {int(m.group(1)) for m in _NUMERIC_REFERENCE_ENTRY_RE.finditer(
            "9. Conclusion\n\n[1] 王小明. 熱傳[J]. 2024.\n"
        )}
        self.assertEqual(listed, {1})

    def test_pdf_source_still_cites(self):
        from report_workflow.nodes.citation_bind import _format_in_text_citation

        citation = _format_in_text_citation({
            "source_role": "research_document",
            "source_file_name": "kording2017.pdf",
            "file_type": "pdf",
        })
        self.assertTrue(citation.startswith("("))

    def test_undated_duplicate_citations_collapse(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        text = "兩者皆為優先補助項目(估價單 (n.d.))(估價單 (n.d.))(估價單 (n.d.))。"
        self.assertEqual(
            _collapse_adjacent_duplicate_citations(text),
            "兩者皆為優先補助項目(估價單 (n.d.))。",
        )


class FactualityBlockMessageTest(unittest.TestCase):
    def test_sentence_level_block_names_the_sentence_and_the_right_artifact(self):
        from report_workflow.nodes.qa_gate import format_blocked_factuality_hint

        hint = format_blocked_factuality_hint(
            {
                "claims": [
                    {"claim_id": "c1", "status": "verified"},
                    {
                        "sentence_id": "sent_21",
                        "status": "blocked",
                        "checker": "FD",
                        "reason": "Wording strength 'measured' is not allowed",
                    },
                ]
            },
            1,
        )
        self.assertIn("sent_21", hint)
        self.assertIn("FD", hint)
        self.assertIn("Wording strength", hint)
        self.assertIn("sentence_map.jsonl", hint)
        self.assertNotIn("(?)", hint)

    def test_claim_level_block_still_names_the_claim(self):
        from report_workflow.nodes.qa_gate import format_blocked_factuality_hint

        hint = format_blocked_factuality_hint(
            {"claims": [{"claim_id": "c7", "status": "blocked", "checker": "FA"}]}, 1
        )
        self.assertIn("c7", hint)
        self.assertIn("claim_matrix.json", hint)


class CaptionNumberingTest(unittest.TestCase):
    """Figures and tables carry independent sequences."""

    def _manifest(self):
        return {
            "figures": [
                {
                    "figure_id": "1",
                    "figure_type": "line",
                    "title": "月別不良率",
                    "path": "",
                    "render_mode": "native_table",
                    "data": {"columns": ["月份", "不良率"], "rows": [["2026-01", "2.50"]]},
                },
                {
                    "figure_id": "2",
                    "figure_type": "table",
                    "title": "月別明細",
                    "path": "",
                    "render_mode": "native_table",
                    "data": {"columns": ["月份", "投產數"], "rows": [["2026-01", "12480"]]},
                },
            ]
        }

    def test_two_tables_number_one_and_two(self):
        from report_workflow.nodes.docx_render import _replace_figure_placeholders

        md = (
            "# 主要發現\n\n不良率自一月的二點五零上升至六月的四點四零，"
            "轉折點落在六月底的模具更換，兩段之間沒有其他製程變更。\n\n"
            "[FIGURE:1]\n\n根因為第三模穴的模具磨耗，並非量測系統問題。\n\n"
            "[FIGURE:2]\n"
        )
        out, replaced, unresolved = _replace_figure_placeholders(md, self._manifest())
        self.assertEqual((replaced, unresolved), (2, []))
        self.assertIn("表 1. 月別不良率", out)
        self.assertIn("表 2. 月別明細", out)

    def test_image_figure_numbering_is_independent_of_figure_id(self):
        from report_workflow.nodes.docx_render import _figure_alt_text

        alt = _figure_alt_text(
            {"figure_id": "7", "title": "月別不良率"},
            "7",
            language="zh",
            display_number="1",
        )
        self.assertTrue(alt.startswith("圖 1."), alt)


class PromptFragmentScanTest(unittest.TestCase):
    """An image path is not publication text."""

    PROMPT = "分析第二產線今年一至九月不良率變化,向廠長說明主因與後續建議"

    def test_prompt_inside_an_image_path_is_not_a_leak(self):
        from report_workflow.nodes.docx_render import _pre_render_sanity_check

        md = (
            "# 主要發現\n\n不良率自一月的 2.50% 上升至六月的 4.40%。\n\n"
            f"![圖 1. 月別不良率](C:/out/{self.PROMPT}--run_1/figures/1.png)\n"
        )
        self.assertEqual(_pre_render_sanity_check(md, {}, [self.PROMPT]), [])

    def test_prompt_in_body_text_is_still_a_leak(self):
        from report_workflow.nodes.docx_render import _pre_render_sanity_check

        md = f"# 主要發現\n\n{self.PROMPT}\n"
        issues = _pre_render_sanity_check(md, {}, [self.PROMPT])
        # The message quotes the offending fragment so the author can find it.
        self.assertTrue(
            [i for i in issues if i.startswith("Raw prompt fragment leaked into publication text")],
            issues,
        )


class CjkTypographyLinkTargetTest(unittest.TestCase):
    """Typography normalization must not rewrite a file path.

    A run directory is named after the user's prompt, so its path routinely
    holds a space between two CJK words. Closing that gap renamed the figure's
    directory, pandoc found nothing there, and the report rendered silently
    without its chart.
    """

    IMAGE = "![](C:/tmp/板式熱交換器效能量測 實驗報告--run_1/figures/1.png)"

    def test_space_inside_an_image_path_survives(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        self.assertEqual(_normalize_cjk_typography(self.IMAGE), self.IMAGE)

    def test_space_inside_a_link_target_survives(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        md = "[說明](C:/報告 目錄/a.md)"
        self.assertEqual(_normalize_cjk_typography(md), md)

    def test_prose_around_a_link_is_still_normalized(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        out = _normalize_cjk_typography(f"轉動。 千分錶\n{self.IMAGE}")
        self.assertIn("轉動。千分錶", out)
        self.assertIn("量測 實驗報告", out)


class CjkTypographyTest(unittest.TestCase):
    def test_chinese_sentence_lines_join_without_space(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        md = "五級荷重下撓度皆隨荷重增加。\n5 N 時實測撓度 1.52 mm,誤差 4.8%。\n每一點皆高於理論值。"
        out = _normalize_cjk_typography(md)
        self.assertEqual(
            out,
            "五級荷重下撓度皆隨荷重增加。5 N 時實測撓度 1.52 mm,誤差 4.8%。每一點皆高於理論值。",
        )

    def test_space_before_citation_marker_closes(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        self.assertEqual(
            _normalize_cjk_typography("誤差 4.8%。 [1]"), "誤差 4.8%。[1]"
        )

    def test_gap_left_by_stripped_marker_closes(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        self.assertEqual(
            _normalize_cjk_typography("受載時會產生微小轉動。 千分錶量測時亦抵住樑面。"),
            "受載時會產生微小轉動。千分錶量測時亦抵住樑面。",
        )

    def test_chinese_latin_spacing_survives(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        text = "撓度自 5 N 的 1.52 mm 增至 25 N 的 7.49 mm。"
        self.assertEqual(_normalize_cjk_typography(text), text)

    def test_paragraph_break_and_headings_survive(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        md = "# 1. 結果與討論\n\n第一段結束。\n第一段續句。\n\n第二段開始。"
        out = _normalize_cjk_typography(md)
        self.assertIn("# 1. 結果與討論", out)
        self.assertIn("第一段結束。第一段續句。", out)
        self.assertIn("\n\n第二段開始。", out)

    def test_table_rows_are_not_joined(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        md = "表 1. 撓度比較\n\n| 荷重(N) | 實測(mm) |\n| --- | --- |\n| 5 | 1.52 |"
        out = _normalize_cjk_typography(md)
        self.assertIn("| 荷重(N) | 實測(mm) |\n| --- | --- |", out)

    def test_code_fence_untouched(self):
        from report_workflow.nodes.docx_render import _normalize_cjk_typography

        md = "說明如下。\n```\n第一行。\n第二行。\n```"
        out = _normalize_cjk_typography(md)
        self.assertIn("第一行。\n第二行。", out)


class DuplicateCitationCollapseTest(unittest.TestCase):
    def test_repeated_numeric_marker_collapses(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        text = "At 10 N and 15 N the deflections were 3.01 mm and 4.48 mm. [1] [1]"
        self.assertEqual(
            _collapse_adjacent_duplicate_citations(text),
            "At 10 N and 15 N the deflections were 3.01 mm and 4.48 mm. [1]",
        )

    def test_three_in_a_row_collapse(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        self.assertEqual(
            _collapse_adjacent_duplicate_citations("Summary. [2] [2] [2]"),
            "Summary. [2]",
        )

    def test_distinct_citations_are_kept(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        text = "Both sources agree. [1] [2]"
        self.assertEqual(_collapse_adjacent_duplicate_citations(text), text)

    def _rendered(self, tmpdir, paragraphs, table_rows=None):
        from docx import Document

        path = str(Path(tmpdir) / "out.docx")
        document = Document()
        for text in paragraphs:
            document.add_paragraph(text)
        if table_rows:
            table = document.add_table(rows=len(table_rows),
                                       cols=len(table_rows[0]))
            for r, row in enumerate(table_rows):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = cell
        document.save(path)
        return path

    def _state_with(self, tmpdir, paragraphs, manifest=None):
        import json

        from docx import Document

        from report_workflow.state import ReportState

        path = str(Path(tmpdir) / "final.docx")
        document = Document()
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)
        state = ReportState.new("report", [], str(Path(tmpdir) / "out"))
        state.spec["report_profile"] = "engineering_lab_report"
        state.output["final_docx_path"] = path
        if manifest is not None:
            manifest_path = str(Path(tmpdir) / "figure_manifest.json")
            with open(manifest_path, "w", encoding="utf-8") as handle:
                json.dump(manifest, handle)
            state.output["figure_manifest_path"] = manifest_path
        return state

    def _captions(self, caps):
        from report_workflow.errors import QAHardBlockError
        from report_workflow.nodes.post_render_validate import run_post_render_validate

        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._state_with(tmpdir, caps, manifest={"figures": [], "errors": []})
            try:
                run_post_render_validate(state)
            except QAHardBlockError as exc:
                return str(exc)
        return ""

    def test_a_chinese_reference_to_a_missing_figure_is_caught(self):
        """Three checks were built on an English-only pattern.

        A reference with no caption, references with no embedded figure at
        all, and references to figures the outline never declared — none had
        ever run against a Chinese report, where the prose says "由圖 3 可見".
        """
        self.assertIn("figure references without matching captions",
                      self._captions(["由圖 3 可見有效度上升。", "其他內容。"]))

    def test_no_space_between_the_word_and_the_number(self):
        """A word boundary does nothing next to CJK, so "圖3" must match too."""
        self.assertIn("figure references without matching captions",
                      self._captions(["由圖3可見有效度上升。", "其他內容。"]))

    def test_english_references_still_behave(self):
        self.assertIn("figure references without matching captions",
                      self._captions(["As Figure 3 shows, effectiveness rises.", "Other."]))

    def test_a_chinese_caption_alone_is_not_a_dangling_reference(self):
        self.assertEqual(self._captions(["圖 1. 流量與有效度", "說明文字。"]), "")

    def test_a_repeated_chinese_caption_is_caught(self):
        """The duplicate check matched "Figure 1." only.

        A Chinese report is the ordinary case here, and two captions both
        reading "圖 1." passed — the same English-shaped rule that dropped
        short CJK headings and called a finished Chinese report incomplete.
        """
        self.assertIn("duplicate figure caption",
                      self._captions(["圖 1. 流量與有效度", "圖 1. 流量與有效度"]))

    def test_simplified_characters_count_too(self):
        self.assertIn("duplicate figure caption",
                      self._captions(["图 1. 流量与有效度", "图 1. 流量与有效度"]))

    def test_english_captions_still_behave(self):
        self.assertIn("duplicate figure caption",
                      self._captions(["Figure 1. Flow", "Figure 1. Flow"]))

    def test_distinct_chinese_captions_pass(self):
        self.assertEqual(self._captions(["圖 1. 流量與有效度", "圖 2. 溫度影響"]), "")

    def test_a_raw_placeholder_never_reaches_the_deliverable(self):
        """The document a student hands in carried "[FIGURE:fig_x]" as text.

        The placeholder check lived inside the hint for a count mismatch, so
        when every figure failed to build the expected count fell to zero,
        matched the zero embedded, and publish reported "validation passed"
        over a document with the raw placeholder still in it.
        """
        from report_workflow.errors import QAHardBlockError
        from report_workflow.nodes.post_render_validate import run_post_render_validate

        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._state_with(
                tmpdir, ["量測數據", "[FIGURE:fig_effectiveness]"],
                manifest={"figures": [], "errors": []})
            with self.assertRaises(QAHardBlockError) as ctx:
                run_post_render_validate(state)
            self.assertIn("fig_effectiveness", str(ctx.exception))

    def test_the_build_error_is_carried_into_the_message(self):
        """FIGURE_BUILD records why and carries on; that reason was unread."""
        from report_workflow.errors import QAHardBlockError
        from report_workflow.nodes.post_render_validate import run_post_render_validate

        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._state_with(
                tmpdir, ["量測數據", "[FIGURE:fig_x]"],
                manifest={"figures": [], "errors": ["fig_x: data must be an object"]})
            with self.assertRaises(QAHardBlockError) as ctx:
                run_post_render_validate(state)
            self.assertIn("data must be an object", str(ctx.exception))

    def test_a_document_without_placeholders_passes(self):
        from report_workflow.nodes.post_render_validate import run_post_render_validate

        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._state_with(
                tmpdir, ["量測數據", "流量為 2.0 L/min 時，實測有效度為 72.4%。"],
                manifest={"figures": [], "errors": []})
            run_post_render_validate(state)

    def test_a_relative_figure_link_resolves_against_its_own_draft(self):
        """A relative link means relative to the file that contains it.

        Resolving against the run directory alone put the path somewhere the
        file was not, and pandoc then emitted the caption from the alt text
        with no image under it — the document went out saying "圖 1." and
        "由圖 1 可見" with nothing between them.
        """
        from report_workflow.nodes.docx_render import _absolutize_image_paths

        with tempfile.TemporaryDirectory() as tmpdir:
            draft_dir = Path(tmpdir) / "drafts"
            (draft_dir / "figures").mkdir(parents=True)
            (draft_dir / "figures" / "fig1.png").write_bytes(b"x")
            run_dir = Path(tmpdir) / "run"
            run_dir.mkdir()
            out = _absolutize_image_paths(
                "![圖 1](figures/fig1.png)", run_dir, draft_dir)
            self.assertIn("drafts/figures/fig1.png", out.replace("\\", "/"))

    def test_an_absolute_figure_link_is_left_alone(self):
        from report_workflow.nodes.docx_render import _absolutize_image_paths

        with tempfile.TemporaryDirectory() as tmpdir:
            original = f"![f]({Path(tmpdir).as_posix()}/a.png)"
            self.assertEqual(
                _absolutize_image_paths(original, Path(tmpdir), Path(tmpdir)),
                original)

    def test_a_figure_that_never_arrived_is_named(self):
        """The repair pass already learned the file was missing and threw the
        finding away, so the figure disappeared without a word."""
        from report_workflow.nodes.docx_render import _validate_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._rendered(tmpdir, ["標題", "圖 1. 有效度隨流量變化。"])
            issues = _validate_docx(
                path, "![圖 1](/nowhere/fig1.png)\n\n圖 1. 有效度隨流量變化。")
            self.assertTrue(any("do not exist" in i for i in issues), issues)

    def test_a_report_with_no_figures_is_not_nagged(self):
        from report_workflow.nodes.docx_render import _validate_docx

        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._rendered(tmpdir, ["標題", "本節沒有任何圖片。"])
            issues = _validate_docx(path, "# 標題\n\n本節沒有任何圖片。")
            self.assertEqual([i for i in issues if "figure" in i], [])

    def test_a_complete_chinese_report_is_not_called_incomplete(self):
        """The floor was 500 characters, which asks "long enough in English".

        Chinese says the same thing in far fewer characters, so a finished
        report was told it was likely incomplete while its English
        translation passed — the same threshold-tuned-in-English mistake as
        the block floor that discarded short CJK headings.
        """
        from report_workflow.nodes.docx_render import _validate_docx

        source = ("# 板式熱交換器效能量測報告\n\n## 1. 實驗目的\n\n"
                  "評估不同流量下的有效度，並與理論模型比較。\n\n"
                  "## 2. 結果\n\n以最小平方法擬合，斜率為 2.45。\n")
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._rendered(tmpdir, [
                "板式熱交換器效能量測報告", "1. 實驗目的",
                "評估不同流量下的有效度，並與理論模型比較。",
                "2. 結果", "以最小平方法擬合，斜率為 2.45。"])
            self.assertEqual(
                [i for i in _validate_docx(path, source) if "chars" in i], [])

    def test_losing_the_content_is_still_reported(self):
        from report_workflow.nodes.docx_render import _validate_docx

        source = "# 標題\n\n" + "本節說明流量與有效度之關係。" * 60
        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._rendered(tmpdir, ["標題", "本節說明流量與有效度之關係。"])
            issues = _validate_docx(path, source)
            self.assertTrue(any("lost in rendering" in i for i in issues), issues)

    def test_a_report_that_is_mostly_a_table_is_counted(self):
        """doc.paragraphs does not reach inside tables, so a results table
        counted as almost nothing."""
        from docx import Document

        from report_workflow.nodes.docx_render import _docx_text_length

        with tempfile.TemporaryDirectory() as tmpdir:
            path = self._rendered(
                tmpdir, ["量測結果"],
                [["流量 (L/min)", "實測有效度 (%)"], ["2.0", "72.4"], ["3.0", "76.1"]])
            document = Document(path)
            paragraphs_only = sum(len(p.text) for p in document.paragraphs)
            self.assertGreater(_docx_text_length(document), paragraphs_only)

    def test_markup_is_not_counted_as_prose(self):
        from report_workflow.nodes.docx_render import _prose_length

        self.assertEqual(_prose_length("# 標題\n\n| a | b |\n| --- | --- |\n"),
                         _prose_length("標題 a b"))

    def test_a_file_name_does_not_name_authors(self):
        """A reference list may not name people who do not exist.

        The tag was built by splitting the file name on underscores and
        joining the pieces as surnames, so "Shah2003_PlateExchangers.pdf"
        was cited as "Shah2003 & PlateExchangers" — a title fragment
        credited as a co-author — and "熱傳學_王小明_2019.pdf" became
        "熱傳學 et al.", the book's title standing in for the author, who
        was in the file name and was dropped.
        """
        from report_workflow.nodes.citation_bind import _format_apa_author_year

        self.assertEqual(
            _format_apa_author_year("Shah2003_PlateExchangers.pdf", "pdf"),
            "Shah2003 (n.d.)")
        self.assertEqual(
            _format_apa_author_year("熱傳學_王小明_2019.pdf", "pdf"), "熱傳學 (n.d.)")
        self.assertEqual(
            _format_apa_author_year("Kays_London_1984.pdf", "pdf"), "Kays (n.d.)")

    def test_a_single_word_file_name_is_unchanged(self):
        from report_workflow.nodes.citation_bind import _format_apa_author_year

        self.assertEqual(
            _format_apa_author_year("Incropera.pdf", "pdf"), "Incropera (n.d.)")

    def test_the_entry_shows_the_whole_file_name_in_the_title(self):
        """Nothing is lost: what the tag stops claiming, the title states."""
        from report_workflow.nodes.citation_bind import _format_apa_reference_entry

        entry = _format_apa_reference_entry("Kays_London_1984.pdf", "pdf", "sid")
        self.assertIn("*Kays_London_1984*", entry)
        self.assertNotIn("et al..", entry)
        self.assertNotIn("&", entry)

    def test_the_no_date_rule_is_kept(self):
        """The year in a file name is still not asserted as a publication date.

        That rule was learned from a fabricated bibliography entry reaching a
        rendered paper, and reading 1984 out of a file name would be the same
        kind of guess wearing a different hat.
        """
        from report_workflow.nodes.citation_bind import _format_apa_reference_entry

        self.assertIn(
            "(n.d.)", _format_apa_reference_entry("Kays_London_1984.pdf", "pdf", "s"))

    def test_author_year_and_source_markers(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        self.assertEqual(
            _collapse_adjacent_duplicate_citations("Claim. (Tsai, 2026) (Tsai, 2026)"),
            "Claim. (Tsai, 2026)",
        )
        self.assertEqual(
            _collapse_adjacent_duplicate_citations("Note. [Source: a.py] [Source: a.py]"),
            "Note. [Source: a.py]",
        )

    def test_markdown_links_untouched(self):
        from report_workflow.nodes.citation_bind import (
            _collapse_adjacent_duplicate_citations,
        )

        text = "See [the guide](https://example.com/2026) [the guide](https://example.com/2026)"
        self.assertEqual(_collapse_adjacent_duplicate_citations(text), text)


class ReaderRubricTest(unittest.TestCase):
    PROFILES = [
        "engineering_lab_report",
        "academic_paper",
        "proposal",
        "business_report",
        "admissions_report",
        "admissions_project_report",
        "custom",
    ]

    def test_every_profile_has_a_rubric(self):
        from report_workflow.nodes.agent_tasks import _reader_rubric_section

        for profile in self.PROFILES:
            section = _reader_rubric_section(profile)
            self.assertIn("How the Reader Grades This", section)
            self.assertIn("entry ticket", section)

    def test_lab_rubric_demands_quantified_comparison(self):
        from report_workflow.nodes.agent_tasks import _reader_rubric_section

        section = _reader_rubric_section("engineering_lab_report")
        self.assertIn("R²", section)
        self.assertIn("professor", section)


class StructureGuidanceTest(unittest.TestCase):
    def test_paragraph_rule_for_every_profile(self):
        from report_workflow.nodes.agent_tasks import _structure_guidance

        for profile in ReaderRubricTest.PROFILES:
            section = _structure_guidance(profile)
            self.assertIn("Context → Content → Conclusion", section)
            self.assertIn("Kording", section)

    def test_lab_discussion_recipe_present(self):
        from report_workflow.nodes.agent_tasks import _structure_guidance

        section = _structure_guidance("engineering_lab_report")
        self.assertIn("quantitatively with theory", section)
        self.assertIn("verdict", section)

    def test_business_lines_lead_with_the_answer(self):
        from report_workflow.nodes.agent_tasks import _structure_guidance

        self.assertIn("Pyramid Principle", _structure_guidance("business_report"))
        self.assertIn("SCQA", _structure_guidance("proposal"))


class EnglishTitleLocalizationTest(unittest.TestCase):
    def test_cjk_only_blueprint_title_falls_back_for_english(self):
        from report_workflow.language import localized_section_title

        self.assertEqual(
            localized_section_title({"title": "封面"}, "cover", "en"), "Cover"
        )
        self.assertEqual(
            localized_section_title({"title": "結果與討論"}, "results_discussion", "en"),
            "Results Discussion",
        )

    def test_zh_document_still_gets_chinese_title(self):
        from report_workflow.language import localized_section_title

        self.assertEqual(
            localized_section_title(
                {"title": "Cover", "title_zh": "封面"}, "cover", "zh"
            ),
            "封面",
        )

    def test_engineering_blueprint_is_bilingual(self):
        import report_workflow

        blueprint = (
            Path(report_workflow.__file__).parent
            / "blueprints"
            / "engineering_lab_report.yaml"
        ).read_text(encoding="utf-8")
        self.assertIn("title: Cover", blueprint)
        self.assertIn("title_zh: 封面", blueprint)
        self.assertIn("title: Results and Discussion", blueprint)
        self.assertIn("title_zh: 結果與討論", blueprint)


class NativeTableRenderTest(unittest.TestCase):
    def _manifest(self):
        return {
            "figures": [
                {
                    "figure_id": "1",
                    "figure_type": "table",
                    "title": "各級荷重之實測與理論撓度",
                    "render_mode": "native_table",
                    "data": {
                        "columns": ["荷重(N)", "實測撓度(mm)"],
                        "rows": [["5", "1.52"], ["10", "3.01"]],
                    },
                }
            ]
        }

    def test_zh_native_table_markdown(self):
        from report_workflow.nodes.docx_render import _replace_figure_placeholders

        md = "五級荷重的量測結果整理如下表,實測撓度與理論撓度的誤差均低於檢查門檻。\n\n[FIGURE:1]\n"
        out, replaced, unresolved = _replace_figure_placeholders(md, self._manifest())
        self.assertEqual(replaced, 1)
        self.assertEqual(unresolved, [])
        self.assertIn("表 1. 各級荷重之實測與理論撓度", out)
        self.assertIn("| 荷重(N) | 實測撓度(mm) |", out)
        self.assertIn("| 5 | 1.52 |", out)
        self.assertNotIn("[FIGURE:1]", out)

    def test_en_native_table_caption(self):
        from report_workflow.nodes.docx_render import _replace_figure_placeholders

        md = "The measured deflections are summarized in the table below.\n\n[FIGURE:1]\n"
        manifest = self._manifest()
        manifest["figures"][0]["title"] = "Deflection summary"
        out, replaced, _ = _replace_figure_placeholders(md, manifest)
        self.assertEqual(replaced, 1)
        self.assertIn("Table 1. Deflection summary", out)


class FigureCaptionLanguageTest(unittest.TestCase):
    def test_zh_caption_prefix(self):
        from report_workflow.nodes.docx_render import _figure_alt_text

        alt = _figure_alt_text(
            {"figure_id": "1", "title": "實測與理論撓度比較"}, "1", language="zh"
        )
        self.assertTrue(alt.startswith("圖 1."), alt)

    def test_en_caption_prefix_unchanged(self):
        from report_workflow.nodes.docx_render import _figure_alt_text

        alt = _figure_alt_text({"figure_id": "1", "title": "Load vs deflection"}, "1")
        self.assertTrue(alt.startswith("Figure 1."), alt)


class ProvenanceGradeTest(unittest.TestCase):
    def test_users_measured_csv_row_grades_high(self):
        from report_workflow.nodes.evidence_normalize import compute_provenance_score

        entry = {"file_type": "csv"}
        block = {
            "content": '{"荷重(N)": "5", "實測撓度(mm)": "1.52", "理論撓度(mm)": "1.45"}',
            "block_type": "csv_row",
        }
        self.assertGreaterEqual(compute_provenance_score(entry, block), 0.7)


class CjkFrontMatterParseTest(unittest.TestCase):
    def test_labelled_cjk_author(self):
        self.assertEqual(_parse_author_from_user_prompt("作者:王小明"), "王小明")
        self.assertEqual(_parse_author_from_user_prompt("Author: 蔡知均"), "蔡知均")

    def test_english_author_still_parses(self):
        self.assertEqual(_parse_author_from_user_prompt("Author: Jane Smith"), "Jane Smith")

    def test_cjk_title_and_affiliation(self):
        self.assertEqual(_parse_title_from_user_prompt("標題:拉伸試驗報告"), "拉伸試驗報告")
        self.assertEqual(
            _parse_affiliation_from_user_prompt("單位:國立成功大學機械工程學系"),
            "國立成功大學機械工程學系",
        )


class CustomReferenceDocResolutionTest(unittest.TestCase):
    def test_defaults_to_builtin_template(self):
        self.assertEqual(_resolve_reference_doc({}), _REFERENCE_DOC)
        self.assertEqual(_resolve_reference_doc({"reference_docx_path": "  "}), _REFERENCE_DOC)

    def test_valid_custom_template_wins(self):
        spec = {"reference_docx_path": str(_REFERENCE_DOC)}
        self.assertEqual(_resolve_reference_doc(spec), _REFERENCE_DOC)

    def test_missing_custom_template_hard_blocks(self):
        with self.assertRaises(QAHardBlockError):
            _resolve_reference_doc({"reference_docx_path": r"Z:\no\such\template.docx"})

    def test_error_messages(self):
        self.assertIn("not found", reference_docx_error(Path(r"Z:\no\such\template.docx")))
        with tempfile.TemporaryDirectory() as td:
            txt = Path(td) / "styles.txt"
            txt.write_text("not a docx", encoding="utf-8")
            self.assertIn("not a .docx", reference_docx_error(txt))
            fake = Path(td) / "fake.docx"
            fake.write_text("still not a zip", encoding="utf-8")
            self.assertIn("not a valid docx", reference_docx_error(fake))
        self.assertIsNone(reference_docx_error(_REFERENCE_DOC))


@unittest.skipIf(_find_pandoc() is None, "pandoc not installed")
class CustomReferenceDocRenderTest(unittest.TestCase):
    def test_output_follows_user_template_styles(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as td:
            custom = Path(td) / "corporate.docx"
            shutil.copy2(_REFERENCE_DOC, custom)
            doc = Document(str(custom))
            doc.styles["Heading 1"].font.name = "Arial Black"
            doc.save(str(custom))

            md = Path(td) / "doc.md"
            md.write_text("# Section\n\nBody text.\n", encoding="utf-8")
            out = Path(td) / "doc.docx"
            self.assertTrue(_render_via_pandoc(str(md), str(out), reference_doc=custom))
            with zipfile.ZipFile(out) as z:
                styles = z.read("word/styles.xml").decode("utf-8")
            self.assertIn("Arial Black", styles)


@unittest.skipIf(_find_pandoc() is None, "pandoc not installed")
class PandocFormatIntegrationTest(unittest.TestCase):
    def test_math_toc_and_footer_in_rendered_docx(self):
        with tempfile.TemporaryDirectory() as td:
            md = Path(td) / "doc.md"
            md.write_text(
                _inject_toc(
                    "Title page\n\n---\n\n# Methods\n\nInline $E = mc^2$ math.\n",
                    has_front_matter=True,
                ),
                encoding="utf-8",
            )
            out = Path(td) / "doc.docx"
            self.assertTrue(_render_via_pandoc(str(md), str(out)))
            with zipfile.ZipFile(out) as z:
                doc = z.read("word/document.xml").decode("utf-8")
                self.assertIn("<m:oMath", doc, "TeX math must render to OMML")
                self.assertIn("TOC", doc)
                self.assertTrue(
                    [n for n in z.namelist() if n.startswith("word/footer")],
                    "page-number footer must survive into rendered output",
                )


if __name__ == "__main__":
    unittest.main()
