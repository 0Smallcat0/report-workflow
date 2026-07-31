import json
import importlib.util
import re
import tempfile
import unittest
from unittest.mock import patch
from pathlib import Path

import yaml
from docx import Document

from report_workflow.agent_wrapper import (
    _normalize_source_files,
    lint_agent_artifacts,
    query_evidence,
    run_engineering_audit,
)
from report_workflow.nodes.artifacts import run_artifacts
from report_workflow.nodes.corpus_build import run_corpus_build
from report_workflow.nodes.citation_bind import audit_sentence_citations, resolve_citations_publication
from report_workflow.nodes.factuality_check import (
    _check_content_overlap,
    _extract_numbers_with_unit,
)
from report_workflow.nodes.front_matter_build import run_front_matter_build
from report_workflow.nodes.evidence_normalize import _determine_source_role
from report_workflow.nodes.post_render_validate import run_post_render_validate
from report_workflow.nodes.section_draft import run_section_draft
from report_workflow.artifact_contract import load_jsonl_without_contract
from report_workflow.nodes.visual_render_check import run_visual_render_check
from report_workflow.state import ReportState, WORKFLOW_RUNS_DIR


class SourceRoleContractTests(unittest.TestCase):
    def test_agent_wrapper_accepts_structured_source_roles(self):
        files, role_map = _normalize_source_files([
            {"path": "old_report.docx", "role": "base_document"},
            "measurements.csv:source_data",
        ])

        self.assertEqual(files, ["old_report.docx", "measurements.csv"])
        self.assertEqual(role_map["old_report.docx"], "base_document")
        self.assertEqual(role_map["measurements.csv"], "source_data")

    def test_the_same_file_attached_twice_counts_once(self):
        """A double drag-select used to triple every row of a CSV.

        Each copy became its own source with its own ids, and the QA gate
        counts evidence entries, so duplicates could carry a thin source set
        past a threshold it should not clear.
        """
        files, _roles = _normalize_source_files([
            "notes.md", "notes.md", "data.csv",
        ])
        self.assertEqual(files, ["notes.md", "data.csv"])

    def test_files_sharing_a_name_are_kept_apart(self):
        """Two paths may well be two different files; only an identical
        resolved path counts as a duplicate."""
        files, _roles = _normalize_source_files(["a/data.csv", "b/data.csv"])
        self.assertEqual(len(files), 2)

    def test_one_file_cannot_serve_two_roles(self):
        with self.assertRaises(ValueError) as ctx:
            _normalize_source_files([
                {"path": "report.docx", "role": "base_document"},
                {"path": "report.docx", "role": "source_data"},
            ])
        self.assertIn("two roles", str(ctx.exception))

    def test_chinese_literature_pdf_is_a_research_document(self):
        """The pdf/docx branch matched English keywords only.

        Both counts came out zero for a Chinese paper, the strict comparison
        failed, and it fell through to primary_source — so a Chinese journal
        article could not be recognised as literature at all, and the
        evidence-policy warning kept asking those users to attach the
        references they had already attached. The md/txt branch above had
        carried the Chinese tokens for a while; this branch, where literature
        actually arrives, had not.
        """
        role = _determine_source_role(
            {"file_name": "王小明2024_結垢.pdf", "file_type": "pdf"},
            {"content": "本論文發表於機械工程學報，第 60 卷，探討結垢對熱傳的影響。"},
        )
        self.assertEqual(role, "research_document")

    def test_chinese_filename_token_is_enough(self):
        role = _determine_source_role(
            {"file_name": "文獻回顧.pdf", "file_type": "pdf"},
            {"content": "任何內容。"},
        )
        self.assertEqual(role, "research_document")

    def test_a_chinese_handout_is_not_literature(self):
        """Instructional material is a primary source, not a paper. Widening
        the literature tokens must not swallow every Chinese document."""
        role = _determine_source_role(
            {"file_name": "handout.docx", "file_type": "docx"},
            {"content": "本講義說明量測的操作條件，實驗步驟與方法如下，結果記錄於下表。"},
        )
        self.assertEqual(role, "primary_source")

    def test_english_classification_is_unchanged(self):
        self.assertEqual(
            _determine_source_role(
                {"file_name": "smith2020.pdf", "file_type": "pdf"},
                {"content": "Smith et al. journal of heat transfer, doi: 10.1000/x"},
            ),
            "research_document",
        )
        self.assertEqual(
            _determine_source_role(
                {"file_name": "lab_data.docx", "file_type": "docx"},
                {"content": "method and result of the experiment with participant data"},
            ),
            "primary_source",
        )

    def test_source_data_markdown_notes_are_project_sources(self):
        role = _determine_source_role(
            {
                "artifact_role": "source_data",
                "file_name": "source_notes.md",
                "file_type": "md",
            },
            {"content": "Transcribed measurements from scanned lab handout."},
        )

        self.assertEqual(role, "internal_project_source")

    def test_corpus_build_resolves_roles_by_absolute_path(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            base = Path(tmpdir) / "base.md"
            base.write_text("# Report\n\nBody", encoding="utf-8")
            source = Path(tmpdir) / "source.txt"
            source.write_text("Measurement source.", encoding="utf-8")

            state = ReportState.new(
                "revise report",
                [str(base), str(source)],
                str(Path(tmpdir) / "out"),
            )
            state.spec["artifact_role_map"] = {
                str(base.resolve()): "base_document",
                str(source.resolve()): "source_data",
            }

            result = run_corpus_build(state)
            roles = {
                entry["file_name"]: entry["artifact_role"]
                for entry in result.sources["source_registry"]
            }

        self.assertEqual(roles["base.md"], "base_document")
        self.assertEqual(roles["source.txt"], "source_data")


class EvidenceQueryContractTests(unittest.TestCase):
    def test_query_evidence_ranks_by_query(self):
        state = ReportState.new("report", [], "out")
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        rows = [
            {
                "evidence_id": "E001",
                "content": "The voltage measurement stayed stable during the trial.",
            },
            {
                "evidence_id": "E002",
                "content": "A separate table describes unrelated attendance data.",
            },
        ]
        (run_dir / "evidence_ledger.jsonl").write_text(
            "".join(json.dumps(row) + "\n" for row in rows),
            encoding="utf-8",
        )

        result = query_evidence(state.job_id, query="voltage stable")

        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["entries"][0]["evidence_id"], "E001")
        self.assertEqual(result["total_matches"], 1)


class ChineseFactualityContractTests(unittest.TestCase):
    def test_numeric_overlap_matches_compact_and_chinese_units(self):
        self.assertIn(("5", "cm"), _extract_numbers_with_unit("length is 5cm"))

        reasons = _check_content_overlap(
            {"claim_text": "試片長度為 5 cm 且電壓量測穩定。"},
            {"content": "電壓量測結果穩定，試片長度為5公分。"},
        )

        self.assertEqual(reasons, [])

    def test_chinese_overlap_blocks_unrelated_evidence(self):
        reasons = _check_content_overlap(
            {"claim_text": "試片溫度快速上升。"},
            {"content": "電壓量測結果穩定，試片長度為5公分。"},
        )

        self.assertTrue(any("Chinese claim terms not in evidence" in reason for reason in reasons))


class StructuredDraftContractTests(unittest.TestCase):
    def test_structured_drafts_generate_markdown_and_sentence_map(self):
        state = ReportState.new("write report", [], "out")
        state.plan["blueprint"] = {
            "section_order": ["results"],
            "sections": {"results": {"required": True}},
        }
        state.plan["outline"] = {
            "sections": {
                "results": {
                    "section_id": "results",
                    "claim_ids": ["c1"],
                }
            }
        }
        state.plan["claim_matrix"] = {
            "claims": [{
                "claim_id": "c1",
                "claim_text": "The pilot program enrolled 42 participants.",
                "evidence_ids": ["E001"],
            }]
        }
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        (run_dir / "structured_drafts.json").write_text(json.dumps({
            "sections": {
                "results": {
                    "title": "Results",
                    "sentences": [{
                        "text": "The pilot program enrolled 42 participants.",
                        "claim_ids": ["c1"],
                        "evidence_ids": ["E001"],
                        "wording_strength": "hedged",
                    }],
                }
            }
        }), encoding="utf-8")

        result = run_section_draft(state)

        section_path = Path(result.drafts["section_drafts"]["results"])
        self.assertIn(
            "The pilot program enrolled 42 participants [CITE:E001].",
            section_path.read_text(encoding="utf-8"),
        )
        rows = load_jsonl_without_contract(result.drafts["sentence_map_path"])
        self.assertEqual(rows[0]["section_id"], "results")
        self.assertEqual(rows[0]["claim_ids"], ["c1"])
        self.assertEqual(rows[0]["evidence_ids"], ["E001"])
        self.assertEqual(rows[0]["citation_ids"], ["E001"])
        self.assertEqual(rows[0]["draft_origin"], "structured_draft")

    def test_structured_drafts_emit_separate_markers_for_multiple_citations(self):
        state = ReportState.new("write report", [], "out")
        state.plan["blueprint"] = {
            "section_order": ["results"],
            "sections": {"results": {"required": True}},
        }
        state.plan["outline"] = {
            "sections": {
                "results": {
                    "section_id": "results",
                    "claim_ids": ["c1"],
                }
            }
        }
        state.plan["claim_matrix"] = {
            "claims": [{
                "claim_id": "c1",
                "claim_text": "The table and notes both support the result.",
                "evidence_ids": ["E001", "E002"],
            }]
        }
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        (run_dir / "structured_drafts.json").write_text(json.dumps({
            "sections": {
                "results": {
                    "title": "Results",
                    "sentences": [{
                        "text": "The table and notes both support the result.",
                        "claim_ids": ["c1"],
                        "evidence_ids": ["E001", "E002"],
                    }],
                }
            }
        }), encoding="utf-8")

        result = run_section_draft(state)

        section_text = Path(result.drafts["section_drafts"]["results"]).read_text(encoding="utf-8")
        self.assertIn("[CITE:E001] [CITE:E002]", section_text)
        self.assertNotIn("[CITE:E001,E002]", section_text)
        rows = load_jsonl_without_contract(result.drafts["sentence_map_path"])
        self.assertEqual(rows[0]["citation_ids"], ["E001", "E002"])

    def test_citation_bind_accepts_comma_delimited_legacy_markers(self):
        evidence = [
            {
                "evidence_id": "E001",
                "source_id": "S1",
                "source_role": "internal_project_source",
                "source_file_name": "source_data.md",
            },
            {
                "evidence_id": "E002",
                "source_id": "S1",
                "source_role": "internal_project_source",
                "source_file_name": "source_data.md",
            },
        ]
        sentence_map = [{
            "sentence_id": "s1",
            "section_id": "results",
            "evidence_ids": ["E001", "E002"],
            "citation_ids": ["E001", "E002"],
        }]

        audit = audit_sentence_citations("Result [CITE:E001,E002].", sentence_map, evidence)
        resolved, resolved_audit, _, _ = resolve_citations_publication(
            "Result [CITE:E001,E002].",
            evidence,
            [],
        )

        self.assertEqual(audit, [])
        self.assertNotIn("[CITE:", resolved)
        self.assertTrue(all(item["resolved"] for item in resolved_audit))


class PostRenderLayoutManifestTests(unittest.TestCase):
    def test_post_render_validate_writes_layout_manifest_and_packages_it(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("render report", [], str(Path(tmpdir) / "out"))
            state.output["renderer_used"] = "python-docx (fallback)"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            docx_path = run_dir / "rendered_report.docx"
            doc = Document()
            doc.add_heading("Results", level=1)
            doc.add_paragraph("The rendered report contains final prose without internal markers.")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Metric"
            table.rows[0].cells[1].text = "Value"
            table.rows[1].cells[0].text = "Participants"
            table.rows[1].cells[1].text = "42"
            doc.save(docx_path)
            state.output["final_docx_path"] = str(docx_path)

            validated = run_post_render_validate(state)
            manifest_path = Path(validated.runtime["post_render_layout_manifest_path"])
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))

            self.assertEqual(manifest["status"], "passed")
            self.assertEqual(manifest["renderer_used"], "python-docx (fallback)")
            self.assertEqual(manifest["counts"]["tables"], 1)
            self.assertEqual(manifest["counts"]["paragraphs"], 2)
            self.assertGreater(manifest["docx"]["file_size_bytes"], 0)
            self.assertEqual(manifest["headings"][0]["text"], "Results")
            self.assertEqual(manifest["tables"][0]["rows"], 2)
            self.assertEqual(manifest["tables"][0]["first_row_preview"], ["Metric", "Value"])

            packaged = run_artifacts(validated)
            roles = {
                item["role"]
                for item in json.loads(
                    Path(packaged.output["artifacts_manifest_path"]).read_text(encoding="utf-8")
                )["files"]
            }
            self.assertIn("qa_post_render_layout_manifest", roles)


class FinalQASummaryContractTests(unittest.TestCase):
    def test_artifacts_write_final_qa_summary_and_package_markdown(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "engineering_lab_report"
            state.qa["qa_decision"] = "pass"
            state.qa["artifact_completeness_status"] = "pass"
            state.output["workflow_success"] = True
            state.output["renderer_used"] = "pandoc"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            docx_path = run_dir / "final.docx"
            docx_path.write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(docx_path)
            state.output["published_report_path"] = str(docx_path)

            qa_summary_path = run_dir / "qa_summary.json"
            qa_summary_path.write_text(json.dumps({
                "qa_decision": "pass",
                "artifact_completeness_status": "pass",
            }), encoding="utf-8")
            state.qa["qa_summary_path"] = str(qa_summary_path)

            factuality_path = run_dir / "factuality_report.json"
            factuality_path.write_text(json.dumps({
                "verified_count": 2,
                "blocked_count": 0,
                "claims": [{"claim_id": "C001", "status": "verified"}],
            }), encoding="utf-8")
            state.qa["factuality_report_path"] = str(factuality_path)

            engineering_path = run_dir / "engineering_audit_report.json"
            engineering_path.write_text(json.dumps({
                "status": "review_recommended",
                "warning_count": 1,
                "issue_count": 1,
                "info_count": 0,
                "measurement_count": 3,
                "table_evidence_count": 1,
                "calculation_count": 1,
                "issues": [{"severity": "warning", "check": "claim_table_value_support"}],
            }), encoding="utf-8")
            state.qa["engineering_audit_report_path"] = str(engineering_path)

            (run_dir / "artifact_lint_report.json").write_text(json.dumps({
                "status": "ok",
                "error_count": 0,
                "warning_count": 0,
                "issues": [],
            }), encoding="utf-8")

            post_render_path = run_dir / "post_render_validate_report.json"
            post_render_path.write_text(json.dumps({
                "status": "passed",
                "paragraph_count": 4,
                "table_count": 1,
                "inline_shape_count": 0,
                "issues": [],
            }), encoding="utf-8")
            state.runtime["post_render_validate_report_path"] = str(post_render_path)

            layout_path = run_dir / "post_render_layout_manifest.json"
            layout_path.write_text(json.dumps({
                "status": "passed",
                "counts": {
                    "paragraphs": 4,
                    "tables": 1,
                    "inline_shapes": 0,
                },
                "issues": [],
            }), encoding="utf-8")
            state.runtime["post_render_layout_manifest_path"] = str(layout_path)

            figure_visual_path = run_dir / "figure_visual_quality_report.json"
            figure_visual_path.write_text(json.dumps({
                "status": "review",
                "issue_count": 1,
                "issues": [{"type": "tick_label_overlap", "figure_id": "fig1"}],
                "figures": [{"figure_id": "fig1", "status": "review"}],
            }), encoding="utf-8")
            state.qa["figure_visual_quality_report_path"] = str(figure_visual_path)

            packaged = run_artifacts(state)

            summary_path = Path(packaged.qa["final_qa_summary_path"])
            markdown_path = Path(packaged.qa["final_qa_summary_md_path"])
            summary = json.loads(summary_path.read_text(encoding="utf-8"))
            manifest = json.loads(
                Path(packaged.output["artifacts_manifest_path"]).read_text(encoding="utf-8")
            )
            roles = {item["role"] for item in manifest["files"]}

            self.assertEqual(summary["overall_status"], "review")
            self.assertEqual(summary["factuality"]["verified_count"], 2)
            self.assertEqual(summary["engineering_audit"]["warning_count"], 1)
            self.assertEqual(summary["engineering_audit"]["table_evidence_count"], 1)
            self.assertEqual(summary["figure_visual_quality"]["issue_count"], 1)
            self.assertEqual(summary["render"]["table_count"], 1)
            self.assertTrue(markdown_path.exists())
            self.assertIn("Engineering audit: review", markdown_path.read_text(encoding="utf-8"))
            self.assertIn("Figure visual quality: review", markdown_path.read_text(encoding="utf-8"))
            self.assertIn("qa_final_qa_summary", roles)
            self.assertIn("qa_final_qa_summary_markdown", roles)
            self.assertIn("qa_figure_visual_quality_report", roles)
            self.assertTrue((Path(packaged.output["published_dir"]) / "qa" / "final_qa_summary.json").exists())
            self.assertTrue((Path(packaged.output["published_dir"]) / "qa" / "figure_visual_quality_report.json").exists())

    def test_skipped_visual_render_check_does_not_downgrade_a_clean_delivery(self):
        """An optional check that did not run is not a finding about the document.

        LibreOffice and Poppler are asked for nowhere in the install
        instructions, so almost every first run skips the visual render check.
        The note explaining why was appended to that report's `issues`, and the
        delivery summary concatenates those into its render-issue list -- so a
        report that passed every gate came back marked "review", citing two
        tools its reader had never been told to install.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "business_report"
            state.qa["qa_decision"] = "pass"
            state.qa["artifact_completeness_status"] = "pass"
            state.output["workflow_success"] = True
            state.output["renderer_used"] = "pandoc"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            docx_path = run_dir / "final.docx"
            docx_path.write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(docx_path)
            state.output["published_report_path"] = str(docx_path)

            qa_summary_path = run_dir / "qa_summary.json"
            qa_summary_path.write_text(json.dumps({
                "qa_decision": "pass",
                "artifact_completeness_status": "pass",
            }), encoding="utf-8")
            state.qa["qa_summary_path"] = str(qa_summary_path)

            factuality_path = run_dir / "factuality_report.json"
            factuality_path.write_text(json.dumps({
                "verified_count": 5,
                "blocked_count": 0,
                "claims": [],
            }), encoding="utf-8")
            state.qa["factuality_report_path"] = str(factuality_path)

            post_render_path = run_dir / "post_render_validate_report.json"
            post_render_path.write_text(json.dumps({
                "status": "passed",
                "paragraph_count": 14,
                "table_count": 1,
                "inline_shape_count": 1,
                "issues": [],
            }), encoding="utf-8")
            state.runtime["post_render_validate_report_path"] = str(post_render_path)

            layout_path = run_dir / "post_render_layout_manifest.json"
            layout_path.write_text(json.dumps({
                "status": "passed",
                "counts": {"paragraphs": 14, "tables": 1, "inline_shapes": 1},
                "issues": [],
            }), encoding="utf-8")
            state.runtime["post_render_layout_manifest_path"] = str(layout_path)

            visual_path = run_dir / "visual_render_check_report.json"
            visual_path.write_text(json.dumps({
                "status": "skipped",
                "issues": [],
                "skipped_reason": "LibreOffice soffice or Poppler pdftoppm not found",
                "pdf_path": "",
                "png_paths": [],
            }), encoding="utf-8")
            state.runtime["visual_render_check_report_path"] = str(visual_path)

            packaged = run_artifacts(state)
            summary = json.loads(
                Path(packaged.qa["final_qa_summary_path"]).read_text(encoding="utf-8")
            )
            markdown = Path(packaged.qa["final_qa_summary_md_path"]).read_text(encoding="utf-8")

            # render_status is what carried the skip into the verdict: any
            # render issue makes it "review", and "review" makes the delivery
            # need review. With nothing wrong with the render, it passes.
            self.assertEqual(summary["render"]["issues"], [])
            self.assertEqual(summary["render"]["status"], "pass")
            self.assertNotIn("## Render Issues", markdown)
            # Still said once, plainly, where it cannot be read as a defect.
            self.assertEqual(summary["render"]["visual_render_status"], "skipped")
            self.assertIn("Visual render check: skipped (LibreOffice", markdown)

    def test_delivery_summary_names_the_packaged_report_not_the_working_copy(self):
        """"Which file do I send?" must not be answered with an intermediate.

        FINAL_PUBLISH set published_report_path to the run directory's
        final.docx -- the very same value as final_docx_path -- and packaging
        never corrected it once report.docx existed inside published/. So the
        delivery summary, the packaged metadata, and the payload handed back to
        the agent all pointed outside the delivery package at a working copy.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "business_report"
            state.qa["qa_decision"] = "pass"
            state.output["workflow_success"] = True
            state.output["renderer_used"] = "pandoc"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            working_copy = run_dir / "final.docx"
            working_copy.write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(working_copy)
            # What FINAL_PUBLISH leaves behind: the working copy, under a name
            # that promises the published one.
            state.output["published_report_path"] = str(working_copy)

            packaged = run_artifacts(state)

            delivered = Path(packaged.output["published_dir"]) / "report.docx"
            self.assertTrue(delivered.exists())
            self.assertEqual(Path(packaged.output["published_report_path"]), delivered)
            self.assertEqual(Path(packaged.output["final_docx_path"]), working_copy)

            summary = json.loads(
                Path(packaged.qa["final_qa_summary_path"]).read_text(encoding="utf-8")
            )
            metadata = json.loads(
                (Path(packaged.output["published_dir"]) / "metadata.json").read_text(encoding="utf-8")
            )
            markdown = Path(packaged.qa["final_qa_summary_md_path"]).read_text(encoding="utf-8")

            self.assertEqual(Path(summary["report"]["published_report_path"]), delivered)
            self.assertEqual(Path(metadata["published_report_path"]), delivered)
            self.assertIn(f"- Report to send: {delivered}", markdown)

    def test_client_readable_note_shows_what_backs_each_claim(self):
        """The file named client-readable has to answer the client's question.

        The claim-to-source mapping was assembled and written only as JSON,
        while the note a person opens carried a QA verdict, two counts, and a
        sentence recommending that JSON. So the pack shipped the answer to
        "what backs this?" in the one format the reader it was written for
        cannot read.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "business_report"
            state.qa["qa_decision"] = "pass"
            state.output["workflow_success"] = True

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            docx_path = run_dir / "final.docx"
            docx_path.write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(docx_path)

            ledger_path = run_dir / "evidence_ledger.jsonl"
            ledger_path.write_text(json.dumps({
                "evidence_id": "ev_median",
                "source_id": "3f7754bd",
                "source_file_name": "pilot_results.csv",
                "evidence_grade": "high",
                "evidence_type": "quantitative",
                "content": "Structured workflow, 42 notes,\n20.0 median minutes per note",
            }) + "\n", encoding="utf-8")
            state.sources["evidence_ledger_path"] = str(ledger_path)

            state.plan["claim_matrix"] = {"claims": [
                {
                    "claim_id": "c_median",
                    "claim_text": "The structured workflow cut the median to 20.0 minutes per note.",
                    "claim_type": "statistical",
                    "evidence_ids": ["ev_median"],
                },
                {
                    "claim_id": "c_orphan",
                    "claim_text": "A claim nothing was registered for.",
                    "claim_type": "factual",
                    "evidence_ids": [],
                },
            ]}

            factuality_path = run_dir / "factuality_report.json"
            factuality_path.write_text(json.dumps({
                "verified_count": 1,
                "blocked_count": 0,
                "claims": [{"claim_id": "c_median", "status": "verified", "checker": "FA"}],
            }), encoding="utf-8")
            state.qa["factuality_report_path"] = str(factuality_path)

            packaged = run_artifacts(state)
            note = (
                Path(packaged.output["published_dir"]) / "traceability"
                / "client_readable_qa_note.md"
            ).read_text(encoding="utf-8")

            self.assertIn("The structured workflow cut the median to 20.0 minutes per note.", note)
            self.assertIn("pilot_results.csv", note)
            # The quoted source text, on one line -- a newline mid-preview used
            # to break out of the bullet it belongs to.
            self.assertIn("Structured workflow, 42 notes, 20.0 median minutes per note", note)
            self.assertIn("Status: verified", note)
            self.assertIn("No supporting evidence is recorded", note)

    def test_traceability_notes_follow_the_document_language(self):
        """The pack's own prose must not be in a language its reader did not ask for.

        A Chinese deliverable renders 目錄, 執行摘要, 表 1., 圖 1. -- every
        heading and caption localized. The three notes under traceability/ are
        read by the same person and were English scaffolding wrapped around
        Chinese quotes, so the only part the reader could follow was the part
        quoted out of their own files.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "business_report"
            state.qa["qa_decision"] = "pass"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            (run_dir / "final.docx").write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(run_dir / "final.docx")

            ledger_path = run_dir / "evidence_ledger.jsonl"
            ledger_path.write_text(json.dumps({
                "evidence_id": "ev_median",
                "source_file_name": "試辦結果.csv",
                "evidence_grade": "high",
                "content": "結構化流程，42 份，每份中位數 20.0 分鐘",
            }, ensure_ascii=False) + "\n", encoding="utf-8")
            state.sources["evidence_ledger_path"] = str(ledger_path)

            state.plan["claim_matrix"] = {"claims": [{
                "claim_id": "c_median",
                "claim_text": "結構化流程把每份作業的批改中位數從 28.0 分鐘降到 20.0 分鐘，退件率也同步下降。",
                "claim_type": "statistical",
                "evidence_ids": ["ev_median"],
            }]}

            packaged = run_artifacts(state)
            trace = Path(packaged.output["published_dir"]) / "traceability"
            note = (trace / "client_readable_qa_note.md").read_text(encoding="utf-8")
            coverage = (trace / "evidence_coverage_summary.md").read_text(encoding="utf-8")
            factuality = (trace / "factuality_summary.md").read_text(encoding="utf-8")

            self.assertIn("每一項主張的依據", note)
            self.assertIn("出自 `試辦結果.csv`", note)
            self.assertNotIn("What Backs Each Claim", note)
            # The neighbours in the same folder, read by the same person.
            self.assertIn("證據涵蓋摘要", coverage)
            self.assertIn("事實查核摘要", factuality)
            # The pipeline's own vocabulary stays put, so these files still line
            # up with the JSON sitting next to them.
            self.assertIn("high", coverage)

    def test_english_delivery_keeps_english_traceability_notes(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "business_report"
            state.qa["qa_decision"] = "pass"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            (run_dir / "final.docx").write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(run_dir / "final.docx")

            ledger_path = run_dir / "evidence_ledger.jsonl"
            ledger_path.write_text(json.dumps({
                "evidence_id": "ev_median",
                "source_file_name": "pilot_results.csv",
                "evidence_grade": "high",
                "content": "Structured workflow, 42 notes, 20.0 median minutes per note",
            }) + "\n", encoding="utf-8")
            state.sources["evidence_ledger_path"] = str(ledger_path)

            state.plan["claim_matrix"] = {"claims": [{
                "claim_id": "c_median",
                "claim_text": "The structured workflow cut the median to 20.0 minutes per note.",
                "claim_type": "statistical",
                "evidence_ids": ["ev_median"],
            }]}

            packaged = run_artifacts(state)
            note = (
                Path(packaged.output["published_dir"]) / "traceability"
                / "client_readable_qa_note.md"
            ).read_text(encoding="utf-8")

            self.assertIn("What Backs Each Claim", note)
            self.assertIn("From `pilot_results.csv`", note)

    def test_truncated_source_quote_says_it_was_truncated(self):
        """A cut quote must not read as the whole of what the source said."""
        from report_workflow.nodes.artifacts import _preview

        self.assertEqual(_preview("short enough"), "short enough")
        cut = _preview("word " * 200)
        self.assertTrue(cut.endswith(" [...]"))
        self.assertLess(len(cut), 260)

    def test_visual_render_check_reports_absent_tools_as_a_skip_not_an_issue(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish report", [], str(Path(tmpdir) / "out"))
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            docx_path = run_dir / "final.docx"
            docx_path.write_bytes(b"docx placeholder")
            state.output["final_docx_path"] = str(docx_path)

            with patch(
                "report_workflow.nodes.visual_render_check._find_executable",
                return_value=None,
            ):
                checked = run_visual_render_check(state)

            report = json.loads(
                Path(checked.runtime["visual_render_check_report_path"]).read_text(encoding="utf-8")
            )
            self.assertEqual(report["status"], "skipped")
            self.assertEqual(report["issues"], [])
            self.assertIn("LibreOffice", report["skipped_reason"])


class TemplateStyleMapContractTests(unittest.TestCase):
    def test_artifacts_write_template_style_map_from_reference_and_rendered_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish styled report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "academic_paper"
            state.spec["reference_template_mode"] = "style_reference"
            state.qa["qa_decision"] = "pass"
            state.qa["artifact_completeness_status"] = "pass"
            state.output["workflow_success"] = True
            state.output["renderer_used"] = "pandoc"

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            reference_path = run_dir / "reference.docx"
            reference = Document()
            reference.add_heading("Reference Heading", level=1)
            reference.add_paragraph("Reference body text.")
            reference.save(reference_path)

            rendered_path = run_dir / "rendered_report.docx"
            rendered = Document()
            rendered.add_heading("Rendered Heading", level=1)
            rendered.add_paragraph("Rendered body text.")
            rendered.save(rendered_path)

            state.output["reference_docx_path"] = str(reference_path)
            state.output["reference_docx_applied"] = True
            state.output["final_docx_path"] = str(rendered_path)
            state.output["published_report_path"] = str(rendered_path)

            packaged = run_artifacts(state)
            style_map_path = Path(packaged.output["template_style_map_path"])
            style_map = json.loads(style_map_path.read_text(encoding="utf-8"))
            manifest = json.loads(
                Path(packaged.output["artifacts_manifest_path"]).read_text(encoding="utf-8")
            )
            roles = {item["role"] for item in manifest["files"]}

            self.assertEqual(style_map["status"], "pass")
            self.assertTrue(style_map["reference_docx_applied"])
            self.assertEqual(style_map["reference_template_mode"], "style_reference")
            self.assertIn("Heading 1", style_map["rendered_docx"]["paragraph_style_counts"])
            self.assertIn("Heading 1", style_map["style_comparison"]["rendered_styles_defined_in_reference"])
            self.assertIn("qa_template_style_map", roles)
            self.assertIn("qa_template_style_map_markdown", roles)
            self.assertTrue((Path(packaged.output["published_dir"]) / "qa" / "template_style_map.md").exists())


class TemplateFieldFillContractTests(unittest.TestCase):
    def test_template_fields_are_rendered_in_front_matter(self):
        state = ReportState.new("write fixed template report", [], "out")
        state.spec["report_profile"] = "custom"
        state.spec["front_matter"] = {
            "title": "Capstone Lab Report",
            "template_fields": {
                "course_name": "Control Systems",
                "student_id": "S12345",
            },
        }
        state.plan["blueprint"] = {"front_matter": {}}

        result = run_front_matter_build(state)

        self.assertIn("**Course Name:** Control Systems", result.plan["front_matter_md"])
        self.assertIn("**Student Id:** S12345", result.plan["front_matter_md"])
        self.assertEqual(result.plan["front_matter"]["template_fields"]["course_name"], "Control Systems")

    def test_artifacts_write_template_field_fill_report(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("publish fixed template report", [], str(Path(tmpdir) / "out"))
            state.status = "completed"
            state.spec["report_profile"] = "engineering_lab_report"
            state.spec["reference_template_mode"] = "fixed_template"
            state.qa["qa_decision"] = "pass"
            state.qa["artifact_completeness_status"] = "pass"
            state.output["workflow_success"] = True
            state.output["renderer_used"] = "pandoc"
            state.plan["front_matter"] = {
                "title": "Capstone Lab Report",
                "author_block": "Example Student",
                "template_fields": {
                    "course_name": "Control Systems",
                    "student_id": "S12345",
                },
            }

            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            reference_path = run_dir / "reference.docx"
            reference = Document()
            reference.add_paragraph("[Course Name]")
            reference.add_paragraph("[Student Id]")
            reference.save(reference_path)

            rendered_path = run_dir / "rendered_report.docx"
            rendered = Document()
            rendered.add_heading("Capstone Lab Report", level=1)
            rendered.add_paragraph("Example Student")
            rendered.add_paragraph("Course Name: Control Systems")
            rendered.add_paragraph("Student Id: S12345")
            rendered.save(rendered_path)

            state.output["reference_docx_path"] = str(reference_path)
            state.output["reference_docx_applied"] = True
            state.output["final_docx_path"] = str(rendered_path)
            state.output["published_report_path"] = str(rendered_path)

            packaged = run_artifacts(state)
            report_path = Path(packaged.output["template_field_fill_report_path"])
            report = json.loads(report_path.read_text(encoding="utf-8"))
            final_summary = json.loads(Path(packaged.output["final_qa_summary_path"]).read_text(encoding="utf-8"))
            manifest = json.loads(
                Path(packaged.output["artifacts_manifest_path"]).read_text(encoding="utf-8")
            )
            roles = {item["role"] for item in manifest["files"]}
            statuses = {field["key"]: field["status"] for field in report["fields"]}

            self.assertEqual(report["status"], "pass")
            self.assertEqual(statuses["course_name"], "filled")
            self.assertEqual(statuses["student_id"], "filled")
            self.assertEqual(final_summary["template_field_fill"]["filled_count"], report["filled_count"])
            self.assertIn("qa_template_field_fill_report", roles)
            self.assertIn("qa_template_field_fill_report_markdown", roles)
            self.assertTrue((Path(packaged.output["published_dir"]) / "qa" / "template_field_fill_report.md").exists())


class ArtifactLintContractTests(unittest.TestCase):
    def test_lint_agent_artifacts_reports_json_paths_and_hints(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("lint report", [], str(Path(tmpdir) / "out"))
            state.plan["blueprint"] = {
                "section_order": ["results"],
                "sections": {"results": {"required": True}},
            }
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({"evidence_id": "E001", "content": "Supported measurement."}) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "The report cites an out-of-run evidence ID.",
                    "evidence_ids": ["E999"],
                }]
            }), encoding="utf-8")
            (run_dir / "outline.json").write_text(json.dumps({
                "sections": {"results": {"claim_ids": ["C404"]}}
            }), encoding="utf-8")
            state.checkpoint("AGENT_TASKS")

            result = lint_agent_artifacts(state.job_id)

            self.assertEqual(result["status"], "issues_found")
            self.assertGreater(result["error_count"], 0)
            report = json.loads(Path(result["report_path"]).read_text(encoding="utf-8"))
            json_paths = {issue["json_path"] for issue in report["issues"]}
            self.assertIn("$.claims[0].evidence_ids", json_paths)
            self.assertIn("$.sections.results.claim_ids", json_paths)
            self.assertTrue(all(issue["hint"] for issue in report["issues"]))

    def test_lint_agent_artifacts_accepts_structured_drafts(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("lint report", [], str(Path(tmpdir) / "out"))
            state.plan["blueprint"] = {
                "section_order": ["results"],
                "sections": {"results": {"required": True}},
            }
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({"evidence_id": "E001", "content": "The pilot enrolled 42 participants."}) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "The pilot enrolled 42 participants.",
                    "evidence_ids": ["E001"],
                }]
            }), encoding="utf-8")
            (run_dir / "outline.json").write_text(json.dumps({
                "sections": {"results": {"claim_ids": ["C001"]}}
            }), encoding="utf-8")
            (run_dir / "structured_drafts.json").write_text(json.dumps({
                "sections": {
                    "results": {
                        "sentences": [{
                            "text": "The pilot enrolled 42 participants.",
                            "claim_ids": ["C001"],
                            "evidence_ids": ["E001"],
                        }]
                    }
                }
            }), encoding="utf-8")
            state.checkpoint("AGENT_TASKS")

            result = lint_agent_artifacts(state.job_id)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["error_count"], 0)
            self.assertTrue(Path(result["report_path"]).exists())


class EngineeringAuditContractTests(unittest.TestCase):
    def test_engineering_audit_flags_unit_support_and_bad_calculation(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("engineering report", [], str(Path(tmpdir) / "out"))
            state.spec["report_profile"] = "engineering_lab_report"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({"evidence_id": "E001", "content": "Measured length was 5 cm."}) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "Measured length was 5 mm.",
                    "evidence_ids": ["E001"],
                }]
            }), encoding="utf-8")
            draft_dir = run_dir / "section_drafts"
            draft_dir.mkdir()
            (draft_dir / "calculations.md").write_text(
                "# Calculations\n\nThe reported length is 5 mm.\n\n5 + 5 = 11 cm\n",
                encoding="utf-8",
            )
            state.checkpoint("AGENT_TASKS")

            result = run_engineering_audit(state.job_id)

            self.assertEqual(result["status"], "review_recommended")
            checks = {issue["check"] for issue in result["issues"]}
            self.assertIn("claim_unit_support", checks)
            self.assertIn("calculation_result", checks)
            self.assertTrue(Path(result["report_path"]).exists())

    def test_engineering_audit_accepts_supported_measurement_and_calculation(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("engineering report", [], str(Path(tmpdir) / "out"))
            state.spec["report_profile"] = "engineering_lab_report"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({"evidence_id": "E001", "content": "Measured length was 5 cm."}) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "Measured length was 5 cm.",
                    "evidence_ids": ["E001"],
                }]
            }), encoding="utf-8")
            draft_dir = run_dir / "section_drafts"
            draft_dir.mkdir()
            (draft_dir / "calculations.md").write_text(
                "# Calculations\n\nThe measured length is 5 cm.\n\n5 + 5 = 10 cm\n",
                encoding="utf-8",
            )
            state.checkpoint("AGENT_TASKS")

            result = run_engineering_audit(state.job_id)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["warning_count"], 0)
            self.assertEqual(result["calculation_count"], 1)
            self.assertGreaterEqual(result["measurement_count"], 2)

    def test_engineering_audit_checks_claim_values_against_table_evidence(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("engineering report", [], str(Path(tmpdir) / "out"))
            state.spec["report_profile"] = "engineering_lab_report"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({
                    "evidence_id": "E001",
                    "source_file_name": "measurements.csv",
                    "file_type": "csv",
                    "granularity": "table_row",
                    "content": json.dumps({
                        "metric": "participants",
                        "value": "42",
                        "unit": "count",
                    }),
                    "table_data": [
                        ["metric", "value", "unit"],
                        ["participants", "42", "count"],
                    ],
                }) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "The participant table reports 41 participants.",
                    "evidence_ids": ["E001"],
                }]
            }), encoding="utf-8")
            draft_dir = run_dir / "section_drafts"
            draft_dir.mkdir()
            (draft_dir / "data.md").write_text("# Data\n\nThe table reports participants.\n", encoding="utf-8")
            state.checkpoint("AGENT_TASKS")

            result = run_engineering_audit(state.job_id)

            self.assertEqual(result["status"], "review_recommended")
            self.assertEqual(result["table_evidence_count"], 1)
            checks = {issue["check"] for issue in result["issues"]}
            self.assertIn("claim_table_value_support", checks)
            report = json.loads(Path(result["report_path"]).read_text(encoding="utf-8"))
            self.assertEqual(report["tables"][0]["row_count"], 2)

    def test_engineering_audit_accepts_matching_table_claim_value(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("engineering report", [], str(Path(tmpdir) / "out"))
            state.spec["report_profile"] = "engineering_lab_report"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                json.dumps({
                    "evidence_id": "E001",
                    "source_file_name": "measurements.csv",
                    "file_type": "csv",
                    "granularity": "table_row",
                    "content": json.dumps({"metric": "participants", "value": "42"}),
                    "table_data": [["metric", "value"], ["participants", "42"]],
                }) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "The participant table reports 42 participants.",
                    "evidence_ids": ["E001"],
                }]
            }), encoding="utf-8")
            draft_dir = run_dir / "section_drafts"
            draft_dir.mkdir()
            (draft_dir / "data.md").write_text("# Data\n\nParticipant count is listed in the source table.\n", encoding="utf-8")
            state.checkpoint("AGENT_TASKS")

            result = run_engineering_audit(state.job_id)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["table_evidence_count"], 1)
            self.assertNotIn("claim_table_value_support", {issue["check"] for issue in result["issues"]})

    def test_engineering_audit_accepts_rounded_table_claim_values_and_page_numbers(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("engineering report", [], str(Path(tmpdir) / "out"))
            state.spec["report_profile"] = "engineering_lab_report"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            evidence_path = run_dir / "evidence_ledger.jsonl"
            evidence_path.write_text(
                "\n".join([
                    json.dumps({
                        "evidence_id": "E001",
                        "source_file_name": "measurements.md",
                        "file_type": "md",
                        "content": "PDF page 12 records airflow 6.0 m/s.",
                    }),
                    json.dumps({
                        "evidence_id": "E002",
                        "source_file_name": "measurements.md",
                        "file_type": "md",
                        "content": "| metric | value |\n|---|---:|\n| COP initial | 1.587793 |\n| COP final | 5.722860 |",
                    }),
                ]) + "\n",
                encoding="utf-8",
            )
            state.sources["evidence_ledger_path"] = str(evidence_path)
            (run_dir / "claim_matrix.json").write_text(json.dumps({
                "claims": [{
                    "claim_id": "C001",
                    "claim_text": "PDF page 12 shows airflow 6.0 m/s and COP rises from 1.588 to 5.723.",
                    "evidence_ids": ["E001", "E002"],
                }]
            }), encoding="utf-8")
            draft_dir = run_dir / "section_drafts"
            draft_dir.mkdir()
            (draft_dir / "data.md").write_text(
                "# Data\n\nPDF page 12 shows airflow 6.0 m/s and rounded COP values.\n",
                encoding="utf-8",
            )
            state.checkpoint("AGENT_TASKS")

            result = run_engineering_audit(state.job_id)

            self.assertNotIn("claim_table_value_support", {issue["check"] for issue in result["issues"]})
            self.assertNotIn("number_without_unit", {issue["check"] for issue in result["issues"]})


class DocumentationContractTests(unittest.TestCase):
    def test_agents_doc_does_not_show_removed_family_cli(self):
        text = Path("AGENTS.md").read_text(encoding="utf-8")

        self.assertNotIn("--family", text)
        self.assertNotIn("state.spec." + "report_family", text)
        legacy_triplet = "|".join(["academic_report", "work_report", "hybrid_report"])
        self.assertNotIn(legacy_triplet, text)

    def test_every_dependency_is_consistent(self):
        """The two lists of the same fact must not drift.

        This named one package and checked it appeared in both files, so
        openpyxl could be added to pyproject and left out of requirements.txt —
        the file the README tells people to install, so the documented setup
        produced an environment that accepts a .xlsx source and then cannot
        open it. Compare all of them, not a favourite.
        """
        import tomllib

        with open("pyproject.toml", "rb") as handle:
            declared = tomllib.load(handle)["project"]["dependencies"]
        listed = [
            line.strip()
            for line in Path("requirements.txt").read_text(encoding="utf-8").splitlines()
            if line.strip() and not line.strip().startswith("#")
        ]
        self.assertEqual(sorted(listed), sorted(declared))

    def test_claim_role_rule_in_the_brief_matches_the_gate(self):
        """The brief has to state this run's rule, not a profile family's.

        It read "**Academic reports**: every claim MUST have a claim_role
        field", while the gate fires on the run's own policy. Four profiles
        enforce it -- engineering_lab_report, academic_paper, admissions_report,
        custom -- and only one of them is an academic paper. So the agent
        writing a lab report, the profile the README leads with, read a rule
        addressed to somebody else, left the field out, and was hard-blocked at
        CLAIM_PLAN by a message that repeated the same mislabel.
        """
        from report_workflow.nodes.agent_tasks import _claim_role_rule
        from report_workflow.policies import get_policy
        from report_workflow.profiles import PROFILE_IDS

        for profile in PROFILE_IDS:
            rule = _claim_role_rule(profile)
            with self.subTest(profile=profile):
                self.assertIn(profile, rule)
                self.assertNotIn("Academic reports", rule)
                if get_policy(profile).claim.role_validation_required:
                    self.assertIn("required", rule)
                    self.assertIn("at most 3", rule)
                else:
                    self.assertIn("optional", rule)

    def test_abstract_guidance_states_this_profile_s_contract(self):
        """The abstract block was a lecture with carve-outs, not a contract.

        It named academic_paper as needing structured headings and "admissions
        and project profiles" as allowed a paragraph, then gave everybody
        "150-250 words". business_report and proposal have no abstract section
        at all and were still told how to write one and which gate would block
        it; of the five that do have one, the range was wrong for three, and
        engineering_lab_report and custom were never told which shape they may
        use even though both accept a plain paragraph.
        """
        from report_workflow.nodes.agent_tasks import _abstract_section
        from report_workflow.policies import get_policy
        from report_workflow.profiles import PROFILE_IDS

        for profile in PROFILE_IDS:
            blueprint = yaml.safe_load(
                Path(f"src/report_workflow/blueprints/{profile}.yaml").read_text(encoding="utf-8")
            )
            section = _abstract_section(profile, blueprint, "en")
            policy = get_policy(profile).abstract
            with self.subTest(profile=profile):
                self.assertNotIn("150-250 words total unless", section)
                if "abstract" not in blueprint.get("sections", {}):
                    self.assertIn("has no abstract section", section)
                    continue
                self.assertIn(str(policy.word_count_max), section)
                if policy.word_count_min:
                    self.assertIn(str(policy.word_count_min), section)
                if policy.structure_required:
                    self.assertIn("**Background:**", section)
                    self.assertIn("hard-blocks at METADATA_GATE", section)
                elif policy.allow_plain_paragraph:
                    self.assertIn("accepts either shape", section)

    def test_abstract_budget_follows_the_cjk_scale(self):
        from report_workflow.nodes.abstract_check import CJK_ABSTRACT_SCALE
        from report_workflow.nodes.agent_tasks import _abstract_section
        from report_workflow.policies import get_policy

        blueprint = yaml.safe_load(
            Path("src/report_workflow/blueprints/engineering_lab_report.yaml").read_text(encoding="utf-8")
        )
        cap = get_policy("engineering_lab_report").abstract.word_count_max
        zh = _abstract_section("engineering_lab_report", blueprint, "zh")
        self.assertIn(f"at most {cap * CJK_ABSTRACT_SCALE} characters", zh)
        self.assertIn(f"at most {cap} words", _abstract_section("engineering_lab_report", blueprint, "en"))

    def test_results_mode_guidance_points_where_the_gate_reads(self):
        """The brief sent authors to a key nothing reads.

        It said "include it in outline.json at the top level" and showed it
        there in the JSON shape, while QA_GATE reads
        outline["sections"]["results"]["results_mode"]. Measured: an outline
        with a top-level architectural_characterization produces the same
        verdict as one that sets nothing at all, so the choice was dropped in
        silence and the blueprint default stood in for it.
        """
        from report_workflow.nodes.agent_tasks import _results_mode_rule, _results_mode_section
        from report_workflow.policies import get_policy
        from report_workflow.profiles import PROFILE_IDS

        for profile in PROFILE_IDS:
            section = _results_mode_section(profile)
            rule = _results_mode_rule(profile)
            with self.subTest(profile=profile):
                self.assertNotIn("at the top level of outline.json:", section)
                if get_policy(profile).results.empirical_strict:
                    self.assertIn("sections.results.results_mode", section)
                    self.assertIn("sections.results.results_mode", rule)
                    self.assertIn("empirical", section)
                else:
                    self.assertIn("not used by", section.lower())
                    self.assertIn("not read", rule)

    def test_top_level_results_mode_is_refused_instead_of_ignored(self):
        from report_workflow.errors import QAHardBlockError
        from report_workflow.nodes.qa_gate import _results_section_reasons

        # The measurement this guard exists for: top level behaves as unset.
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("x", [], tmpdir)
            state.spec["report_profile"] = "academic_paper"
            merged = Path(tmpdir) / "merged.md"
            merged.write_text(
                "# Results\n\nThe structured workflow is faster and reduces effort.\n",
                encoding="utf-8",
            )
            state.drafts["merged_draft_md"] = str(merged)
            state.plan["blueprint"] = {"sections": {"results": {}}}

            state.plan["outline"] = {
                "results_mode": "architectural_characterization",
                "sections": {"results": {"section_id": "results"}},
            }
            self.assertTrue(_results_section_reasons(state))

            state.plan["outline"] = {
                "sections": {
                    "results": {
                        "section_id": "results",
                        "results_mode": "architectural_characterization",
                    }
                }
            }
            self.assertFalse(_results_section_reasons(state))

        # So OUTLINE_PLAN refuses the stranded key rather than dropping it.
        from report_workflow.nodes.outline_plan import run_outline_plan

        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("x", [], tmpdir)
            state.spec["report_profile"] = "academic_paper"
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            (run_dir / "outline.json").write_text(json.dumps({
                "results_mode": "empirical",
                "sections": {"results": {"section_id": "results", "claim_ids": []}},
            }), encoding="utf-8")
            with self.assertRaises(QAHardBlockError) as caught:
                run_outline_plan(state)
            self.assertIn("sections.results.results_mode", str(caught.exception))

    def test_claim_role_block_names_the_profile_that_blocked(self):
        from report_workflow.errors import QAHardBlockError
        from report_workflow.nodes.claim_plan import _validate_claim_matrix

        with self.assertRaises(QAHardBlockError) as caught:
            _validate_claim_matrix(
                {"claims": [{"claim_id": "c1", "claim_text": "x", "evidence_ids": ["e1"]}]},
                "engineering_lab_report",
            )
        message = str(caught.exception)
        self.assertIn("engineering_lab_report requires claim_role", message)
        self.assertNotIn("Academic reports", message)

    def test_readme_test_count_badge_matches_the_suite(self):
        """A number on the landing page that nothing checks will rot.

        The badge said 496 while the suite had grown past 750 -- a claim about
        the product, made to every visitor, that no gate applied to. This
        repository exists so that unverifiable claims do not reach a reader;
        the one on its own front page gets the same treatment.
        """
        badge = re.search(
            r"badge/tests-(\d+)%20passing",
            Path("README.md").read_text(encoding="utf-8"),
        )
        self.assertIsNotNone(badge, "README no longer carries a test-count badge")
        counted = sum(
            len(re.findall(r"^\s+def test_\w+", path.read_text(encoding="utf-8"), re.M))
            for path in sorted(Path("tests").glob("test_*.py"))
        )
        self.assertEqual(counted, int(badge.group(1)))

    def test_short_skill_documents_yaml_tool_surface(self):
        skill_text = Path("agent_skill/SKILL.md").read_text(encoding="utf-8")
        skill_yaml = yaml.safe_load(Path("agent_skill/skill.yaml").read_text(encoding="utf-8"))

        tool_names = [tool["name"] for tool in skill_yaml["tools"]]
        for tool_name in tool_names:
            with self.subTest(tool=tool_name):
                self.assertIn(f"`{tool_name}`", skill_text)

    def test_short_skill_points_to_reference_library(self):
        skill_text = Path("agent_skill/SKILL.md").read_text(encoding="utf-8")

        # SKILL.md stays a lean hub that links every reference file one level deep.
        skill_pointers = [
            "## Reference Library",
            "reference/setup-and-preflight.md",
            "reference/profiles.md",
            "reference/tools.md",
            "reference/authoring.md",
            "reference/figures.md",
            "reference/engineering-lab.md",
            "reference/revision.md",
            "reference/benchmarking.md",
            "python scripts/render_skill_docs.py --check",
            "python scripts/sync_codex_skill.py --write",
        ]
        for pointer in skill_pointers:
            with self.subTest(pointer=pointer):
                self.assertIn(pointer, skill_text)

    def test_reference_files_document_operational_guardrails(self):
        reference = Path("agent_skill/reference")

        # Detailed guardrails live in one-level-deep reference files, not SKILL.md.
        expected = {
            "setup-and-preflight.md": [
                "## Preflight Decision Examples",
                "allow_degraded_render=True",
                '"pandoc": "accept_degraded"',
                "enable_research=True",
                'notebooklm_notebook_id="notebook-id-from-user"',
            ],
            "authoring.md": [
                "`structured_drafts.json`",
                "remap_agent_artifacts(job_id=..., previous_job_id=...)",
                "Validation Failure Repair",
            ],
            "revision.md": [
                "`preview_revision_diff`",
                "`submit_revision_plan`",
            ],
            "engineering-lab.md": [
                "Chinese Engineering Publish Checklist",
                "`template_field_fill_report_path`",
                "`final_qa_summary_path`",
            ],
        }
        for file_name, phrases in expected.items():
            text = (reference / file_name).read_text(encoding="utf-8")
            for phrase in phrases:
                with self.subTest(file=file_name, phrase=phrase):
                    self.assertIn(phrase, text)

    def test_generated_skill_doc_blocks_are_current(self):
        script_path = Path("scripts/render_skill_docs.py")
        spec = importlib.util.spec_from_file_location("render_skill_docs", script_path)
        module = importlib.util.module_from_spec(spec)
        assert spec.loader is not None
        spec.loader.exec_module(module)

        results = module.render_docs(Path("."), write=False)
        self.assertTrue(results)
        for result in results:
            with self.subTest(path=str(result["path"])):
                self.assertFalse(result["changed"])

    def test_sync_codex_skill_script_copies_skill_files(self):
        script_path = Path("scripts/sync_codex_skill.py")
        spec = importlib.util.spec_from_file_location("sync_codex_skill", script_path)
        module = importlib.util.module_from_spec(spec)
        assert spec.loader is not None
        spec.loader.exec_module(module)

        with tempfile.TemporaryDirectory() as tmpdir:
            dest_dir = Path(tmpdir) / "report-workflow"

            dry_run_ops = module.sync_skill(Path("agent_skill"), dest_dir, write=False)
            # SKILL.md + skill.yaml + every reference/ file.
            self.assertGreaterEqual(len(dry_run_ops), 3)
            self.assertFalse(dest_dir.exists())

            write_ops = module.sync_skill(Path("agent_skill"), dest_dir, write=True)
            self.assertEqual(len(write_ops), len(dry_run_ops))
            for source, dest in write_ops:
                self.assertTrue(dest.exists())
                self.assertEqual(
                    source.read_text(encoding="utf-8"),
                    dest.read_text(encoding="utf-8"),
                )

            synced_names = {dest.name for _, dest in write_ops}
            self.assertIn("SKILL.md", synced_names)
            self.assertIn("skill.yaml", synced_names)
            self.assertIn("tools.md", synced_names)
            self.assertTrue((dest_dir / "reference" / "tools.md").exists())
            self.assertFalse((dest_dir / "agent_instructions.md").exists())


class DeliverableNamingTests(unittest.TestCase):
    """What the author actually collects, pinned by name.

    Packaging copies the rendered document into the published directory as
    report.docx and the draft beside it as report.md. Nothing tested that,
    and the names are the whole interface between this tool and the person
    submitting the work — renaming them would break every workflow built on
    them without a single test going red.
    """

    def _package(self, tmpdir, *, with_docx=True):
        from report_workflow.nodes.artifacts import run_artifacts
        from report_workflow.state import ReportState, published_dir_for, run_dir_for

        state = ReportState.new("報告", [], str(Path(tmpdir) / "out"))
        state.spec["report_profile"] = "engineering_lab_report"
        run_dir = run_dir_for(state)
        run_dir.mkdir(parents=True, exist_ok=True)
        draft = run_dir / "merged_draft.md"
        draft.write_text("# 報告\n\n內容。\n", encoding="utf-8")
        state.drafts["merged_draft_md"] = str(draft)
        if with_docx:
            docx = run_dir / "rendered_report.docx"
            docx.write_bytes(b"PK\x03\x04placeholder")
            state.output["final_docx_path"] = str(docx)
        run_artifacts(state)
        return published_dir_for(state)

    def test_the_document_is_collected_as_report_docx(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            published = self._package(tmpdir)
            self.assertTrue((published / "report.docx").exists())
            self.assertTrue((published / "report.md").exists())

    def test_a_missing_render_leaves_no_document_behind(self):
        """Documented, not fixed: the copy skips a missing source silently.

        RENDER_QA hard-blocks a missing DOCX before packaging runs, so this
        is a latent hazard rather than an observed defect — recorded here so
        the silence is at least visible to the next reader.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            published = self._package(tmpdir, with_docx=False)
            self.assertFalse((published / "report.docx").exists())
            self.assertTrue((published / "report.md").exists())


class ProfilesStayDistinctTests(unittest.TestCase):
    """Two profiles over one source must not collapse into one shape.

    Run the same CSV as an engineering lab report and as a business report:
    thirteen sections against four, APA citations against none. The parts
    that differ are the parts a reader sees. What must not differ is the
    traceability underneath — the internal trace map is built before the
    publication citation layer and is not gated on citation style, so a
    business report with no visible bibliography still carries a full
    claim-to-evidence trace. That is the guarantee this tool exists for, and
    it would be quiet to lose.
    """

    def test_the_two_profiles_do_not_share_a_citation_policy(self):
        from report_workflow.policies.policy_pack import get_policy

        lab = get_policy("engineering_lab_report").citation
        business = get_policy("business_report").citation
        self.assertEqual(lab.style, "APA")
        self.assertEqual(business.style, "none")
        self.assertTrue(lab.source_marker_hard_block)
        self.assertFalse(business.source_marker_hard_block)

    def test_a_business_report_needs_no_reference_section(self):
        """Its blueprint has none, and with style 'none' that is correct
        rather than an omission — a business report is not a paper."""
        from report_workflow.policies.policy_pack import get_policy

        self.assertEqual(get_policy("business_report").citation.style, "none")

    def test_the_trace_map_is_not_written_by_the_citation_layer(self):
        """Layer 1 runs before Layer 2 and knows nothing about style.

        Pinned by position: the internal trace map must be assigned before
        publication citation resolution begins, so no future style branch
        can end up deciding whether a report is traceable at all.
        """
        import inspect

        from report_workflow.nodes import citation_bind

        source = inspect.getsource(citation_bind.run_citation_bind)
        trace_at = source.index('state.citations["internal_trace_path"]')
        publication_at = source.index("Layer 2")
        self.assertLess(trace_at, publication_at)


if __name__ == "__main__":
    unittest.main()
