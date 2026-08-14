import json
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from report_workflow.nodes.base_document_parse import _parse_markdown_sections
from report_workflow.nodes.docx_render import _style_tables_post_render
from report_workflow.nodes.style_pass import _apply_style_polish
from report_workflow.nodes.admissions_monograph_polish import polish_admissions_monograph
from report_workflow.nodes.reference_reality_check import _classify_reference
from report_workflow.nodes.front_matter_build import _build_front_matter
from report_workflow.nodes.front_matter_build import run_front_matter_build
from report_workflow.nodes.final_publish import run_final_publish
from report_workflow.profiles import infer_report_profile, normalize_profile_id
from report_workflow.nodes.reference_verify import _is_publication_reference_candidate
from report_workflow.nodes.project_identity_gate import run_project_identity_gate
from report_workflow.nodes.figure_quality import run_figure_quality
from report_workflow.nodes.admissions_tone_gate import run_admissions_tone_gate
from report_workflow.nodes.reference_relevance_gate import run_reference_relevance_gate
from report_workflow.nodes.claim_plan import run_claim_plan
from report_workflow.nodes.evidence_normalize import run_evidence_normalize
from report_workflow.nodes.revision_apply import run_revision_apply
from report_workflow.agent_wrapper import submit_and_publish_report
from report_workflow.artifact_contract import (
    find_repo_hygiene_issues,
    load_artifact_contract,
    remap_evidence_ids,
    stable_evidence_id,
    validate_evidence_ledger_provenance,
    write_base_document_integrity,
)
from report_workflow.errors import QAHardBlockError
from report_workflow.state import ReportState, WORKFLOW_RUNS_DIR


# The XML Word writes for an equation. python-docx cannot build one, so the
# fixtures below paste it in as Word would have: inline within a sentence, and
# set on its own line.
_MATH_NAMESPACES = (
    'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
)
_INLINE_OMML = f"""<m:oMath {_MATH_NAMESPACES}>
  <m:r><m:t>Re</m:t></m:r>
  <m:r><m:t>=</m:t></m:r>
  <m:f>
    <m:num><m:r><m:t>&#961;VD</m:t></m:r></m:num>
    <m:den><m:r><m:t>&#956;</m:t></m:r></m:den>
  </m:f>
</m:oMath>"""
_DISPLAY_OMML = f"""<m:oMathPara {_MATH_NAMESPACES}>
  <m:oMath>
    <m:r><m:t>C_D</m:t></m:r>
    <m:r><m:t>=</m:t></m:r>
    <m:f>
      <m:num><m:r><m:t>2F</m:t></m:r></m:num>
      <m:den>
        <m:r><m:t>&#961;</m:t></m:r>
        <m:sSup>
          <m:e><m:r><m:t>U</m:t></m:r></m:e>
          <m:sup><m:r><m:t>2</m:t></m:r></m:sup>
        </m:sSup>
        <m:r><m:t>A</m:t></m:r>
      </m:den>
    </m:f>
  </m:oMath>
</m:oMathPara>"""


# A text box, written the way Word writes one: the modern shape under
# mc:Choice and the same words again under mc:Fallback for older readers.
_TEXTBOX_NAMESPACES = " ".join([
    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
    'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"',
    'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"',
    'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"',
    'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"',
    'xmlns:v="urn:schemas-microsoft-com:vml"',
])
_TEXTBOX_BODY = (
    "<w:p><w:r><w:t>國立成功大學機械工程學系</w:t></w:r></w:p>"
    "<w:p><w:r><w:t>流體力學實驗 第三組</w:t></w:r></w:p>"
)
_TEXTBOX_RUN = f"""<w:r {_TEXTBOX_NAMESPACES}>
  <mc:AlternateContent>
    <mc:Choice Requires="wps">
      <w:drawing><wp:inline><a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp><wps:txbx><w:txbxContent>{_TEXTBOX_BODY}</w:txbxContent></wps:txbx></wps:wsp>
        </a:graphicData>
      </a:graphic></wp:inline></w:drawing>
    </mc:Choice>
    <mc:Fallback>
      <w:pict><v:shape><v:textbox><w:txbxContent>{_TEXTBOX_BODY}</w:txbxContent></v:textbox></v:shape></w:pict>
    </mc:Fallback>
  </mc:AlternateContent>
</w:r>"""
_TEXTBOX_LINES = "國立成功大學機械工程學系\n流體力學實驗 第三組"


def _strategy_project_identity():
    return {
        "required_terms": ["deterministic compilation", "StrategyIR", "AST", "Taiwan equities"],
        "required_context_terms": ["compiler architecture", "domain-specific intermediate representation"],
        "forbidden_terms": ["U.S. equity markets", "StrategySpec JSON"],
        "canonical_title_terms": ["deterministic", "compilation", "StrategyIR", "compiler"],
        "domain_context": "Taiwan equities",
        "author_metadata": {},
    }


class MarkdownBaseDocumentParseTests(unittest.TestCase):
    def test_markdown_base_document_sections_are_extracted(self):
        md = """# Project Title

**Author:** Example

## Abstract

Abstract body.

## 1. Introduction

Intro body.

## 2. Research Scope and Design Framing

Scope body.

## 3. Methods

Methods body.

## References

Reference body.
"""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "base.md"
            path.write_text(md, encoding="utf-8")

            sections = _parse_markdown_sections(str(path))

        self.assertEqual(sections["abstract"], "Abstract body.")
        self.assertEqual(sections["introduction"], "Intro body.")
        self.assertEqual(sections["research_scope"], "Scope body.")
        self.assertEqual(sections["methods"], "Methods body.")
        self.assertEqual(sections["references"], "Reference body.")
        self.assertIn("Project Title", sections["preamble"])

    def test_markdown_subsections_remain_with_parent_section(self):
        md = """# Title

## 3. Methods

### 3.1 Research Design

Design body.

### 3.2 Corpus

Corpus body.

## 4. Results

### 4.1 Graph Structure

Results body.
"""
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "base.md"
            path.write_text(md, encoding="utf-8")
            sections = _parse_markdown_sections(str(path))

        self.assertIn("### 3.1 Research Design", sections["methods"])
        self.assertIn("Corpus body.", sections["methods"])
        self.assertIn("### 4.1 Graph Structure", sections["results"])


class StylePassMarkdownPreservationTests(unittest.TestCase):
    def test_style_pass_preserves_tables_figures_and_headings(self):
        md = """# Results

This section utilizes the results.

![Figure 1. Hub ranking](figures/hub.png)

Figure 1. Hub ranking from graph analysis.

| Metric | Value |
|--------|-------|
| Graph Nodes | 5,171 |
| Graph Edges | 9,764 |

## Details

- Item one
- Item two
"""
        polished, _issues = _apply_style_polish(md)

        self.assertIn("![Figure 1. Hub ranking](figures/hub.png)", polished)
        self.assertIn("| Metric | Value |", polished)
        self.assertIn("| Graph Nodes | 5,171 |", polished)
        self.assertIn("## Details", polished)
        self.assertIn("- Item one", polished)
        self.assertIn("uses the results", polished)


class DocxTableStyleFallbackTests(unittest.TestCase):
    def test_table_styling_does_not_require_table_grid_style(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "table.docx"
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Metric"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Graph Nodes"
            table.cell(1, 1).text = "5,171"
            doc.save(path)

            _style_tables_post_render(str(path))

            rendered = Document(str(path))
            self.assertEqual(len(rendered.tables), 1)
            self.assertEqual(rendered.tables[0].cell(1, 1).text, "5,171")


class AdmissionsMonographPolishTests(unittest.TestCase):
    def test_polish_removes_stock_report_scaffolding_without_touching_tables(self):
        md = """# Introduction

This work presents a deterministic compilation architecture. The paper is organized as follows. Section 2 establishes scope.

| Metric | Value |
|--------|-------|
| Graph Nodes | 5,171 |
"""
        polished, changed = polish_admissions_monograph(md)

        self.assertIn("This project develops a deterministic compilation architecture.", polished)
        self.assertNotIn("The paper is organized as follows", polished)
        self.assertIn("| Graph Nodes | 5,171 |", polished)
        self.assertTrue(changed)


class ReferenceRealityClassificationTests(unittest.TestCase):
    def test_book_reference_is_not_forced_to_human_review(self):
        category, reason = _classify_reference({
            "ref_id": "Aho, A. V. (2006). *Compilers: Principles, Techniques, and Tools*. Addison-Wesley.",
            "raw": "Aho, A. V. (2006). *Compilers: Principles, Techniques, and Tools*. Addison-Wesley.",
            "status": "skipped",
            "checks": [{"type": "skipped", "reason": "no_doi_or_arxiv"}],
            "errors": [],
        })
        self.assertEqual(category, "publisher_or_book_plausible")

    def test_failed_doi_is_hard_failure_category(self):
        category, reason = _classify_reference({
            "ref_id": "Bad DOI",
            "status": "failed",
            "checks": [{"type": "doi", "verified": False}],
            "errors": ["DOI not found"],
        })
        self.assertEqual(category, "failed_external_verification")

    def test_year_like_journal_volume_is_not_arxiv_candidate(self):
        self.assertTrue(
            _is_publication_reference_candidate(
                "Cytron, R. (1991). Efficiently computing static single assignment form. "
                "*ACM Transactions on Programming Languages and Systems*, 13(4), 451-490. "
                "https://doi.org/10.1145/115372.115320"
            )
        )


class FrontMatterPrecedenceTests(unittest.TestCase):
    def test_structured_title_wins_over_base_preamble_title(self):
        state = ReportState.new("revise report", ["base.md"], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "revise_existing"
        state.spec["front_matter"] = {
            "title": "Short Explicit Title",
            "author_block": "Author Name",
            "affiliation_block": "Affiliation",
            "correspondence": "email@example.com",
            "keywords": ["deterministic compilation", "StrategyIR", "AST compilation"],
        }
        state.plan["blueprint"] = {"front_matter": {}}
        state.sources["base_document_sections"] = {
            "preamble": "# Very Long Base Title: A Subtitle That Should Not Override Explicit Structured Metadata\n\n---"
        }

        front_matter = _build_front_matter(state)

        self.assertEqual(front_matter["title"], "Short Explicit Title")

    def test_markdown_bold_front_matter_is_cleaned(self):
        state = ReportState.new("revise report", ["base.md"], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"front_matter": {}}
        state.sources["base_document_sections"] = {
            "preamble": (
                "# Deterministic Compilation Architecture\n\n---\n\n"
                "**Author:** Example Student\n"
                "**Affiliation:** Department of Engineering, Example University\n"
                "**Correspondence:** student@example.edu\n"
                "**Keywords:** deterministic compilation, StrategyIR, AST compilation\n"
            )
        }

        front_matter = _build_front_matter(state)

        self.assertEqual(front_matter["author_block"], "Example Student")
        self.assertEqual(
            front_matter["affiliation_block"],
            "Department of Engineering, Example University",
        )
        self.assertEqual(front_matter["correspondence"], "student@example.edu")
        self.assertEqual(front_matter["keywords"][0], "deterministic compilation")

    def test_generic_research_metadata_hard_fails(self):
        state = ReportState.new("write report", ["source.md"], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_report"
        state.spec["front_matter"] = {
            "title": "Deterministic Compilation Architecture",
            "author_block": "Research Author",
            "affiliation_block": "Department of Computer Science, Research University",
            "correspondence": "research@university.edu",
            "keywords": ["deterministic compilation", "StrategyIR", "AST compilation"],
        }
        state.plan["blueprint"] = {"front_matter": {}}

        with self.assertRaises(QAHardBlockError):
            run_front_matter_build(state)

    def test_admissions_project_report_allows_missing_publication_metadata(self):
        state = ReportState.new("write admissions project report", ["source.md"], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_project_report"
        state.plan["blueprint"] = {
            "front_matter": {
                "title": "Quant Strategy Validation Platform",
                "keywords": ["Taiwan equities", "validation"],
            }
        }

        result = run_front_matter_build(state)

        self.assertEqual(result.plan["front_matter"]["title"], "Quant Strategy Validation Platform")


class QualityGateContractTests(unittest.TestCase):
    def test_stable_evidence_id_repeats_for_same_source_span(self):
        entry = {"source_id": "abc123", "file_name": "source.md"}
        block = {"line_start": 10, "line_end": 12, "content_hash": "deadbeefcafebabe", "content": "same"}

        self.assertEqual(stable_evidence_id(entry, block), stable_evidence_id(entry, block))

    def test_claim_plan_blocks_cross_run_evidence_ids_early(self):
        state = ReportState.new("report", [], "out")
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        evidence_id = "E_current_123"
        (run_dir / "evidence_ledger.jsonl").write_text(
            json.dumps({"evidence_id": evidence_id}) + "\n",
            encoding="utf-8",
        )
        state.sources["evidence_ledger_path"] = str(run_dir / "evidence_ledger.jsonl")
        (run_dir / "claim_matrix.json").write_text(json.dumps({
            "claims": [{
                "claim_id": "c1",
                "claim_text": "Claim copied from another run.",
                "evidence_ids": ["E_old_999"],
                "claim_role": "primary",
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_claim_plan(state)
        self.assertIn("remap-evidence", str(ctx.exception))

    def test_remap_evidence_dry_run_and_write_updates_artifacts(self):
        old = ReportState.new("old", [], "out")
        new = ReportState.new("new", [], "out")
        old_dir = WORKFLOW_RUNS_DIR / old.job_id
        new_dir = WORKFLOW_RUNS_DIR / new.job_id
        old_ev = {
            "evidence_id": "old123",
            "source_file_name": "source.md",
            "line_start": 1,
            "line_end": 2,
            "content_hash": "hash-a",
            "quote": "same quote",
        }
        new_ev = {**old_ev, "evidence_id": "new456"}
        (old_dir / "evidence_ledger.jsonl").write_text(json.dumps(old_ev) + "\n", encoding="utf-8")
        (new_dir / "evidence_ledger.jsonl").write_text(json.dumps(new_ev) + "\n", encoding="utf-8")
        (new_dir / "source_registry.json").write_text("[]", encoding="utf-8")
        (new_dir / "claim_matrix.json").write_text(json.dumps({
            "claims": [{"claim_id": "c1", "claim_text": "Claim", "evidence_ids": ["old123"], "claim_role": "primary"}]
        }), encoding="utf-8")
        (new_dir / "outline.json").write_text(json.dumps({
            "sections": {"results": {"claim_ids": ["c1"]}}
        }), encoding="utf-8")
        (new_dir / "sentence_map.jsonl").write_text(
            json.dumps({"sentence_id": "s1", "section_id": "results", "claim_ids": ["c1"], "evidence_ids": ["old123"], "citation_ids": ["old123"]}) + "\n",
            encoding="utf-8",
        )
        section_dir = new_dir / "section_drafts"
        section_dir.mkdir(exist_ok=True)
        (section_dir / "results.md").write_text("Result [CITE:old123].", encoding="utf-8")

        dry = remap_evidence_ids(new.job_id, old.job_id, write=False)
        self.assertEqual(dry["mapping"]["old123"], "new456")
        self.assertIn("old123", (new_dir / "claim_matrix.json").read_text(encoding="utf-8"))

        written = remap_evidence_ids(new.job_id, old.job_id, write=True)
        self.assertEqual(written["status"], "ok")
        self.assertIn("new456", (new_dir / "claim_matrix.json").read_text(encoding="utf-8"))
        self.assertIn("[CITE:new456]", (section_dir / "results.md").read_text(encoding="utf-8"))
        self.assertEqual(load_artifact_contract(new_dir / "claim_matrix.json")["job_id"], new.job_id)

    def test_revision_apply_blocks_no_op_change(self):
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.sources["base_document_sections"] = {"abstract": "This work presents a system."}
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "abstract",
                "change_type": "replace",
                "original_text": "This work presents",
                "new_text": "This work presents",
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError):
            run_revision_apply(state)

    def test_revision_apply_blocks_figure_reference_deletion_without_reason(self):
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"section_order": ["results"]}
        state.sources["base_document_sections"] = {"results": "Figure 1 shows the architecture."}
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        sections_path = run_dir / "base_document_sections.json"
        sections_path.write_text(json.dumps(state.sources["base_document_sections"]), encoding="utf-8")
        state.sources["base_document_sections_path"] = str(sections_path)
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "results",
                "change_type": "delete",
                "original_text": "Figure 1 shows the architecture.",
                "new_text": "",
                "editorial": True,
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_revision_apply(state)
        self.assertIn("remove_figure_reference", str(ctx.exception))

    def test_revision_apply_blocks_chinese_figure_and_table_deletion(self):
        # The same guard, on the language this tool is mostly used in. Matching
        # only "Figure"/"Table" left a Chinese report with no protection at all.
        for label, body, expected in (
            ("figure", "圖 1. 系統架構\n如圖 1 所示。", "remove_figure_reference"),
            ("table", "表 1. 參數設定\n如表 1 所示。", "remove_table_reference"),
        ):
            with self.subTest(label):
                state = ReportState.new("revise", [], "out")
                state.spec["task_intent"] = "revise_existing"
                state.plan["blueprint"] = {"section_order": ["results"]}
                state.sources["base_document_sections"] = {"results": body}
                run_dir = WORKFLOW_RUNS_DIR / state.job_id
                sections_path = run_dir / "base_document_sections.json"
                sections_path.write_text(
                    json.dumps(state.sources["base_document_sections"], ensure_ascii=False),
                    encoding="utf-8",
                )
                state.sources["base_document_sections_path"] = str(sections_path)
                (run_dir / "revision_plan.json").write_text(json.dumps({
                    "changes": [{
                        "section_id": "results",
                        "change_type": "delete",
                        "original_text": body,
                        "new_text": "",
                        "editorial": True,
                    }]
                }, ensure_ascii=False), encoding="utf-8")

                with self.assertRaises(QAHardBlockError) as ctx:
                    run_revision_apply(state)
                self.assertIn(expected, str(ctx.exception))

    def test_a_title_row_above_the_header_does_not_become_the_header(self):
        # A monthly export opens with its own name in A1. Read as the header,
        # the title named the first column, the rest became Unnamed, the real
        # header turned into a citable data row, and in CSV the one-cell header
        # truncated every row to a single column — two measurements gone.
        import tempfile
        from report_workflow.parsers.structured_parser import parse_csv

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "月報.csv"
        path.write_text(
            "2025 年第一季產線月報\n\n月份,不良率(%),產量(件)\n1,1.8,1230\n",
            encoding="utf-8",
        )
        rows = parse_csv(str(path))
        self.assertEqual(rows, [{"月份": "1", "不良率(%)": "1.8", "產量(件)": "1230"}])

    def test_the_same_file_attached_twice_does_not_clear_the_evidence_bar(self):
        # Dragging the same export in from two folders registers two sources
        # and doubles the ledger — six entries for three readings. The bar is
        # about how much material there is, so it counts distinct content. That
        # is enforced by one set comprehension and explained by a comment;
        # nothing went red if someone counted entries again.
        from report_workflow.nodes.qa_gate import _source_diversity_reasons

        readings = [
            {"evidence_id": f"E_a_{i}", "content": f'{{"月份": "{i}", "不良率(%)": "1.{i}"}}'}
            for i in range(1, 4)
        ]
        duplicated = readings + [
            {**row, "evidence_id": row["evidence_id"].replace("E_a_", "E_b_")}
            for row in readings
        ]

        import tempfile

        tmpdir = Path(tempfile.mkdtemp())
        ledger = tmpdir / "evidence_ledger.jsonl"
        ledger.write_text(
            "\n".join(json.dumps(row, ensure_ascii=False) for row in duplicated) + "\n",
            encoding="utf-8",
        )
        state = ReportState.new("分析不良率", [], str(tmpdir / "out"))
        state.spec["report_profile"] = "academic_paper"
        state.sources["evidence_ledger_path"] = str(ledger)
        state.plan["claim_matrix"] = {"claims": []}
        reasons = " ".join(_source_diversity_reasons(state))
        self.assertIn("distinct", reasons)
        self.assertIn("found 3", reasons)

    def test_an_unsupported_type_is_described_to_the_author_not_the_maintainer(self):
        # Lecture slides are what a student has. They were told "agent fallback
        # parser is not implemented in the local MVP; deterministic parser could
        # not handle file_type='pptx'" — three phrases about this build and none
        # about their file. The unsupported branch left result as None, so the
        # specific reason never outranked the fallback's wording.
        import tempfile
        import zipfile
        from report_workflow.nodes.source_parse import parse_single_source

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "slides.pptx"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("[Content_Types].xml", "<?xml version='1.0'?><Types/>")

        result = parse_single_source({"file_path": str(path), "file_type": "pptx"})
        self.assertFalse(result["success"])
        self.assertIn(".pptx files are not read by this tool", result["error"])
        self.assertIn("pdf", result["error"])
        self.assertNotIn("local MVP", result["error"])

    def test_a_scanned_pdf_says_it_is_a_scan(self):
        # A department hands out scanned PDFs. "The file contains no readable
        # content" is equally true of an empty file, so the author was told
        # something that reads like "your file is broken" instead of the one
        # fact that lets them act: there are pages, there is no text layer.
        import tempfile
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_pdf import PdfPages
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "掃描.pdf"
        with PdfPages(str(path)) as pdf:
            figure = plt.figure(figsize=(4, 4))
            axes = figure.add_axes([0, 0, 1, 1])
            axes.axis("off")
            axes.imshow([[0, 255], [255, 0]], cmap="gray")
            pdf.savefig(figure)
            plt.close(figure)

        result = parse_semi_structured(str(path), "pdf")
        self.assertFalse(result["success"])
        self.assertIn("no text layer", result["error"])
        self.assertIn("1 page", result["error"])

    def test_a_text_pdf_still_parses(self):
        import tempfile
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_pdf import PdfPages
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "text.pdf"
        with PdfPages(str(path)) as pdf:
            figure = plt.figure(figsize=(4, 3))
            figure.text(0.1, 0.5, "defect rate 1.8%")
            pdf.savefig(figure)
            plt.close(figure)

        result = parse_semi_structured(str(path), "pdf")
        self.assertTrue(result["success"])
        self.assertTrue(result["blocks"])

    def test_a_merged_word_header_keeps_both_readings(self):
        # A spec sheet merges one label across two columns. python-docx hands
        # back the label twice with a newline between, which broke the header
        # into rows of different widths; the row A1 / 5 / 7 then reached the
        # ledger as one reading keyed on text nobody wrote.
        import tempfile
        from docx import Document
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured
        from report_workflow.nodes.evidence_normalize import _table_row_blocks

        tmpdir = Path(tempfile.mkdtemp())
        document = Document()
        table = document.add_table(rows=2, cols=3)
        header = table.rows[0].cells
        header[0].text = "機台"
        header[1].text = "解析度(μm)"
        header[2].text = "解析度(μm)"
        header[1].merge(header[2])
        body = table.rows[1].cells
        body[0].text = "A1"
        body[1].text = "5"
        body[2].text = "7"
        path = tmpdir / "規格.docx"
        document.save(str(path))

        blocks = parse_semi_structured(str(path), "docx")["blocks"]
        table_block = next(b for b in blocks if b["block_type"] == "table")
        self.assertNotIn("\n", table_block["table_data"][0][1])

        rows = _table_row_blocks(table_block)
        self.assertEqual(
            json.loads(rows[0]["content"]),
            {"機台": "A1", "解析度(μm)": "5", "解析度(μm) [2]": "7"},
        )

    def test_a_word_equation_survives_ingestion(self):
        # A lab report states its theory in Word equations. Reading only the
        # runs left the sentence promising a definition it no longer carried,
        # and an equation set on its own line produced no block at all.
        import tempfile
        from docx import Document
        from docx.oxml import parse_xml
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured

        tmpdir = Path(tempfile.mkdtemp())
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("雷諾數定義為 ")
        paragraph._p.append(parse_xml(_INLINE_OMML))
        paragraph.add_run(" ，取 D = 25 mm。")
        document.add_paragraph()._p.append(parse_xml(_DISPLAY_OMML))
        path = tmpdir / "理論.docx"
        document.save(str(path))

        blocks = parse_semi_structured(str(path), "docx")["blocks"]
        contents = [block["content"] for block in blocks]
        self.assertEqual(contents[0], "雷諾數定義為 Re=ρVD/μ ，取 D = 25 mm。")
        # The denominator carries an exponent, so writing it inline without
        # brackets would read as 2F over ρ, times U squared, times A.
        self.assertEqual(contents[1], "C_D=2F/(ρU^2A)")

    def test_a_cover_built_out_of_text_boxes_arrives_once(self):
        # A text box's paragraphs are not body paragraphs, so a cover page built
        # out of them was invisible. Reading every descendant instead brought it
        # in twice, because Word writes the same words under mc:Choice and
        # mc:Fallback, and the revision reader added a third copy by treating
        # those paragraphs as top-level ones.
        import tempfile
        from docx import Document
        from docx.oxml import parse_xml
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured
        from report_workflow.nodes.base_document_parse import _parse_docx_section

        tmpdir = Path(tempfile.mkdtemp())
        document = Document()
        document.add_paragraph()._p.append(parse_xml(_TEXTBOX_RUN))
        document.add_paragraph("1. 實驗目的")
        path = tmpdir / "封面.docx"
        document.save(str(path))

        blocks = parse_semi_structured(str(path), "docx")["blocks"]
        self.assertEqual(blocks[0]["content"], _TEXTBOX_LINES)

        sections, _titles = _parse_docx_section(str(path))
        body = "\n".join(sections.values())
        self.assertEqual(body.count("國立成功大學機械工程學系"), 1)
        self.assertIn(_TEXTBOX_LINES, body)

    def test_a_word_equation_is_not_flattened_when_revising(self):
        # This reader collected every "}t", which reaches m:t as well, and
        # joined them with nothing: ρVD over μ came back as ρVDμ and would have
        # been written into the author's own report as if they had typed it.
        import tempfile
        from docx import Document
        from docx.oxml import parse_xml
        from report_workflow.nodes.base_document_parse import _parse_docx_section

        tmpdir = Path(tempfile.mkdtemp())
        document = Document()
        document.add_heading("理論背景", level=1)
        paragraph = document.add_paragraph()
        paragraph.add_run("雷諾數定義為 ")
        paragraph._p.append(parse_xml(_INLINE_OMML))
        path = tmpdir / "原報告.docx"
        document.save(str(path))

        sections, _titles = _parse_docx_section(str(path))
        body = "\n".join(sections.values())
        self.assertIn("Re=ρVD/μ", body)
        self.assertNotIn("ρVDμ", body)

    def test_every_sheet_of_a_workbook_reaches_the_records(self):
        # One year per tab is how these workbooks arrive. Only the first sheet
        # was read, so the other year's rows never reached the ledger and
        # nothing said they existed.
        import tempfile
        from openpyxl import Workbook
        from report_workflow.parsers.structured_parser import parse_xlsx

        tmpdir = Path(tempfile.mkdtemp())
        book = Workbook()
        first = book.active
        first.title = "2025"
        first.append(["月份", "不良率(%)"])
        first.append([1, 1.8])
        second = book.create_sheet("2024")
        second.append(["月份", "不良率(%)"])
        second.append([1, 4.2])
        path = tmpdir / "年度比較.xlsx"
        book.save(str(path))

        records = parse_xlsx(str(path))
        self.assertEqual(
            records,
            [
                {"sheet": "2025", "月份": 1, "不良率(%)": 1.8},
                {"sheet": "2024", "月份": 1, "不良率(%)": 4.2},
            ],
        )

    def test_a_vaults_bookkeeping_is_not_evidence(self):
        # A note exported from Obsidian opens with front matter. Read as
        # content it became one citable block, so a claim could trace to
        # "created: 2026-07-12" or to the note's own summary and clear the
        # gate with nothing measured behind it — and the entry count that asks
        # whether there is enough material got a free one. Revising such a note
        # printed the same block into the report as its preamble.
        import tempfile
        from report_workflow.parsers.semi_structured_parser import (
            _front_matter_end,
            parse_semi_structured,
        )
        from report_workflow.nodes.base_document_parse import _parse_markdown_sections

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "筆記.md"
        path.write_text(
            "---\ntype: lab-note\ntags:\n  - 熱傳\nstatus: active\n"
            "created: 2026-07-12\nsummary: 四個流量點的有效度量測。\n---\n\n"
            "# 有效度量測\n\n每點取三次平均。\n",
            encoding="utf-8",
        )

        blocks = parse_semi_structured(str(path), "md")["blocks"]
        joined = "\n".join(block["content"] for block in blocks)
        self.assertNotIn("created:", joined)
        self.assertNotIn("lab-note", joined)
        self.assertIn("每點取三次平均。", joined)

        sections = _parse_markdown_sections(str(path))
        self.assertNotIn("preamble", sections)

        # A document that opens with a horizontal rule is left exactly as it is.
        self.assertEqual(_front_matter_end(["---\n", "\n", "# Title\n"]), 0)
        self.assertEqual(_front_matter_end(["---\n", "A sentence.\n"]), 0)

    def test_a_hidden_note_to_self_does_not_reach_the_report(self):
        # %%...%% is Obsidian's comment: the author cannot see it in reading
        # view, which is the point of it. "The 4.0 point is not calibrated, do
        # not use the number" was citable evidence, and revising the note
        # printed it into the document they hand in. The tag line above it was
        # typed as a heading, which CommonMark does not call one either.
        import tempfile
        from report_workflow.parsers.semi_structured_parser import parse_semi_structured
        from report_workflow.nodes.base_document_parse import _parse_markdown_sections

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "筆記.md"
        path.write_text(
            "# 有效度量測\n\n#熱傳 #待複查\n\n每點取三次平均。\n\n"
            "%%4.0 那點還沒校正，數字先不要用%%\n\n說明結束。\n",
            encoding="utf-8",
        )

        blocks = parse_semi_structured(str(path), "md")["blocks"]
        joined = "\n".join(block["content"] for block in blocks)
        self.assertNotIn("還沒校正", joined)
        self.assertIn("每點取三次平均。", joined)
        tag_block = next(b for b in blocks if "#熱傳" in b["content"])
        self.assertEqual(tag_block["block_type"], "paragraph")
        # A real heading is still a heading.
        self.assertEqual(blocks[0]["block_type"], "heading")

        sections = _parse_markdown_sections(str(path))
        self.assertNotIn("還沒校正", "\n".join(sections.values()))

    def test_a_merged_group_column_keeps_every_row_its_group(self):
        # A course data sheet merges 組別 down the runs it covers. pandas sees
        # the top-left value and blanks beside it, so four of six readings
        # reached the ledger with no way to say whose they were — while the run
        # nobody measured must stay unmeasured, which is why the merged ranges
        # are read out of the file instead of gaps being carried down.
        import tempfile
        from openpyxl import Workbook
        from report_workflow.parsers.structured_parser import parse_structured

        tmpdir = Path(tempfile.mkdtemp())
        book = Workbook()
        sheet = book.active
        sheet.append(["組別", "試次", "流量(L/min)", "效率(%)"])
        for row in [
            ("A", 1, 2.0, 58.1), (None, 2, 2.5, 60.3), (None, 3, 3.0, 62.0),
            ("B", 1, 2.0, 57.4), (None, 2, 2.5, 59.8), (None, 3, 3.0, None),
        ]:
            sheet.append(row)
        sheet.merge_cells("A2:A4")
        sheet.merge_cells("A5:A7")
        path = tmpdir / "實驗數據.xlsx"
        book.save(str(path))

        blocks = parse_structured(str(path), "xlsx")["blocks"]
        contents = [json.loads(block["content"]) for block in blocks]
        self.assertEqual([row["組別"] for row in contents], list("AAABBB"))
        # 3.0 measured alongside 2.5 must not come back as 3 — asserted on the
        # text the ledger actually carries, because 3 == 3.0 in Python and the
        # comparison this used to make could not tell them apart at all.
        self.assertIn('"流量(L/min)": 3.0', blocks[2]["content"])
        # The last run was never taken. null, because NaN is not JSON and the
        # ledger is read by whoever checks the report.
        self.assertIsNone(contents[5]["效率(%)"])
        self.assertNotIn("NaN", json.dumps(contents))

    def test_a_stacked_header_does_not_make_the_sheet_unreadable(self):
        # 溫度(°C) spans two sub-columns while 機台 and 有效度(%) are merged down
        # across both header rows — the ordinary shape of a course data sheet.
        # Filling those downward merges as data wrote a column's own name into
        # the first reading; pandas refused a string in a float column and the
        # whole sheet came back with success: False and no evidence at all.
        import tempfile
        from openpyxl import Workbook
        from report_workflow.parsers.structured_parser import parse_structured

        tmpdir = Path(tempfile.mkdtemp())
        book = Workbook()
        sheet = book.active
        sheet.append(["機台", "溫度(°C)", None, "有效度(%)"])
        sheet.append([None, "入口", "出口", None])
        sheet.append(["A1", 60.0, 45.2, 72.4])
        sheet.append(["A2", 60.0, 43.1, 76.1])
        sheet.merge_cells("B1:C1")
        sheet.merge_cells("A1:A2")
        sheet.merge_cells("D1:D2")
        path = tmpdir / "量測.xlsx"
        book.save(str(path))

        result = parse_structured(str(path), "xlsx")
        self.assertTrue(result["success"], result.get("error"))
        rows = [json.loads(block["content"]) for block in result["blocks"]]
        readings = [row for row in rows if row.get("機台") in ("A1", "A2")]
        self.assertEqual(len(readings), 2)
        self.assertEqual(readings[0]["有效度(%)"], 72.4)
        self.assertEqual(readings[1]["溫度(°C) [2]"], 43.1)

    def test_a_merged_xlsx_header_keeps_both_readings(self):
        # The same defect the DOCX reader had: one label merged across two
        # columns left the second keyed on "Unnamed: 2", text nobody wrote.
        import tempfile
        from openpyxl import Workbook
        from report_workflow.parsers.structured_parser import parse_xlsx

        tmpdir = Path(tempfile.mkdtemp())
        book = Workbook()
        sheet = book.active
        sheet.append(["機台", "解析度(μm)", None])
        sheet.append(["A1", 5, 7])
        sheet.merge_cells("B1:C1")
        path = tmpdir / "規格.xlsx"
        book.save(str(path))

        self.assertEqual(
            parse_xlsx(str(path)),
            [{"機台": "A1", "解析度(μm)": 5, "解析度(μm) [2]": 7}],
        )

    def test_a_single_sheet_workbook_gains_no_extra_column(self):
        import tempfile
        from openpyxl import Workbook
        from report_workflow.parsers.structured_parser import parse_xlsx

        tmpdir = Path(tempfile.mkdtemp())
        book = Workbook()
        sheet = book.active
        sheet.append(["月份", "不良率(%)"])
        sheet.append([1, 1.8])
        path = tmpdir / "單分頁.xlsx"
        book.save(str(path))

        self.assertEqual(parse_xlsx(str(path)), [{"月份": 1, "不良率(%)": 1.8}])

    def test_a_header_with_a_blank_corner_is_not_skipped(self):
        # The guard must not reach past a real header: a matrix table leaves
        # its corner cell empty and still holds two or more values.
        import tempfile
        from report_workflow.parsers.structured_parser import parse_csv

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "matrix.csv"
        path.write_text(",甲線,乙線\n1月,1.8,2.1\n", encoding="utf-8")
        rows = parse_csv(str(path))
        self.assertEqual(rows, [{"": "1月", "甲線": "1.8", "乙線": "2.1"}])

    def test_a_plain_table_is_untouched(self):
        import tempfile
        from report_workflow.parsers.structured_parser import parse_csv

        tmpdir = Path(tempfile.mkdtemp())
        path = tmpdir / "plain.csv"
        path.write_text("月份,不良率(%)\n1,1.8\n", encoding="utf-8")
        self.assertEqual(parse_csv(str(path)), [{"月份": "1", "不良率(%)": "1.8"}])

    def test_a_source_missing_at_publish_time_is_not_skipped_in_silence(self):
        # A source moved between prepare and publish was skipped without a
        # word: the bundle shipped short, artifacts.json listed only what made
        # it, and the report kept citing claims traced to the absent file.
        import tempfile
        import uuid
        from report_workflow.nodes.artifacts import run_artifacts

        tmpdir = Path(tempfile.mkdtemp())
        present = tmpdir / "kept.csv"
        present.write_text("a,b\n1,2\n", encoding="utf-8")
        absent = tmpdir / "gone.csv"

        state = ReportState.new("report", [], str(tmpdir / "out"))
        state.job_id = f"test_missing_source_{uuid.uuid4().hex}"
        state.qa["qa_decision"] = "pass"
        state.qa["artifact_completeness_status"] = "pass"
        state.spec["uploaded_files"] = [str(present), str(absent)]

        with self.assertRaises(QAHardBlockError) as ctx:
            run_artifacts(state)
        self.assertIn("gone.csv", str(ctx.exception))

    def test_a_missing_base_document_is_named_as_a_base_document(self):
        # A base document is not a cited source — it is the document being
        # revised. The role comes from the registry, which stores resolved
        # paths while uploaded_files keeps the caller's spelling, so the two
        # are compared resolved; matching raw strings labelled every base
        # document a plain source.
        import tempfile
        import uuid
        from report_workflow.nodes.artifacts import run_artifacts

        tmpdir = Path(tempfile.mkdtemp())
        absent = tmpdir / "原稿.docx"

        state = ReportState.new("revise", [], str(tmpdir / "out"))
        state.job_id = f"test_missing_base_{uuid.uuid4().hex}"
        state.qa["qa_decision"] = "pass"
        state.qa["artifact_completeness_status"] = "pass"
        # The caller's spelling goes through a redundant "." segment; the
        # registry holds the resolved form, as the real pipeline writes it.
        state.spec["uploaded_files"] = [str(tmpdir / "." / "原稿.docx")]
        state.sources["source_registry"] = [{
            "source_id": "S", "file_name": "原稿.docx",
            "file_path": str(absent.resolve()), "artifact_role": "base_document",
        }]

        with self.assertRaises(QAHardBlockError) as ctx:
            run_artifacts(state)
        self.assertIn("base_document", str(ctx.exception))

    def test_publishing_still_succeeds_when_every_source_is_present(self):
        import tempfile
        import uuid
        from report_workflow.nodes.artifacts import run_artifacts

        tmpdir = Path(tempfile.mkdtemp())
        for name in ("a.csv", "b.csv"):
            (tmpdir / name).write_text("x,y\n1,2\n", encoding="utf-8")

        state = ReportState.new("report", [], str(tmpdir / "out"))
        state.job_id = f"test_present_sources_{uuid.uuid4().hex}"
        state.qa["qa_decision"] = "pass"
        state.qa["artifact_completeness_status"] = "pass"
        state.spec["uploaded_files"] = [str(tmpdir / "a.csv"), str(tmpdir / "b.csv")]

        state = run_artifacts(state)
        published = Path(state.output["published_dir"]) / "sources"
        self.assertEqual(sorted(p.name for p in published.iterdir()), ["a.csv", "b.csv"])

    def test_a_source_edited_after_prepare_is_caught_before_publishing(self):
        # The bundle is copied from the original path at publish time, so a
        # source fixed in between ships beside evidence quoting text it no
        # longer contains — and 72.4 to 82.4 is the same number of bytes, which
        # is why file_size, recorded since the start, could never see it.
        import tempfile
        from report_workflow.artifact_contract import _hash_bytes
        from report_workflow.nodes.artifacts import _drifted_sources

        tmpdir = Path(tempfile.mkdtemp())
        source = tmpdir / "量測數據.csv"
        source.write_text("流量,有效度\n2.0,72.4\n", encoding="utf-8")
        entry = {
            "file_name": source.name,
            "file_path": str(source),
            "file_size": source.stat().st_size,
            "content_hash": _hash_bytes(source.read_bytes()),
        }

        self.assertEqual(_drifted_sources([entry]), [])

        source.write_text("流量,有效度\n2.0,82.4\n", encoding="utf-8")
        self.assertEqual(source.stat().st_size, entry["file_size"])
        drifted = _drifted_sources([entry])
        self.assertTrue(drifted)
        self.assertIn("量測數據.csv", drifted[0])

        # A run recorded before the hash was kept must not start failing, and a
        # file that is gone is left to the check that already names it.
        self.assertEqual(_drifted_sources([{k: v for k, v in entry.items()
                                            if k != "content_hash"}]), [])
        self.assertEqual(_drifted_sources([{**entry, "file_path": str(tmpdir / "gone.csv")}]), [])

    def test_same_named_sources_both_reach_the_published_bundle(self):
        # Both files landed on published/sources/月報.csv, so the bundle kept
        # whichever was copied last. A reader checking the report's 2024 claim
        # opened the survivor and read 2025's numbers.
        from report_workflow.nodes.artifacts import _unique_source_names

        # Built with the running platform's separator: a hardcoded Windows path
        # is one long filename on Linux, where nothing has a folder to be told
        # apart by, so this passed here and failed on CI.
        first = os.path.join("a", "2024", "月報.csv")
        second = os.path.join("a", "2025", "月報.csv")
        spec = os.path.join("a", "規格.pdf")
        names = _unique_source_names([first, second, spec])
        self.assertEqual(names[first], "2024_月報.csv")
        self.assertEqual(names[second], "2025_月報.csv")
        self.assertEqual(names[spec], "規格.pdf")

    def test_identically_foldered_sources_still_get_distinct_names(self):
        from report_workflow.nodes.artifacts import _unique_source_names

        names = _unique_source_names([r"C:\a\data\月報.csv", r"C:\b\data\月報.csv"])
        self.assertEqual(len(set(names.values())), 2)

    def test_a_blocked_render_names_the_text_that_tripped_it(self):
        # "Raw prompt fragment leaked into publication text" with no fragment
        # quoted leaves nothing to search the drafts for.
        from report_workflow.nodes.docx_render import _pre_render_sanity_check

        prompt = "比較 2024 與 2025 產線不良率"
        issues = _pre_render_sanity_check(
            f"# 報告\n\n{prompt}\n", forbidden_fragments=[prompt]
        )
        leak = [i for i in issues if "Raw prompt fragment" in i]
        self.assertTrue(leak, issues)
        self.assertIn(prompt, leak[0])

    def test_same_named_sources_are_told_apart_in_the_brief(self):
        # Monthly exports arrive as 2024/月報.csv and 2025/月報.csv. Shown as
        # "月報.csv" twice, an author picking between two rows of numbers has
        # nothing to pick by, and citing the wrong year passes every gate.
        from report_workflow.nodes.agent_tasks import _source_labels

        first = os.path.join("a", "2024", "月報.csv")
        second = os.path.join("a", "2025", "月報.csv")
        spec = os.path.join("a", "規格.pdf")
        labels = _source_labels([
            {"source_file_name": "月報.csv", "source_file_path": first},
            {"source_file_name": "月報.csv", "source_file_path": second},
            {"source_file_name": "規格.pdf", "source_file_path": spec},
        ])
        self.assertEqual(labels[first], "2024/月報.csv")
        self.assertEqual(labels[second], "2025/月報.csv")
        self.assertEqual(labels[spec], "規格.pdf")

    def test_a_lone_source_keeps_its_plain_name(self):
        from report_workflow.nodes.agent_tasks import _source_labels

        labels = _source_labels(
            [{"source_file_name": "月報.csv", "source_file_path": r"C:\a\2024\月報.csv"}]
        )
        self.assertEqual(labels[r"C:\a\2024\月報.csv"], "月報.csv")

    def test_identical_parent_names_still_tell_the_sources_apart(self):
        """One folder per experiment, each with its own data/ inside.

        This used to try one parent and then print the whole absolute path —
        136 characters per row in a table whose job is to keep the agent's
        context small. What it has to be is distinct; being long was never the
        requirement, and the full path remains where the walk ends.
        """
        from report_workflow.nodes.agent_tasks import _source_labels

        first = os.path.join("實驗一", "data", "量測.csv")
        second = os.path.join("實驗二", "data", "量測.csv")
        third = os.path.join("實驗三", "data", "量測.csv")
        labels = _source_labels([
            {"source_file_name": "量測.csv", "source_file_path": first},
            {"source_file_name": "量測.csv", "source_file_path": second},
            {"source_file_name": "量測.csv", "source_file_path": third},
        ])
        self.assertEqual(len(set(labels.values())), 3)
        self.assertEqual(labels[first], "實驗一/data/量測.csv")
        for label in labels.values():
            self.assertLess(len(label), 30)

    def test_removing_a_whole_section_does_not_demand_a_figure_decision(self):
        # remove_section is already a declared structural change, recorded in
        # removed_sections. Counting its figures left the author one exit:
        # figure_preservation_decision='remove_because_no_source_asset', which
        # asserts the asset does not exist when it plainly does — a false entry
        # in the audit trail the gate exists to protect.
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"section_order": ["results"]}
        state.sources["base_document_sections"] = {
            "results": "導入後不良率下降。\n圖 1. 月別不良率。",
            "discussion": "樣本期間偏短。",
        }
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        sections_path = run_dir / "base_document_sections.json"
        sections_path.write_text(
            json.dumps(state.sources["base_document_sections"], ensure_ascii=False),
            encoding="utf-8",
        )
        state.sources["base_document_sections_path"] = str(sections_path)
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{"section_id": "results", "change_type": "remove_section"}]
        }, ensure_ascii=False), encoding="utf-8")

        state = run_revision_apply(state)
        report = json.loads(
            Path(state.runtime["revision_diff_report_path"]).read_text(encoding="utf-8")
        )
        self.assertEqual(report["removed_sections"], ["results"])

    def test_a_caption_deleted_without_removing_its_section_still_blocks(self):
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"section_order": ["results"]}
        state.sources["base_document_sections"] = {
            "results": "導入後不良率下降。\n圖 1. 月別不良率。",
        }
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        sections_path = run_dir / "base_document_sections.json"
        sections_path.write_text(
            json.dumps(state.sources["base_document_sections"], ensure_ascii=False),
            encoding="utf-8",
        )
        state.sources["base_document_sections_path"] = str(sections_path)
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "results",
                "change_type": "delete",
                "original_text": "圖 1. 月別不良率。",
                "new_text": "",
                "editorial": True,
            }]
        }, ensure_ascii=False), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_revision_apply(state)
        self.assertIn("remove_figure_reference", str(ctx.exception))

    def test_chinese_reference_ids_ignore_ordinary_prose(self):
        # 圖 and 表 are common morphemes; a prose mention is not a reference.
        # Counting them would hard-block honest sentences, which is worse than
        # missing an uncaptioned figure.
        from report_workflow.nodes.revision_apply import _reference_ids

        prose = {"s": "\n".join([
            "本團隊發表 2 篇相關論文。",
            "受訪者代表 3 家供應商。",
            "量表 5 點尺度計分。",
            "試圖 3 次後放棄。",
        ])}
        self.assertEqual(_reference_ids(prose, "Figure"), set())
        self.assertEqual(_reference_ids(prose, "Table"), set())

        captioned = {"s": "圖 1. 效能對流量。\n表 2. 參數設定\n本團隊發表 2 篇論文。"}
        self.assertEqual(_reference_ids(captioned, "Figure"), {"1"})
        self.assertEqual(_reference_ids(captioned, "Table"), {"2"})

    def test_revision_apply_requires_preservation_decision_for_figure_removal(self):
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"section_order": ["results"]}
        state.sources["base_document_sections"] = {"results": "Figure 1 shows the architecture."}
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        sections_path = run_dir / "base_document_sections.json"
        sections_path.write_text(json.dumps(state.sources["base_document_sections"]), encoding="utf-8")
        state.sources["base_document_sections_path"] = str(sections_path)
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "results",
                "change_type": "delete",
                "original_text": "Figure 1 shows the architecture.",
                "new_text": "",
                "change_reason": "remove_figure_reference",
                "editorial": True,
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_revision_apply(state)
        self.assertIn("figure_preservation_decision", str(ctx.exception))

    def test_revision_apply_blocks_figure_removal_when_prompt_requires_preserve(self):
        state = ReportState.new("revise and preserve figure/table references", [], "out")
        state.spec["task_intent"] = "revise_existing"
        state.plan["blueprint"] = {"section_order": ["results"]}
        state.sources["base_document_sections"] = {"results": "Figure 1 shows the architecture."}
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        sections_path = run_dir / "base_document_sections.json"
        sections_path.write_text(json.dumps(state.sources["base_document_sections"]), encoding="utf-8")
        state.sources["base_document_sections_path"] = str(sections_path)
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "results",
                "change_type": "delete",
                "original_text": "Figure 1 shows the architecture.",
                "new_text": "",
                "change_reason": "remove_figure_reference",
                "figure_preservation_decision": "remove_because_no_source_asset",
                "editorial": True,
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_revision_apply(state)
        self.assertIn("preservation", str(ctx.exception).lower())

    def test_editing_the_base_document_after_prepare_is_caught(self):
        # Fixing what looks like a typo between prepare and the revision being
        # applied is an ordinary thing to do, and 79.3 -> 89.3 changes neither
        # the file's size nor — as far as this check went — anything it
        # recorded. The revision would have been applied to text the file no
        # longer contained.
        import tempfile
        from report_workflow.artifact_contract import validate_base_document_integrity

        tmpdir = Path(tempfile.mkdtemp())
        base = tmpdir / "報告.md"
        base.write_text("# 結果\n\n量測有效度為 79.3%。\n", encoding="utf-8")
        sections = {"結果": "量測有效度為 79.3%。"}
        state = ReportState.new("revise", [], str(tmpdir / "out"))
        entry = {"file_path": str(base), "file_name": base.name}
        integrity_path = Path(write_base_document_integrity(state, sections, entry))

        # Nothing changed: the check must stay out of the way.
        validate_base_document_integrity(state, sections)

        base.write_text("# 結果\n\n量測有效度為 89.3%。\n", encoding="utf-8")
        with self.assertRaises(QAHardBlockError) as ctx:
            validate_base_document_integrity(state, sections)
        self.assertIn("報告.md", str(ctx.exception))

        # A run recorded before the field existed must not start failing.
        payload = json.loads(integrity_path.read_text(encoding="utf-8"))
        payload.pop("source_content_hash")
        integrity_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        validate_base_document_integrity(state, sections)

    def test_base_document_integrity_blocks_direct_section_mutation(self):
        state = ReportState.new("revise", [], "out")
        state.spec["task_intent"] = "revise_existing"
        original = {"results": "Figure 1 shows the architecture."}
        write_base_document_integrity(
            state,
            original,
            {"file_path": __file__, "file_name": "base.md"},
        )
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        (run_dir / "base_document_sections.json").write_text(
            json.dumps({"results": "The architecture is described."}),
            encoding="utf-8",
        )
        state.sources["base_document_sections_path"] = str(run_dir / "base_document_sections.json")
        (run_dir / "revision_plan.json").write_text(json.dumps({
            "changes": [{
                "section_id": "results",
                "change_type": "replace",
                "original_text": "architecture",
                "new_text": "compiler architecture",
            }]
        }), encoding="utf-8")

        with self.assertRaises(QAHardBlockError) as ctx:
            run_revision_apply(state)
        self.assertIn("immutable parse snapshot", str(ctx.exception))

    def test_evidence_provenance_blocks_hand_authored_primary_source(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            ledger = Path(tmpdir) / "evidence_ledger.jsonl"
            ledger.write_text(json.dumps({
                "evidence_id": "E003",
                "source_role": "primary_source",
                "content": "Hand-authored evidence without parser trace.",
            }) + "\n", encoding="utf-8")
            with self.assertRaises(QAHardBlockError):
                validate_evidence_ledger_provenance(ledger)

    def test_evidence_provenance_blocks_base_document_role(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            ledger = Path(tmpdir) / "evidence_ledger.jsonl"
            ledger.write_text(json.dumps({
                "evidence_id": "E001",
                "source_role": "base_document",
                "source_id": "src1",
                "source_file_name": "revised_report.md",
                "content": "Copied from base document.",
                "content_hash": "abc123",
            }) + "\n", encoding="utf-8")
            with self.assertRaises(QAHardBlockError):
                validate_evidence_ledger_provenance(ledger)

    def test_evidence_normalize_adds_missing_content_hash(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = ReportState.new("normalize parser blocks", [], tmpdir)
            state.sources["source_registry"] = [{
                "source_id": "S001",
                "file_name": "sample.pdf",
                "file_path": __file__,
                "file_type": "pdf",
                "file_size": 100,
                "artifact_role": "source_data",
                "parsed_content": [{
                    "block_id": "B001",
                    "block_type": "paragraph",
                    "content": "This method paragraph has enough source content to become evidence.",
                }],
            }]

            run_evidence_normalize(state)

            ledger = Path(state.sources["evidence_ledger_path"])
            row = json.loads(ledger.read_text(encoding="utf-8").strip())
            self.assertEqual(len(row["content_hash"]), 16)
            validate_evidence_ledger_provenance(ledger)

    def test_repo_hygiene_detects_orphan_repair_script(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            (root / "fix_claims.py").write_text("print('temp')", encoding="utf-8")
            issues = find_repo_hygiene_issues(root)
        self.assertTrue(any("fix_claims.py" in issue for issue in issues))

    def test_final_publish_blocks_root_orphan_scripts(self):
        state = ReportState.new("publish", [], "out")
        state.qa["qa_decision"] = "pass"
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            docx = root / "report.docx"
            doc = Document()
            doc.add_paragraph("Report")
            doc.save(docx)
            (root / "strip_figures.py").write_text("print('temp')", encoding="utf-8")
            state.output["final_docx_path"] = str(docx)
            state.output["output_dir"] = str(root / "out")
            with patch("report_workflow.nodes.final_publish.find_repo_hygiene_issues", return_value=[str(root / "strip_figures.py")]):
                with self.assertRaises(QAHardBlockError):
                    run_final_publish(state)

    def test_admissions_project_report_inference_and_reference_policy(self):
        profile = infer_report_profile(
            "Write an admissions-facing academic project report from internal architecture docs."
        )
        self.assertEqual(profile, "admissions_project_report")

    def test_admissions_profiles_are_reachable_in_chinese(self):
        """備審資料 is the name of the document, not a translation exercise.

        Every other profile carried its Chinese vocabulary; this branch carried
        only 申請, a generic verb. So the two admissions profiles and the gates
        written for them could not be selected by the words their users type,
        and 備審資料 fell through to academic_paper.
        """
        self.assertEqual(infer_report_profile("備審資料 自傳與讀書計畫"), "admissions_report")
        self.assertEqual(infer_report_profile("學習歷程檔案"), "admissions_report")
        self.assertEqual(
            infer_report_profile("研究所推甄備審 專題成果"), "admissions_project_report")
        # 專題 on its own is a course project, not an application.
        self.assertEqual(infer_report_profile("機械設計專題報告"), "academic_paper")
        # The profiles that already worked keep working.
        self.assertEqual(infer_report_profile("幫我寫熱傳學實驗報告"), "engineering_lab_report")
        self.assertEqual(infer_report_profile("企劃書"), "proposal")

    def test_report_profile_alias_normalizes_to_admissions_project(self):
        profile = normalize_profile_id("admissions project report")
        self.assertEqual(profile, "admissions_project_report")

    def test_reference_relevance_allows_internal_only_for_admissions_project_report(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_project_report"
        state.spec["task_intent"] = "new_draft"
        with tempfile.TemporaryDirectory() as tmpdir:
            report = Path(tmpdir) / "reference_verify_report.json"
            report.write_text(json.dumps({
                "references": [{
                    "ref_id": "Tsai, C.-C. (2026). *LLM Debate Architecture* (Architecture Documentation, Version 2026-04-13).",
                    "raw": "Tsai, C.-C. (2026). *LLM Debate Architecture* (Architecture Documentation, Version 2026-04-13).",
                    "status": "project_source",
                    "checks": [{"type": "project_source"}],
                }]
            }), encoding="utf-8")
            state.runtime["reference_verify_report_path"] = str(report)
            result = run_reference_relevance_gate(state)
        self.assertTrue(result.runtime["reference_relevance_report_path"])

    def test_submit_and_publish_report_returns_rendered_but_not_publishable(self):
        state = ReportState.new("report", [], "out")
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        doc = Document()
        doc.add_paragraph("Rendered but not published")
        rendered = run_dir / "rendered_report.docx"
        doc.save(rendered)
        result = submit_and_publish_report(state.job_id)
        self.assertTrue(result["rendered_but_not_publishable"])
        self.assertTrue(result["not_final_deliverable"])

    def test_project_identity_gate_blocks_topic_drift(self):
        state = ReportState.new("write admissions report", [], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "new_draft"
        state.spec["project_identity"] = _strategy_project_identity()
        state.plan["front_matter"] = {
            "title": "A Trustworthy Verification Framework for LLM-Assisted Quantitative Strategy Generation in U.S. Equity Markets"
        }
        state.plan["thesis_statement"] = "StrategySpec JSON with JSON Schema validation supports a seven-phase pipeline."
        with tempfile.TemporaryDirectory() as tmpdir:
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Abstract\n\nA StrategySpec JSON framework for U.S. equity markets is described.\n\n"
                "# Introduction\n\nThe report focuses on U.S. equity markets and JSON Schema validation.\n",
                encoding="utf-8",
            )
            state.drafts["merged_draft_md"] = str(draft)
            with self.assertRaises(QAHardBlockError):
                run_project_identity_gate(state)

    def test_project_identity_gate_passes_project_spine(self):
        state = ReportState.new("write admissions report", [], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "new_draft"
        state.spec["project_identity"] = _strategy_project_identity()
        state.plan["front_matter"] = {
            "title": "Deterministic Compilation Architecture for StrategyIR Trading Systems"
        }
        state.plan["thesis_statement"] = (
            "Deterministic compilation with StrategyIR and AST construction supports "
            "auditable Taiwan equities strategy generation."
        )
        with tempfile.TemporaryDirectory() as tmpdir:
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Abstract\n\nDeterministic compilation uses StrategyIR and AST construction for Taiwan equities.\n\n"
                "# Introduction\n\nThe project centers on compiler architecture, StrategyIR, and Taiwan equities verification.\n",
                encoding="utf-8",
            )
            state.drafts["merged_draft_md"] = str(draft)
            result = run_project_identity_gate(state)
        self.assertTrue(result.runtime["project_identity_report_path"])

    def test_project_identity_gate_uses_outline_thesis_when_scope_freeze_does_not_set_one(self):
        state = ReportState.new("write admissions project report", [], "out")
        state.spec["report_profile"] = "admissions_project_report"
        state.spec["task_intent"] = "new_draft"
        state.spec["project_identity"] = {
            "required_terms": ["structured workflow", "evidence handling", "QA thinking"],
            "required_context_terms": ["auditable artifacts"],
            "forbidden_terms": [],
            "canonical_title_terms": ["workflow"],
            "domain_context": "structured workflow",
            "author_metadata": {},
        }
        state.plan["front_matter"] = {"title": "Structured Workflow Evidence Project"}
        state.plan["outline"] = {
            "thesis_statement": (
                "The structured workflow demonstrates evidence handling and QA thinking "
                "through auditable artifacts."
            )
        }
        state.plan["primary_contribution"] = "A measured outcome claim without the identity spine."
        with tempfile.TemporaryDirectory() as tmpdir:
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Abstract\n\nThe structured workflow turns evidence handling and QA thinking into auditable artifacts.\n\n"
                "# Introduction\n\nThe structured workflow project centers evidence handling, QA thinking, and auditable artifacts.\n",
                encoding="utf-8",
            )
            state.drafts["merged_draft_md"] = str(draft)
            result = run_project_identity_gate(state)
        self.assertTrue(result.runtime["project_identity_report_path"])

    def test_project_identity_gate_accepts_distributed_domain_context(self):
        state = ReportState.new("write admissions report", [], "out")
        state.spec["report_profile"] = "academic_paper"
        state.spec["report_profile"] = "admissions_project_report"
        state.plan["front_matter"] = {
            "title": "Quant: Taiwan Equity Strategy Validation"
        }
        state.plan["thesis_statement"] = (
            "Quant supports Taiwan equities, LLM validation, and historical replay "
            "for a graduate admissions project."
        )
        with tempfile.TemporaryDirectory() as tmpdir:
            run_dir = WORKFLOW_RUNS_DIR / state.job_id
            identity_path = run_dir / "project_identity.json"
            identity_path.write_text(json.dumps({
                "required_terms": ["Quant", "Taiwan equities", "historical replay"],
                "required_context_terms": ["validation"],
                "forbidden_terms": [],
                "canonical_title_terms": ["Quant"],
                "domain_context": "Taiwan equities and graduate admissions project introduction",
                "author_metadata": {},
            }), encoding="utf-8")
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Abstract\n\nQuant is a validation platform for Taiwan equities and historical replay.\n\n"
                "# Introduction\n\nThe project introduces a graduate admissions report about Taiwan equities, "
                "validation, and historical replay.\n",
                encoding="utf-8",
            )
            state.drafts["merged_draft_md"] = str(draft)
            result = run_project_identity_gate(state)
        self.assertTrue(result.runtime["project_identity_report_path"])

    def test_figure_quality_blocks_undeclared_prose_reference(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "academic_paper"
        state.plan["outline"] = {"sections": {"results": {"figure_ids": []}}}
        with tempfile.TemporaryDirectory() as tmpdir:
            merged = Path(tmpdir) / "merged.md"
            merged.write_text("# Results\n\nThe architecture is summarized in Figure 3.\n", encoding="utf-8")
            state.drafts["merged_draft_md"] = str(merged)
            with self.assertRaises(QAHardBlockError):
                run_figure_quality(state)

    def test_admissions_tone_gate_blocks_meta_reader_phrase(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_report"
        with tempfile.TemporaryDirectory() as tmpdir:
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Discussion\n\nFor an admissions committee, the relevant evidence is clear.\n",
                encoding="utf-8",
            )
            state.drafts["publication_style_draft"] = str(draft)
            with self.assertRaises(QAHardBlockError):
                run_admissions_tone_gate(state)

    def test_admissions_tone_gate_uses_project_identity_instead_of_hardcoded_project_terms(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "new_draft"
        state.plan["project_identity"] = {
            "required_terms": ["structured workflow", "evidence handling", "QA thinking"],
            "required_context_terms": ["auditable artifacts"],
            "canonical_title_terms": ["workflow"],
            "domain_context": "technical communication project",
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            draft = Path(tmpdir) / "draft.md"
            draft.write_text(
                "# Conclusion\n\nThe project demonstrates evidence handling, QA thinking, "
                "and structured workflow judgment through auditable artifacts.\n",
                encoding="utf-8",
            )
            state.drafts["publication_style_draft"] = str(draft)
            result = run_admissions_tone_gate(state)
        self.assertTrue(result.runtime["admissions_tone_report_path"])

    def test_reference_relevance_gate_blocks_generic_reference(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_report"
        state.plan["project_identity"] = {
            "required_terms": ["deterministic compilation", "StrategyIR"],
            "required_context_terms": ["compiler architecture"],
            "canonical_title_terms": ["compiler"],
            "forbidden_terms": [],
            "domain_context": "Taiwan equities",
            "author_metadata": {},
        }
        state.plan["claim_matrix"] = {
            "claims": [{
                "claim_id": "c1",
                "claim_text": "StrategyIR supports deterministic compiler architecture.",
                "topic_tags": ["compiler architecture"],
            }]
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            report = Path(tmpdir) / "reference_verify_report.json"
            report.write_text(json.dumps({
                "references": [{
                    "ref_id": "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine.",
                    "status": "verified",
                    "checks": [{"type": "doi", "verified": True}],
                }]
            }), encoding="utf-8")
            state.runtime["reference_verify_report_path"] = str(report)
            with self.assertRaises(QAHardBlockError):
                run_reference_relevance_gate(state)

    def test_reference_relevance_gate_accepts_role_supported_reference(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_report"
        with tempfile.TemporaryDirectory() as tmpdir:
            report = Path(tmpdir) / "reference_verify_report.json"
            report.write_text(json.dumps({
                "references": [{
                    "ref_id": "Cytron, R. (1991). Efficiently computing static single assignment form.",
                    "status": "verified",
                    "reference_role": "compiler_foundation",
                    "supports_claim_ids": ["c1"],
                    "checks": [{"type": "doi", "verified": True}],
                }]
            }), encoding="utf-8")
            state.runtime["reference_verify_report_path"] = str(report)
            result = run_reference_relevance_gate(state)
        self.assertTrue(result.runtime["reference_relevance_report_path"])

    def test_reference_relevance_gate_blocks_internal_only_admissions_without_name_error(self):
        state = ReportState.new("report", [], "out")
        state.spec["report_profile"] = "admissions_report"
        state.spec["task_intent"] = "new_draft"
        with tempfile.TemporaryDirectory() as tmpdir:
            report = Path(tmpdir) / "reference_verify_report.json"
            report.write_text(json.dumps({
                "references": [{
                    "ref_id": "Internal project note.",
                    "raw": "Internal project note.",
                    "status": "excluded",
                    "checks": [{"type": "internal"}],
                }]
            }), encoding="utf-8")
            state.runtime["reference_verify_report_path"] = str(report)
            with self.assertRaises(QAHardBlockError) as raised:
                run_reference_relevance_gate(state)
        self.assertIn("bibliography is not publication-bearing", str(raised.exception))


class DeliveredHeadingAndProvenanceTests(unittest.TestCase):
    """Two things the reader sees that named the tool instead of the subject."""

    @staticmethod
    def _blueprint_section_ids() -> set[str]:
        import yaml

        from report_workflow import blueprints

        section_ids: set[str] = set()
        for path in Path(blueprints.__file__).parent.glob("*.yaml"):
            data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
            section_ids |= set((data.get("sections") or {}).keys())
        return section_ids

    def test_no_blueprint_section_id_survives_as_a_visible_heading(self):
        """A draft that opens with its own id must not ship that id as a heading.

        `## executive_summary` reached a delivered document because the strip
        compared against "Executive Summary" and "executive summary" and never
        against the id itself. Every multi-word id in every blueprint was
        exposed; single-word ids were not, which is why it took a real run to
        find.
        """
        from report_workflow.nodes.merge_draft import _strip_duplicate_title_heading

        section_ids = self._blueprint_section_ids()
        self.assertTrue(section_ids, "no blueprint sections found to check")
        for section_id in sorted(section_ids):
            with self.subTest(section_id=section_id):
                content = f"# {section_id}\n\nBody text.\n"
                stripped = _strip_duplicate_title_heading(
                    content, section_id, {}, section_id.replace("_", " ").title()
                )
                self.assertNotIn(
                    section_id,
                    stripped,
                    f"{section_id!r} survives into the merged draft as a heading",
                )

    def test_table_provenance_names_its_file_once(self):
        """One table, one attribution — not "amazon.csv amazon.csv (544 rows)".

        The span the ledger records already opens with the file name, so
        prefixing it printed the file twice under every table in the document.
        """
        from report_workflow.nodes.source_tables import provenance_line

        table = {
            "source_file_name": "amazon_classified.csv",
            "source_span": "amazon_classified.csv (544 rows)",
        }
        for language in ("zh", "en"):
            with self.subTest(language=language):
                line = provenance_line(table, language)
                self.assertEqual(
                    line.count("amazon_classified.csv"),
                    1,
                    f"provenance line repeats the file name: {line!r}",
                )
                self.assertIn("(544 rows)", line)

    def test_the_source_list_names_its_file_once_too(self):
        """The same duplication, one renderer over.

        The table line was repaired and the source list was not, so every
        `[S<n>]` entry in the delivered document still read
        "amazon_classified.csv amazon_classified.csv (544 rows)".
        """
        from report_workflow.nodes.citation_bind import _format_source_trace_entry

        line = _format_source_trace_entry(
            {
                "source_file_name": "amazon_classified.csv",
                "source_span": "amazon_classified.csv (544 rows)",
                "quote": "544 listings",
            },
            1,
        )
        self.assertEqual(line.count("amazon_classified.csv"), 1, line)
        self.assertIn("(544 rows)", line)

    def test_the_source_list_still_names_the_file_when_the_span_omits_it(self):
        from report_workflow.nodes.citation_bind import _format_source_trace_entry

        line = _format_source_trace_entry(
            {"source_file_name": "notes.md", "source_span": "lines 4-9"}, 2
        )
        self.assertEqual(line, "[S2] notes.md lines 4-9")

    def test_table_provenance_still_prefixes_a_span_without_the_file_name(self):
        """The prefix is dropped only when the span already carries the name."""
        from report_workflow.nodes.source_tables import provenance_line

        line = provenance_line(
            {"source_file_name": "amazon_classified.csv", "source_span": "rows 1-544"},
            "en",
        )
        self.assertEqual(line, "Source: amazon_classified.csv rows 1-544")


if __name__ == "__main__":
    unittest.main()
