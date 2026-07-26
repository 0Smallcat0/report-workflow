"""Tests for base_document_diff and revision_apply enhancements."""
import json
import tempfile
import unittest
from pathlib import Path

from report_workflow.nodes.base_document_diff import (
    compute_revision_diff,
    detect_overlapping_changes,
    compute_section_diff_summary,
    write_diff_report,
)
from report_workflow.agent_wrapper import (
    submit_revision_plan,
    preview_revision_diff,
)
from report_workflow.state import register_job_run


class BaseDocumentImageCarryTests(unittest.TestCase):
    """A figure is content the author put there, so a revision must keep it.

    Reading only `w:t` text nodes dropped every embedded image: changing one
    word in a report deleted its chart and left the caption standing over
    nothing — the same shape as a References heading with no entries.
    """

    def _fixture(self, tmpdir):
        from docx import Document
        from docx.shared import Inches

        import base64

        png = Path(tmpdir) / "chart.png"
        # Smallest valid PNG: 1x1, transparent.
        png.write_bytes(base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk"
            "YPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
        ))
        doc = Document()
        doc.add_heading("1. 結果", level=1)
        doc.add_paragraph("量測結果如下。")
        doc.add_picture(str(png), width=Inches(2))
        doc.add_paragraph("圖 1. 效能對流量。")
        path = str(Path(tmpdir) / "base.docx")
        doc.save(path)
        return path

    def test_image_is_extracted_and_linked_in_place(self):
        from report_workflow.nodes.base_document_parse import _parse_docx_section

        with tempfile.TemporaryDirectory() as tmpdir:
            media_dir = Path(tmpdir) / "base_media"
            sections, _titles = _parse_docx_section(
                self._fixture(tmpdir), media_dir=media_dir
            )
            carrying = [sid for sid, text in sections.items() if "![](" in text]
            self.assertEqual(len(carrying), 1, f"expected one section with an image: {sections}")
            body = sections[carrying[0]]
            self.assertLess(body.index("![]("), body.index("圖 1."),
                            "image must stay above its caption")
            self.assertTrue(list(media_dir.iterdir()), "no media was extracted")

    def test_without_media_dir_nothing_is_extracted(self):
        """The default stays text-only, so callers that never render images
        (and the .md/.txt paths) are unaffected."""
        from report_workflow.nodes.base_document_parse import _parse_docx_section

        with tempfile.TemporaryDirectory() as tmpdir:
            sections, _titles = _parse_docx_section(self._fixture(tmpdir))
            self.assertFalse([t for t in sections.values() if "![](" in t])


class GeneratedTocIngestTests(unittest.TestCase):
    """Reading back our own output must not re-ingest our own scaffolding.

    A rendered report's front matter is title page, then the generated table of
    contents. Revising it captured both into the preamble while the next render
    added a fresh TOC, so the scaffolding accumulated one copy per revision.
    """

    def _drop(self, text):
        from report_workflow.nodes.base_document_parse import _drop_generated_toc

        return _drop_generated_toc(text)

    def test_generated_toc_is_removed_and_title_page_kept(self):
        from report_workflow.nodes.docx_render import _TOC_PLACEHOLDERS, _TOC_TITLES

        preamble = "\n".join([
            "熱傳學實驗", "國立成功大學機械工程學系", "蔡知均", "2026 年 7 月",
            _TOC_TITLES["zh"], _TOC_PLACEHOLDERS["zh"],
        ])
        cleaned = self._drop(preamble)
        self.assertEqual(cleaned.split("\n"),
                         ["熱傳學實驗", "國立成功大學機械工程學系", "蔡知均", "2026 年 7 月"])

    def test_english_scaffolding_is_removed_too(self):
        from report_workflow.nodes.docx_render import _TOC_PLACEHOLDERS, _TOC_TITLES

        cleaned = self._drop(
            f"Heat Transfer Lab\nJ. Tsai\n{_TOC_TITLES['en']}\n{_TOC_PLACEHOLDERS['en']}"
        )
        self.assertEqual(cleaned.split("\n"), ["Heat Transfer Lab", "J. Tsai"])

    def test_a_real_heading_named_like_the_toc_is_kept(self):
        """"目錄" is an ordinary word. Without the placeholder under it, an
        author's own line by that name is theirs, not ours."""
        cleaned = self._drop("報告封面\n目錄\n本章說明目錄的編排方式。")
        self.assertIn("目錄", cleaned.split("\n"))

    def test_repeated_revision_does_not_accumulate(self):
        from report_workflow.nodes.docx_render import _TOC_PLACEHOLDERS, _TOC_TITLES

        once = self._drop(f"封面\n{_TOC_TITLES['zh']}\n{_TOC_PLACEHOLDERS['zh']}")
        twice = self._drop(f"{once}\n{_TOC_TITLES['zh']}\n{_TOC_PLACEHOLDERS['zh']}")
        self.assertEqual(once, twice)


class BaseDocumentHeadingRoundTripTests(unittest.TestCase):
    """Revising a document must not rewrite its headings.

    Section ids are slugs of the heading text, and the titles map was built as
    an identity over those slugs, so every numbered heading came back out with
    an underscore in it: "1. 實驗目的" rendered as "1._實驗目的". Ten headings in
    a real lab report, in the DOCX the author would hand in.
    """

    HEADINGS = ["摘要", "1. 實驗目的", "2. 需求與規格矩陣", "10. 附錄"]

    def _parse(self, tmpdir):
        from docx import Document

        from report_workflow.nodes.base_document_parse import _parse_docx_section

        doc = Document()
        for heading in self.HEADINGS:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(f"{heading} 的內文。")
        path = str(Path(tmpdir) / "base.docx")
        doc.save(path)
        return _parse_docx_section(path)

    def test_titles_keep_the_heading_text_verbatim(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            _sections, titles = self._parse(tmpdir)
            self.assertEqual(sorted(titles.values()), sorted(self.HEADINGS))

    def test_no_title_carries_a_slug_underscore(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            _sections, titles = self._parse(tmpdir)
            self.assertFalse(
                [title for title in titles.values() if "_" in title],
                f"slug leaked into a display title: {titles}",
            )

    def test_every_addressable_section_has_a_title(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            sections, titles = self._parse(tmpdir)
            addressable = [sid for sid in sections if sid != "preamble"]
            self.assertTrue(addressable)
            for section_id in addressable:
                self.assertIn(section_id, titles)


class ComputeRevisionDiffTests(unittest.TestCase):
    """Test compute_revision_diff function."""

    def setUp(self):
        self.base = {
            "introduction": "The system uses AST parsing to build intermediate representations.",
            "results": "We analyzed 388 files with 5,171 nodes and 9,764 edges.",
            "discussion": "The hub structure suggests a modular design pattern.",
        }

    def test_all_valid_changes(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "388 files",
                "new_text": "400 files",
            },
            {
                "section_id": "discussion",
                "change_type": "replace",
                "original_text": "modular design",
                "new_text": "highly modular design",
            },
        ]}
        result = compute_revision_diff(self.base, plan)
        self.assertEqual(result["total_changes"], 2)
        self.assertEqual(result["valid_changes"], 2)
        self.assertEqual(len(result["unresolvable"]), 0)
        self.assertEqual(len(result["conflicts"]), 0)

    def test_unresolvable_original_text(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "THIS TEXT DOES NOT EXIST",
                "new_text": "replacement",
            },
        ]}
        result = compute_revision_diff(self.base, plan)
        self.assertEqual(result["valid_changes"], 0)
        self.assertEqual(len(result["unresolvable"]), 1)
        self.assertIn("not found", result["unresolvable"][0]["reason"])

    def test_unknown_section(self):
        plan = {"changes": [
            {
                "section_id": "nonexistent_section",
                "change_type": "replace",
                "original_text": "foo",
                "new_text": "bar",
            },
        ]}
        result = compute_revision_diff(self.base, plan)
        self.assertEqual(len(result["unresolvable"]), 1)
        self.assertIn("not found", result["unresolvable"][0]["reason"])

    def test_unknown_change_type(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "merge",
                "original_text": "388",
                "new_text": "400",
            },
        ]}
        result = compute_revision_diff(self.base, plan)
        self.assertEqual(len(result["unresolvable"]), 1)

    def test_context_preview_populated(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "5,171 nodes",
                "new_text": "5,200 nodes",
            },
        ]}
        result = compute_revision_diff(self.base, plan)
        preview = result["preview"][0]
        self.assertEqual(preview["status"], "valid")
        self.assertTrue(len(preview["context_before"]) > 0)


class DetectOverlappingChangesTests(unittest.TestCase):
    """Test detect_overlapping_changes function."""

    def test_no_overlap_different_sections(self):
        changes = [
            {"section_id": "intro", "change_type": "replace", "original_text": "hello"},
            {"section_id": "results", "change_type": "replace", "original_text": "world"},
        ]
        base = {"intro": "hello world", "results": "hello world"}
        conflicts = detect_overlapping_changes(changes, base)
        self.assertEqual(len(conflicts), 0)

    def test_overlap_same_section(self):
        changes = [
            {"section_id": "results", "change_type": "replace", "original_text": "hello world"},
            {"section_id": "results", "change_type": "replace", "original_text": "world foo"},
        ]
        base = {"results": "hello world foo bar"}
        conflicts = detect_overlapping_changes(changes, base)
        self.assertEqual(len(conflicts), 1)
        self.assertEqual(conflicts[0], (0, 1))

    def test_no_overlap_same_section(self):
        changes = [
            {"section_id": "results", "change_type": "replace", "original_text": "hello"},
            {"section_id": "results", "change_type": "replace", "original_text": "bar"},
        ]
        base = {"results": "hello world foo bar"}
        conflicts = detect_overlapping_changes(changes, base)
        self.assertEqual(len(conflicts), 0)

    def test_no_base_sections_returns_empty(self):
        conflicts = detect_overlapping_changes([{"section_id": "a"}], None)
        self.assertEqual(conflicts, [])


class ComputeSectionDiffSummaryTests(unittest.TestCase):
    """Test compute_section_diff_summary function."""

    def test_identical_text(self):
        text = "Line 1\nLine 2\nLine 3"
        result = compute_section_diff_summary(text, text)
        self.assertEqual(result["similarity_ratio"], 1.0)
        self.assertEqual(result["added_lines"], 0)
        self.assertEqual(result["removed_lines"], 0)

    def test_small_change_high_ratio(self):
        old = "Line 1\nLine 2\nLine 3\nLine 4\nLine 5"
        new = "Line 1\nLine 2 modified\nLine 3\nLine 4\nLine 5"
        result = compute_section_diff_summary(old, new)
        self.assertGreater(result["similarity_ratio"], 0.7)

    def test_large_change_low_ratio(self):
        old = "Original line 1\nOriginal line 2\nOriginal line 3"
        new = "Completely different A\nCompletely different B\nCompletely different C"
        result = compute_section_diff_summary(old, new)
        self.assertLess(result["similarity_ratio"], 0.3)

    def test_added_lines_counted(self):
        old = "Line 1"
        new = "Line 1\nLine 2\nLine 3"
        result = compute_section_diff_summary(old, new)
        self.assertGreater(result["added_lines"], 0)


class WriteDiffReportTests(unittest.TestCase):
    """Test write_diff_report function."""

    def test_writes_json(self):
        job_id = "test_diff_report"
        run_dir = Path(tempfile.mkdtemp()) / f"diff-report--{job_id}"
        run_dir.mkdir(parents=True, exist_ok=True)
        register_job_run(job_id, run_dir)

        try:
            path = write_diff_report(job_id, {"test": True})
            self.assertTrue(Path(path).exists())
            with open(path) as f:
                data = json.load(f)
            self.assertTrue(data["test"])
        finally:
            import shutil
            shutil.rmtree(run_dir, ignore_errors=True)


class AgentToolRevisionTests(unittest.TestCase):
    """Test submit_revision_plan and preview_revision_diff agent tools."""

    def setUp(self):
        self.job_id = "test_revision_tools"
        self.tmpdir = tempfile.mkdtemp()
        self.run_dir = Path(self.tmpdir) / f"revision-tools--{self.job_id}"
        self.run_dir.mkdir(parents=True, exist_ok=True)
        register_job_run(self.job_id, self.run_dir)

        # Write base_document_sections
        base = {"results": "We found 388 files with 5,171 nodes."}
        (self.run_dir / "base_document_sections.json").write_text(
            json.dumps(base), encoding="utf-8"
        )

    def tearDown(self):
        import shutil
        shutil.rmtree(self.tmpdir, ignore_errors=True)

    def test_submit_validates_valid_plan(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "388 files",
                "new_text": "400 files",
            },
        ]}
        (self.run_dir / "revision_plan.json").write_text(
            json.dumps(plan), encoding="utf-8"
        )
        result = submit_revision_plan(self.job_id)
        self.assertEqual(result["status"], "ok")
        self.assertIn("preview", result)

    def test_submit_detects_invalid_plan(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "NONEXISTENT TEXT",
                "new_text": "replacement",
            },
        ]}
        (self.run_dir / "revision_plan.json").write_text(
            json.dumps(plan), encoding="utf-8"
        )
        result = submit_revision_plan(self.job_id)
        self.assertEqual(result["status"], "validation_failed")

    def test_preview_is_readonly(self):
        plan = {"changes": [
            {
                "section_id": "results",
                "change_type": "replace",
                "original_text": "388 files",
                "new_text": "400 files",
            },
        ]}
        (self.run_dir / "revision_plan.json").write_text(
            json.dumps(plan), encoding="utf-8"
        )
        result = preview_revision_diff(self.job_id)
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["valid_changes"], 1)
        # Should NOT write any report file (readonly)
        self.assertNotIn("diff_report_path", result)

    def test_missing_base_document(self):
        (self.run_dir / "base_document_sections.json").unlink()
        result = submit_revision_plan(self.job_id)
        self.assertEqual(result["status"], "error")

    def test_missing_revision_plan(self):
        result = submit_revision_plan(self.job_id)
        self.assertEqual(result["status"], "error")


if __name__ == "__main__":
    unittest.main()
