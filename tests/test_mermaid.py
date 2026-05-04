"""Tests for Mermaid diagram auto-conversion in docx_render."""
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch, MagicMock

from docx import Document

from report_workflow.nodes.docx_render import (
    _find_mmdc,
    _convert_mermaid_blocks,
    _MERMAID_BLOCK_RE,
    _replace_figure_placeholders,
    run_docx_render,
)
from report_workflow.state import ReportState, WORKFLOW_RUNS_DIR


class MermaidBlockRegexTests(unittest.TestCase):
    """Test the mermaid code fence regex."""

    def test_matches_simple_block(self):
        md = '```mermaid\ngraph LR\n  A --> B\n```'
        matches = _MERMAID_BLOCK_RE.findall(md)
        self.assertEqual(len(matches), 1)
        self.assertIn("graph LR", matches[0])

    def test_matches_multiple_blocks(self):
        md = '```mermaid\ngraph LR\n  A-->B\n```\n\nSome text.\n\n```mermaid\nsequenceDiagram\n  A->>B: Hello\n```'
        matches = _MERMAID_BLOCK_RE.findall(md)
        self.assertEqual(len(matches), 2)

    def test_no_match_for_other_code_fences(self):
        md = '```python\nprint("hello")\n```'
        matches = _MERMAID_BLOCK_RE.findall(md)
        self.assertEqual(len(matches), 0)

    def test_no_match_for_plain_text(self):
        md = 'This is just text with no code fences.'
        matches = _MERMAID_BLOCK_RE.findall(md)
        self.assertEqual(len(matches), 0)


class FindMmdcTests(unittest.TestCase):
    """Test _find_mmdc discovery."""

    @patch("report_workflow.nodes.docx_render.Path.exists", return_value=False)
    @patch("shutil.which", return_value=None)
    def test_returns_none_when_not_installed(self, mock_which, mock_exists):
        result = _find_mmdc()
        self.assertIsNone(result)

    @patch("shutil.which", return_value="/usr/local/bin/mmdc")
    def test_returns_path_when_on_path(self, mock_which):
        result = _find_mmdc()
        self.assertEqual(result, "/usr/local/bin/mmdc")


class ConvertMermaidBlocksTests(unittest.TestCase):
    """Test _convert_mermaid_blocks function."""

    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp())

    def test_no_mermaid_blocks_passthrough(self):
        md = "# Title\n\nSome text with no mermaid.\n"
        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(result, md)
        self.assertEqual(count, 0)

    @patch("report_workflow.nodes.docx_render._find_mmdc", return_value=None)
    def test_mmdc_not_installed_passthrough(self, mock_find):
        md = '```mermaid\ngraph LR\n  A-->B\n```'
        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(result, md)
        self.assertEqual(count, 0)

    @patch("report_workflow.nodes.docx_render._find_mmdc", return_value="/usr/bin/mmdc")
    @patch("subprocess.run")
    def test_successful_conversion(self, mock_run, mock_find):
        md = 'Before\n\n```mermaid\ngraph LR\n  A-->B\n```\n\nAfter'

        # Mock successful mmdc run
        mock_run.return_value = MagicMock(returncode=0, stderr="")

        # Create the expected output PNG so the check passes
        figures_dir = self.tmpdir / "figures"
        figures_dir.mkdir(parents=True, exist_ok=True)
        (figures_dir / "mermaid_figure_1.png").write_bytes(b"fake png")

        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(count, 1)
        self.assertIn("![Figure 1]", result)
        self.assertNotIn("```mermaid", result)
        self.assertIn("Before", result)
        self.assertIn("After", result)

    @patch("report_workflow.nodes.docx_render._find_mmdc", return_value="/usr/bin/mmdc")
    @patch("subprocess.run")
    def test_failed_conversion_preserves_block(self, mock_run, mock_find):
        md = '```mermaid\ngraph LR\n  A-->B\n```'

        # Mock failed mmdc run
        mock_run.return_value = MagicMock(returncode=1, stderr="error")

        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(count, 0)
        self.assertIn("```mermaid", result)

    @patch("report_workflow.nodes.docx_render._find_mmdc", return_value="/usr/bin/mmdc")
    @patch("subprocess.run")
    def test_multiple_blocks_numbered(self, mock_run, mock_find):
        md = '```mermaid\ngraph A\n```\n\n```mermaid\ngraph B\n```\n\n```mermaid\ngraph C\n```'

        mock_run.return_value = MagicMock(returncode=0, stderr="")

        # Create all expected PNGs
        figures_dir = self.tmpdir / "figures"
        figures_dir.mkdir(parents=True, exist_ok=True)
        for i in range(1, 4):
            (figures_dir / f"mermaid_figure_{i}.png").write_bytes(b"fake")

        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(count, 3)
        self.assertIn("![Figure 1]", result)
        self.assertIn("![Figure 2]", result)
        self.assertIn("![Figure 3]", result)

    @patch("report_workflow.nodes.docx_render._find_mmdc", return_value="/usr/bin/mmdc")
    @patch("subprocess.run", side_effect=Exception("subprocess crash"))
    def test_exception_preserves_block(self, mock_run, mock_find):
        md = '```mermaid\ngraph LR\n  A-->B\n```'
        result, count = _convert_mermaid_blocks(md, self.tmpdir)
        self.assertEqual(count, 0)
        self.assertIn("```mermaid", result)


class FigurePlaceholderReplacementTests(unittest.TestCase):
    """Test figure manifest placeholder conversion for DOCX rendering."""

    def setUp(self):
        self.tmpdir = Path(tempfile.mkdtemp())

    def test_replaces_known_placeholder_with_markdown_image(self):
        image_path = self.tmpdir / "figure_1.png"
        image_path.write_bytes(b"fake png")
        manifest = {
            "figures": [
                {
                    "figure_id": "1",
                    "title": "Voltage trend",
                    "path": str(image_path),
                }
            ]
        }
        md = "# Results\n\n[FIGURE:1]\n\nFigure 1: Voltage trend.\n"

        result, replaced, unresolved = _replace_figure_placeholders(md, manifest)

        self.assertEqual(replaced, 1)
        self.assertEqual(unresolved, [])
        self.assertNotIn("[FIGURE:1]", result)
        self.assertIn("![Figure 1. Voltage trend]", result)
        self.assertIn(image_path.resolve().as_posix(), result)

    def test_unknown_placeholder_is_preserved_for_validation(self):
        result, replaced, unresolved = _replace_figure_placeholders("[FIGURE:2]", {"figures": []})

        self.assertEqual(replaced, 0)
        self.assertEqual(unresolved, ["2"])
        self.assertEqual(result, "[FIGURE:2]")

    @patch("report_workflow.nodes.docx_render._validate_docx", return_value=[])
    @patch("report_workflow.nodes.docx_render._style_tables_post_render", return_value=None)
    @patch("report_workflow.nodes.docx_render._repair_missing_figures", return_value=0)
    def test_run_docx_render_writes_resolved_placeholder_to_pandoc_input(
        self,
        mock_repair,
        mock_style,
        mock_validate,
    ):
        state = ReportState.new("report", [], "out")
        state.qa["qa_decision"] = "pass"
        state.spec["report_profile"] = "business_report"
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        run_dir.mkdir(parents=True, exist_ok=True)

        draft_path = run_dir / "publication.md"
        draft_path.write_text(
            "# Results\n\n[FIGURE:1]\n\nFigure 1: Voltage trend.\n\n"
            "The measurement is shown in Figure 1.\n",
            encoding="utf-8",
        )
        state.drafts["publication_style_draft"] = str(draft_path)

        image_path = run_dir / "figures" / "figure_1.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(b"fake png")
        manifest_path = run_dir / "figure_manifest.json"
        manifest_path.write_text(
            json.dumps(
                {
                    "generated_count": 1,
                    "figures": [
                        {
                            "figure_id": "1",
                            "title": "Voltage trend",
                            "path": str(image_path),
                        }
                    ],
                }
            ),
            encoding="utf-8",
        )
        state.output["figure_manifest_path"] = str(manifest_path)

        def fake_pandoc(md_path, output_path, toc=True, number_sections=False):
            doc = Document()
            doc.add_heading("Results", level=1)
            doc.add_paragraph("Rendered body.")
            doc.add_paragraph("Figure 1: Voltage trend.")
            doc.save(output_path)
            return True

        with patch("report_workflow.nodes.docx_render._render_via_pandoc", side_effect=fake_pandoc):
            result = run_docx_render(state)

        pandoc_input = (run_dir / "pandoc_input.md").read_text(encoding="utf-8")
        self.assertIn("![Figure 1. Voltage trend]", pandoc_input)
        self.assertNotIn("[FIGURE:1]", pandoc_input)
        self.assertEqual(result.output.get("figure_placeholders_resolved"), 1)


if __name__ == "__main__":
    unittest.main()
