"""Regression tests: a rendered report must never end with a dangling
"References" heading that has nothing under it.

Found by auditing a rendered benchmark report (2026-07-15): the fixture drafts
carried no references, the reference list file was empty, and the document
still rendered a bare "References" heading as its final paragraph. Both render
paths are covered — the markdown pre-pass that pandoc consumes
(``_split_body_references``) and the python-docx fallback
(``_add_hanging_indent_references``).
"""
import unittest

from docx import Document

from report_workflow.nodes.docx_render import (
    _add_hanging_indent_references,
    _split_body_references,
)


class ReferenceHeadingLevelTests(unittest.TestCase):
    """The generated reference list is a top-level section like any other.

    Found by rendering a real lab report (2026-07-27): every section came out
    at Heading 1 and the appended references at Heading 2, so a Word table of
    contents nested them under the appendix. The level was pinned in three
    places at once — the writers, this matcher, and its `add_heading` call.
    """

    ENTRIES = (
        "\n\n- Doe, J. (2025). A real source. Journal of Examples.\n"
        "- Roe, R. (2024). Another source. Example Press.\n"
    )

    def _heading(self, ref_md):
        doc = Document()
        _add_hanging_indent_references(doc, ref_md)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        self.assertEqual(len(headings), 1, "expected exactly one References heading")
        return headings[0]

    def test_generated_list_renders_at_its_declared_level(self):
        self.assertEqual(self._heading("# References" + self.ENTRIES).style.name, "Heading 1")

    def test_authored_level_is_preserved_rather_than_forced(self):
        self.assertEqual(self._heading("### References" + self.ENTRIES).style.name, "Heading 3")

    def test_writers_agree_on_one_heading(self):
        from report_workflow.nodes.citation_bind import REFERENCE_LIST_HEADING

        self.assertTrue(REFERENCE_LIST_HEADING.startswith("# "))
        self.assertEqual(
            self._heading(REFERENCE_LIST_HEADING + self.ENTRIES).style.name, "Heading 1"
        )


class SplitBodyReferencesTests(unittest.TestCase):
    def test_empty_trailing_references_section_is_removed(self):
        md = "# Title\n\nBody prose.\n\n## References\n"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(refs_md, "")
        self.assertNotIn("References", cleaned)
        self.assertIn("Body prose.", cleaned)

    def test_whitespace_only_references_section_is_removed(self):
        md = "# Title\n\nBody prose.\n\n## References\n\n   \n\n"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(refs_md, "")
        self.assertNotIn("References", cleaned)

    def test_populated_references_section_is_extracted(self):
        md = (
            "# Title\n\nBody prose.\n\n## References\n\n"
            "- Doe, J. (2025). A real source. Journal of Examples.\n"
            "- Roe, R. (2024). Another source. Example Press.\n"
        )
        cleaned, refs_md = _split_body_references(md)
        self.assertNotIn("## References", cleaned)
        self.assertIn("## References", refs_md)
        self.assertIn("Doe, J. (2025)", refs_md)
        self.assertIn("Roe, R. (2024)", refs_md)

    def test_document_without_references_is_unchanged(self):
        md = "# Title\n\nBody prose only.\n"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(cleaned, md)
        self.assertEqual(refs_md, "")

    def test_mid_document_empty_references_does_not_swallow_next_section(self):
        md = "# Title\n\n## References\n\n## Appendix\n\nAppendix body.\n"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(refs_md, "")
        self.assertIn("## Appendix", cleaned)
        self.assertIn("Appendix body.", cleaned)
        self.assertNotIn("## References", cleaned)

    def test_h1_empty_references_is_removed(self):
        # Upstream drafts carry "# References" (H1) before heading
        # normalization; the empty-section guard must catch that level too.
        md = "# Title\n\nBody prose.\n\n# References\n"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(refs_md, "")
        self.assertNotIn("References", cleaned)

    def test_empty_references_at_eof_without_newline_is_removed(self):
        # Real drafts can end exactly at the heading with no trailing newline
        # (the case that actually shipped a dangling heading).
        md = "# Title\n\nBody prose.\n\n# References"
        cleaned, refs_md = _split_body_references(md)
        self.assertEqual(refs_md, "")
        self.assertNotIn("References", cleaned)

    def test_h1_populated_references_are_extracted(self):
        md = (
            "# Title\n\nBody.\n\n# References\n\n"
            "- Doe, J. (2025). A real source. Journal of Examples.\n"
        )
        cleaned, refs_md = _split_body_references(md)
        self.assertNotIn("References", cleaned)
        self.assertIn("Doe, J. (2025)", refs_md)


class HangingIndentReferencesTests(unittest.TestCase):
    def test_heading_only_reference_md_adds_nothing(self):
        doc = Document()
        before = len(doc.paragraphs)
        _add_hanging_indent_references(doc, "## References\n\n")
        self.assertEqual(len(doc.paragraphs), before)
        self.assertNotIn("References", [p.text for p in doc.paragraphs])

    def test_populated_reference_md_adds_heading_and_entries(self):
        doc = Document()
        _add_hanging_indent_references(
            doc,
            "## References\n\n- Doe, J. (2025). A real source.\n",
        )
        texts = [p.text for p in doc.paragraphs]
        self.assertIn("References", texts)
        self.assertTrue(any("Doe, J. (2025)" in t for t in texts))


if __name__ == "__main__":
    unittest.main()
