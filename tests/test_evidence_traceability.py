"""Traceability, evidence identity, and figure-selection regressions.

Each case here comes from one run of a real 16 KB market report through the
business_report profile. The pipeline completed and published; what it
published could not be checked by its reader, and several of its gates were
answering about the wrong row.
"""
import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from report_workflow.agent_wrapper import query_evidence
from report_workflow.errors import QAHardBlockError
from report_workflow.nodes.citation_bind import (
    SOURCE_LIST_HEADING,
    resolve_citations_publication,
)
from report_workflow.nodes.figure_recommend import recommend_figures_from_evidence
from report_workflow.nodes.source_tables import (
    collect_source_tables,
    replace_table_placeholders,
)
from report_workflow.nodes.outline_plan import run_outline_plan
from report_workflow.preflight_decisions import feature_flags_from_decisions
from report_workflow.run_workflow import prepare_workflow
from report_workflow.state import ReportState, WORKFLOW_RUNS_DIR


#: A source shaped like the one that produced these defects: several
#: independent categories, prose that names its houses, hedged figures, and a
#: table per category. Written out at roughly the same paragraph length,
#: because the ledger's per-row overhead only amortises against real prose.
MARKET_REPORT = """# 回收利用經濟性分析

## 鋰電池

濕法冶金製程的鋰鈷鎳回收率可達 95% 以上（來源：IMARC，2026），但實際良率取決於
前處理放電與破碎環節的雜質控制，不同廠區之間的差距可以達到十個百分點以上，因此
單一數字不足以代表整個品類的回收表現。處理成本則受碳酸鋰價格週期支配，價格低檔
時黑粉的收購意願明顯下降，見下表。

| 時點 | 價格（USD/噸） |
|---|---|
| 2022-11 峰值 | ~80,000–85,000 |
| 2025-06 低點 | ~8,259 |
| 2026-02-05 | ~19,800 |
| 2026-07 | ~21,570–21,760 |

方法：以每噸黑粉處理成本乘以年處理量推算年度毛利，再扣除前處理與環保法遵支出；
推算採用各廠區公開年報（來源：各公司年報，2026）揭露的產能利用率，而非設計產能，
否則會系統性高估收益。

## 塑膠

機械回收的 PET 再生料價格長期貼著原生料走（來源：OECD，2026），價差收斂時再生
料需求隨即轉弱，這使得再生廠的毛利對原油價格的敏感度高於對回收量的敏感度。化學
回收雖然可以處理混合料，但單位能耗仍高，尚未在無補貼條件下達到規模經濟。

| 製程 | 產出率 | 噸成本 |
|---|---|---|
| 機械回收 | 72% | 210 |
| 溶劑回收 | 58% | 380 |
| 化學解聚 | 46% | 540 |

## 紡織

紡織品回收的瓶頸在於分選而非處理（來源：Ellen MacArthur Foundation，2026）。
混紡材質無法以單一製程處理，人工分選成本佔總成本的比重過半，自動化光譜分選設備
的投資回收期在目前的處理費水準下仍偏長。

| 材質 | 分選準確率 | 噸成本 |
|---|---|---|
| 純棉 | 91% | 145 |
| 聚酯 | 88% | 168 |
| 混紡 | 54% | 315 |

## 紙張

舊瓦楞紙的到廠價格由區域供需決定（來源：Fastmarkets，2026），出口導向的港口城
市與內陸城市之間常年存在價差，這個價差比品質分級造成的價差更大。白紙等級的回收
料因為印刷油墨較少，脫墨成本低，毛利率長期高於混合紙。

| 品類 | 回收率 | 噸成本 |
|---|---|---|
| 舊瓦楞紙 | 68% | 92 |
| 混合紙 | 41% | 138 |
| 白紙 | 77% | 61 |
"""


def _all_packages_present(name: str, *args, **kwargs):
    return object()


def _prepare_market_report(tmpdir: str) -> tuple[ReportState, Path]:
    src = Path(tmpdir) / "recycling.md"
    src.write_text(MARKET_REPORT, encoding="utf-8")
    with patch(
        "report_workflow.preflight.importlib.util.find_spec",
        side_effect=_all_packages_present,
    ):
        state = prepare_workflow(
            "analyse recycling economics",
            [str(src)],
            str(Path(tmpdir) / "out"),
            report_profile="business_report",
        )
    return state, src


def _ledger(state: ReportState) -> list[dict]:
    path = Path(state.sources["evidence_ledger_path"])
    return [
        json.loads(line)
        for line in path.read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


class EvidenceIdentityTests(unittest.TestCase):
    """Every row of a table used to answer to one id and one hash.

    Four rows of a price table shared `E_2f8b56ad_34113958b0`, so the gate that
    resolves a cited id read whichever row came first: a statistical claim
    citing the quantitative row was refused because a qualitative row of the
    same table answered for it, and the message named an id the author had no
    way to disambiguate.
    """

    def test_table_rows_get_distinct_ids_and_hashes(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = [row for row in _ledger(state) if row.get("block_type") == "table_row"]
            self.assertGreaterEqual(len(rows), 6, "the two source tables should yield rows")

            ids = [row["evidence_id"] for row in rows]
            hashes = [row["content_hash"] for row in rows]
            self.assertEqual(len(set(ids)), len(ids), "table rows share an evidence_id")
            self.assertEqual(len(set(hashes)), len(hashes), "table rows share a content_hash")

    def test_row_hash_is_of_the_row_not_the_table(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = [row for row in _ledger(state) if row.get("block_type") == "table_row"]
            by_table: dict[str, list[str]] = {}
            for row in rows:
                parent = row["block_id"].rsplit("_r", 1)[0]
                by_table.setdefault(parent, []).append(row["content_hash"])
            for parent, digests in by_table.items():
                self.assertEqual(
                    len(set(digests)), len(digests), f"{parent} reuses one hash across rows"
                )


class LedgerSizeTests(unittest.TestCase):
    """cross_references listed nearly every id on every row.

    A 16 KB source produced a 409 KB ledger, and query_evidence — whose whole
    purpose is answering without loading the ledger — returned the entire id
    space for one row.
    """

    def test_cross_references_are_a_rounding_error_in_the_ledger(self):
        """They were about 40% of it, and rose with the square of the row count."""
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            ledger_path = Path(state.sources["evidence_ledger_path"])
            ledger_size = ledger_path.stat().st_size
            rows = _ledger(state)
            reference_bytes = sum(
                len(json.dumps(row.get("cross_references", []), ensure_ascii=False))
                for row in rows
            )
            self.assertLess(
                reference_bytes,
                ledger_size * 0.1,
                "cross_references are still a tenth of the ledger",
            )

    def test_no_row_links_outside_its_own_table(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = _ledger(state)
            table_of = {
                row["evidence_id"]: row["block_id"].rsplit("_r", 1)[0]
                for row in rows
                if row.get("block_type") == "table_row"
            }
            for row in rows:
                for reference in row.get("cross_references", []):
                    self.assertEqual(
                        table_of.get(reference),
                        table_of.get(row["evidence_id"]),
                        f"{row['evidence_id']} links to a row of another table",
                    )

    def test_cross_references_stay_within_the_table(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = _ledger(state)
            all_ids = {row["evidence_id"] for row in rows}
            for row in rows:
                refs = row.get("cross_references", [])
                self.assertLess(
                    len(refs),
                    len(all_ids) - 1,
                    f"{row['evidence_id']} references the whole ledger",
                )

    def test_query_evidence_omits_cross_references(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            first = _ledger(state)[0]["evidence_id"]
            result = query_evidence(
                state.job_id,
                evidence_ids=[first],
                workspace_root=state.output.get("workspace_root"),
            )
            self.assertEqual(result["status"], "ok")
            self.assertTrue(result["entries"])
            for entry in result["entries"]:
                self.assertNotIn("cross_references", entry)


class EvidenceGradeTests(unittest.TestCase):
    """source_role decided the grade, so one source graded uniformly medium.

    FD forbids `measured` wording below high grade, so a report that named a
    house beside every figure still had to hedge every sentence it published.
    """

    def test_grades_vary_within_one_internal_source(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            grades = {row["evidence_grade"] for row in _ledger(state)}
            self.assertIn("high", grades, "no row reached high grade")
            self.assertGreater(len(grades), 1, "every row graded the same")

    def test_attribution_and_method_are_recorded_per_row(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = _ledger(state)
            self.assertTrue(
                any(row["contains_citations"] for row in rows),
                "no row registered its (來源:...) attribution",
            )
            self.assertTrue(
                any(row["contains_methodology"] for row in rows),
                "no row registered its stated derivation",
            )


class CitationRenderingTests(unittest.TestCase):
    """[CITE:] resolved to the empty string for project sources.

    The delivered DOCX carried no citation marker, no reference entry, and no
    source list: the reader could check no figure in it against anything.
    """

    EVIDENCE = [{
        "evidence_id": "E_recy_001",
        "source_id": "S001",
        "source_role": "internal_project_source",
        "source_file_name": "recycling.md",
        "source_span": "line 38-43",
        "file_type": "md",
        "quote": "濕法冶金製程的鋰鈷鎳回收率可達 95% 以上",
    }]

    def test_marker_is_visible_and_backed_by_a_source_entry(self):
        resolved, _audit, _refs, _internal, source_refs = resolve_citations_publication(
            "濕法冶金回收率可達 95% 以上 [CITE:E_recy_001]。",
            self.EVIDENCE,
            [],
        )
        self.assertIn("[S1]", resolved)
        self.assertNotIn("[CITE:", resolved)
        self.assertEqual(len(source_refs), 1)
        self.assertIn("recycling.md", source_refs[0])
        self.assertIn("line 38-43", source_refs[0])
        self.assertIn("濕法冶金製程的鋰鈷鎳回收率", source_refs[0])
        # The file, the line span and the quote are what a reader can check.
        # The ledger's own handle is internal: it belongs in the audit
        # appendix, not in the delivered source list.
        self.assertNotIn("E_recy_001", source_refs[0])

    def test_repeated_citations_of_one_row_share_a_number(self):
        _resolved, _audit, _refs, _internal, source_refs = resolve_citations_publication(
            "A [CITE:E_recy_001]. B [CITE:E_recy_001].",
            self.EVIDENCE,
            [],
        )
        self.assertEqual(len(source_refs), 1)

    def test_no_space_is_left_before_full_width_punctuation(self):
        resolved, _audit, _refs, _internal, _source_refs = resolve_citations_publication(
            "完全被舊瓦楞紙價格週期支配 [CITE:E_recy_001] 。",
            self.EVIDENCE,
            [],
        )
        self.assertNotIn(" 。", resolved)

    def test_source_list_heading_is_a_top_level_section(self):
        self.assertTrue(SOURCE_LIST_HEADING.startswith("# "))


class FigureSelectionTests(unittest.TestCase):
    """Every recommendation came back `table` with the same confidence.

    "No reliable numeric measure column was detected" was reported for a price
    column reading `~80,000–85,000`: the parser stopped at the tilde and the
    en dash, so an ordered time series lost its chart.
    """

    PRICE_TABLE = [{
        "evidence_id": "E_price",
        "source_id": "S001",
        "source_file_name": "recycling.md",
        "granularity": "table",
        "content": "碳酸鋰價格",
        "table_data": [
            ["時點", "價格（USD/噸）"],
            ["2022-11", "~80,000–85,000"],
            ["2025-06", "~8,259"],
            ["2026-02", "~19,800"],
            ["2026-07", "~21,570–21,760"],
        ],
    }]

    def _state(self, tmpdir: str) -> ReportState:
        state = ReportState.new("chart it", [], str(Path(tmpdir) / "out"))
        state.spec["report_profile"] = "business_report"
        state.plan["blueprint"] = {
            "sections": {"findings": {"section_type": "findings"}},
            "section_order": ["findings"],
        }
        return state

    def test_hedged_price_series_is_recommended_as_a_line(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            recommendations = recommend_figures_from_evidence(
                self._state(tmpdir), self.PRICE_TABLE
            )
            self.assertTrue(recommendations)
            self.assertEqual(recommendations[0]["recommended_figure_type"], "line")

    def test_range_cells_are_declared_not_passed_off_as_exact(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            recommendation = recommend_figures_from_evidence(
                self._state(tmpdir), self.PRICE_TABLE
            )[0]
            summary = recommendation["data_profile"]["summary"]
            self.assertGreater(summary["range_value_count"], 0)
            self.assertTrue(
                any("midpoint" in warning for warning in recommendation["selection_warnings"]),
                "a chart drawn from midpoints did not say so",
            )

    def test_every_shape_a_real_table_writes_is_read_as_a_number(self):
        """These six killed whole columns, and a dead column kills the chart."""
        from report_workflow.nodes.figure_utils import parse_measure

        cases = {
            "~8,259": 8259.0,
            "21,570–21,760": 21665.0,
            "約 500": 500.0,
            "1,000": 1000.0,
            "US$417/噸": 417.0,
            "±5%": 5.0,
        }
        for text, expected in cases.items():
            with self.subTest(cell=text):
                measure = parse_measure(text)
                self.assertIsNotNone(measure, f"{text!r} was not read as a number")
                self.assertAlmostEqual(measure.value, expected)

    def test_a_tolerance_is_kept_as_an_interval(self):
        from report_workflow.nodes.figure_utils import parse_measure

        measure = parse_measure("5 ± 0.3")
        self.assertAlmostEqual(measure.value, 5.0)
        self.assertAlmostEqual(measure.low, 4.7)
        self.assertAlmostEqual(measure.high, 5.3)
        self.assertAlmostEqual(measure.tolerance, 0.3)
        self.assertTrue(measure.is_uncertain)

    def test_a_magnitude_suffix_is_refused_not_silently_rescaled(self):
        """Reading "12.4bn" as 12.4 is a wrong number stated confidently."""
        from report_workflow.nodes.figure_utils import parse_measure, unparsed_reason

        self.assertIsNone(parse_measure("12.4bn"))
        self.assertIn("ambiguous", unparsed_reason("12.4bn"))
        self.assertIsNotNone(parse_measure("417 USD/t"), "a denominator is not a magnitude")

    def test_the_caption_says_the_plotted_values_are_not_exact(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            recommendation = recommend_figures_from_evidence(
                self._state(tmpdir), self.PRICE_TABLE
            )[0]
            title = recommendation["figure_plan"]["title"]
            self.assertTrue(
                "區間中點" in title or "range midpoints" in title,
                f"the caption does not disclose the midpoints: {title!r}",
            )

    def test_a_table_with_no_measure_column_still_recommends_a_table(self):
        """The reverse test: widening the parser must not turn prose numeric."""
        table = [{
            "evidence_id": "E_text",
            "source_id": "S001",
            "source_file_name": "recycling.md",
            "granularity": "table",
            "content": "分類與負責單位",
            "table_data": [
                ["品類", "主管機關", "備註"],
                ["電池", "環境部", "須申報"],
                ["塑膠", "經濟部", "自願性"],
                ["紡織", "環境部", "研議中"],
            ],
        }]
        with tempfile.TemporaryDirectory() as tmpdir:
            recommendation = recommend_figures_from_evidence(self._state(tmpdir), table)[0]
            self.assertEqual(recommendation["recommended_figure_type"], "table")

    def test_unreadable_cells_are_named_rather_than_the_column_abandoned(self):
        table = [{
            "evidence_id": "E_mixed",
            "source_id": "S001",
            "source_file_name": "recycling.md",
            "granularity": "table",
            "content": "市場規模",
            "table_data": [
                ["品類", "市場規模"],
                ["電池", "12.4bn"],
                ["塑膠", "8.1bn"],
                ["紡織", "3.2bn"],
            ],
        }]
        with tempfile.TemporaryDirectory() as tmpdir:
            recommendation = recommend_figures_from_evidence(self._state(tmpdir), table)[0]
            warnings = " ".join(recommendation["selection_warnings"])
            self.assertIn("12.4bn", warnings)
            self.assertIn("市場規模", warnings)


class OutlineFigureContractTests(unittest.TestCase):
    """A figure planned and never used only failed after the DOCX existed.

    "expected 4 Word table(s), found 1" arrived from POST_RENDER_VALIDATE,
    with the render already paid for and nothing naming the entries to remove.
    """

    def _outline_state(self, tmpdir: str, used_figures: list[str]) -> ReportState:
        state, _src = _prepare_market_report(tmpdir)
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        state.plan["claim_matrix"] = {
            "claims": [{"claim_id": "c1"}, {"claim_id": "c_limit_1"}, {"claim_id": "c_limit_2"}]
        }
        sections = {
            section_id: {
                "section_id": section_id,
                "claim_ids": (
                    ["c_limit_1", "c_limit_2"] if section_id == "limitations" else ["c1"]
                ),
                "figure_ids": used_figures if section_id == "findings" else [],
            }
            for section_id in state.plan["blueprint"]["section_order"]
        }
        if "limitations" in sections:
            sections["limitations"]["undermines"] = ["c1"]
        (run_dir / "outline.json").write_text(
            json.dumps({"sections": sections}), encoding="utf-8"
        )
        section_dir = run_dir / "section_drafts"
        section_dir.mkdir(exist_ok=True)
        (section_dir / "figure_plan.json").write_text(
            json.dumps({"figures": [
                {"figure_id": "fig_batteries"},
                {"figure_id": "fig_paper"},
            ]}),
            encoding="utf-8",
        )
        return state

    def test_unused_planned_figures_block_at_the_outline_stage(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._outline_state(tmpdir, ["fig_batteries"])
            with self.assertRaises(QAHardBlockError) as ctx:
                run_outline_plan(state)
            message = str(ctx.exception)
            self.assertIn("fig_paper", message)
            self.assertIn("figure_plan.json", message)

    def test_outline_naming_an_unplanned_figure_blocks_too(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state = self._outline_state(tmpdir, ["fig_batteries", "fig_paper", "fig_ghost"])
            with self.assertRaises(QAHardBlockError) as ctx:
                run_outline_plan(state)
            self.assertIn("fig_ghost", str(ctx.exception))


class OutlineSubsectionTests(unittest.TestCase):
    """Four independent categories had one flat findings section to live in."""

    def _write_outline(self, state: ReportState, sections: dict) -> None:
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        # Prepare writes a figure_plan.json of its own recommendations, and the
        # outline gate now requires the two to agree; this suite is about
        # subsections, so give findings whatever was planned.
        plan_path = run_dir / "section_drafts" / "figure_plan.json"
        if plan_path.exists():
            planned = json.loads(plan_path.read_text(encoding="utf-8")).get("figures", [])
            sections["findings"]["figure_ids"] = [
                str(figure.get("figure_id")) for figure in planned if figure.get("figure_id")
            ]
        (run_dir / "outline.json").write_text(
            json.dumps({"sections": sections}), encoding="utf-8"
        )

    def test_subsection_claims_and_figures_roll_up_to_the_section(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            state.plan["claim_matrix"] = {
                "claims": [
                    {"claim_id": "c1"}, {"claim_id": "c2"},
                    {"claim_id": "c_limit_1"}, {"claim_id": "c_limit_2"},
                ]
            }
            sections = {
                section_id: {"section_id": section_id, "claim_ids": ["c1"], "figure_ids": []}
                for section_id in state.plan["blueprint"]["section_order"]
            }
            sections["limitations"] = {
                "section_id": "limitations",
                "claim_ids": ["c_limit_1", "c_limit_2"],
                "undermines": ["c1"],
                "figure_ids": [],
            }
            sections["findings"] = {
                "section_id": "findings",
                "claim_ids": [],
                "subsections": [
                    {"subsection_id": "batteries", "title": "鋰電池", "claim_ids": ["c1"]},
                    {"subsection_id": "paper", "title": "紙張", "claim_ids": ["c2"]},
                ],
            }
            self._write_outline(state, sections)
            state = run_outline_plan(state)
            findings = state.plan["outline"]["sections"]["findings"]
            self.assertEqual(findings["claim_ids"], ["c1", "c2"])
            self.assertEqual(len(findings["subsections"]), 2)

    def test_a_subsection_without_a_title_is_refused(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            state.plan["claim_matrix"] = {"claims": [{"claim_id": "c1"}]}
            sections = {
                section_id: {"section_id": section_id, "claim_ids": ["c1"]}
                for section_id in state.plan["blueprint"]["section_order"]
            }
            sections["findings"]["subsections"] = [
                {"subsection_id": "batteries", "claim_ids": ["c1"]}
            ]
            self._write_outline(state, sections)
            with self.assertRaises(QAHardBlockError) as ctx:
                run_outline_plan(state)
            self.assertIn("title", str(ctx.exception))


class SourceTableTests(unittest.TestCase):
    """A source's tables had no route back into the deliverable.

    Ingestion split each table into one citable row per line and nothing put
    them back, so a source carrying four tables produced a document carrying
    none. The author's alternatives were retyping the numbers — unchecked by
    any gate, which defeats the ledger — or routing the table through the
    chart recommender, which answers a different question.
    """

    def test_a_split_table_is_rebuilt_from_its_rows(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            tables = collect_source_tables(_ledger(state))
            self.assertGreaterEqual(len(tables), 4, "the four source tables were not rebuilt")

            price = next(
                table for table in tables.values()
                if any("價格" in header for header in table["headers"])
            )
            self.assertEqual(len(price["rows"]), 4)
            self.assertEqual(price["rows"][0][0], "2022-11 峰值")
            self.assertIn("~80,000–85,000", price["rows"][0][1])
            self.assertTrue(price["source_file_name"].endswith("recycling.md"))
            self.assertTrue(price["evidence_ids"])

    def test_a_marker_expands_to_the_table_and_its_source(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            rows = _ledger(state)
            table_id = next(iter(collect_source_tables(rows)))
            markdown, placed, unresolved = replace_table_placeholders(
                f"見下表。\n\n[TABLE:{table_id} 回收成本]\n", rows
            )
            self.assertEqual(placed, 1)
            self.assertEqual(unresolved, [])
            self.assertIn("| ---", markdown)
            self.assertIn("回收成本", markdown)
            self.assertIn("來源：", markdown, "the table shipped without its provenance")

    def test_an_unknown_table_id_is_reported_not_swallowed(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            state, _src = _prepare_market_report(tmpdir)
            markdown, placed, unresolved = replace_table_placeholders(
                "[TABLE:md_no_such_table]", _ledger(state)
            )
            self.assertEqual(placed, 0)
            self.assertEqual(unresolved, ["md_no_such_table"])
            self.assertIn("[TABLE:md_no_such_table]", markdown)


class DeliveredDocumentTests(unittest.TestCase):
    """The unit test proves the resolver emits markers; this proves they ship.

    The failure being guarded was entirely in the wiring: the resolver was
    called, the markers were computed away to the empty string, and the
    delivered DOCX carried no trace of any source. Asserting on the rendered
    file is the only assertion that would have caught it.
    """

    def _author_and_publish(self, tmpdir: str):
        from docx import Document
        from report_workflow.run_workflow import render_workflow, validate_workflow

        state, _src = _prepare_market_report(tmpdir)
        run_dir = WORKFLOW_RUNS_DIR / state.job_id
        rows = _ledger(state)
        table_id = next(iter(collect_source_tables(rows)), "")

        # Prose blocks only. This synthetic author writes each claim as a
        # prefix of the evidence it cites, which is a sentence when the
        # evidence is a paragraph and a truncated JSON record when it is a
        # table row — and a heading asserts nothing at all. FE runs on the
        # publish path now and is right to refuse both; picking them was the
        # fixture speaking in a way no author would.
        prose = [
            row for row in rows
            if row.get("block_type") not in {"heading", "table_row", "csv_row", "data_row"}
        ]
        quantitative = [row for row in prose if "statistical" in row["allowed_claim_types"]][:2]
        qualitative = [row for row in prose if "statistical" not in row["allowed_claim_types"]][:1]
        picked = quantitative + qualitative
        claims = [
            {
                "claim_id": f"c{index}",
                "claim_text": row["content"][:120],
                "claim_type": (
                    "statistical" if "statistical" in row["allowed_claim_types"] else "factual"
                ),
                "risk_level": "low",
                "status": "supported",
                "evidence_ids": [row["evidence_id"]],
                "requires_hedged_wording": True,
                "claim_role": "primary",
            }
            for index, row in enumerate(picked, start=1)
        ]
        # The counter-evidence section this blueprint requires has to carry
        # claims of its own and name conclusions it does not carry, so it gets
        # its own copies rather than sharing the body's.
        body_claim_ids = [claim["claim_id"] for claim in claims]
        claims.extend(
            {**claims[-1], "claim_id": f"c_limit_{index}", "claim_role": "supporting"}
            for index in (1, 2)
        )
        counter_claim_ids = ["c_limit_1", "c_limit_2"]
        (run_dir / "claim_matrix.json").write_text(
            json.dumps({"claims": claims}, ensure_ascii=False), encoding="utf-8"
        )

        # Publish no figures: this suite is about citations, and the outline
        # gate now requires figure_plan.json and the outline to agree.
        plan_path = run_dir / "section_drafts" / "figure_plan.json"
        if plan_path.exists():
            payload = json.loads(plan_path.read_text(encoding="utf-8"))
            payload["figures"] = []
            plan_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

        order = state.plan["blueprint"]["section_order"]
        claimless = {"references", "appendix"}

        def _section_claim_ids(section_id: str) -> list[str]:
            if section_id in claimless:
                return []
            return counter_claim_ids if section_id == "limitations" else body_claim_ids

        sections = {
            section_id: {
                "section_id": section_id,
                "goals": f"cover {section_id}",
                "claim_ids": _section_claim_ids(section_id),
                "paragraph_order": ["state the supported claims"],
                "figure_ids": [],
            }
            for section_id in order
        }
        if "limitations" in sections:
            sections["limitations"]["undermines"] = body_claim_ids[:1]
        (run_dir / "outline.json").write_text(
            json.dumps({"sections": sections}, ensure_ascii=False), encoding="utf-8"
        )

        section_dir = run_dir / "section_drafts"
        section_dir.mkdir(exist_ok=True)
        sentence_rows = []
        for section_id in order:
            if section_id in claimless:
                (section_dir / f"{section_id}.md").write_text(
                    f"# {section_id}\n\n本節無正文。\n", encoding="utf-8"
                )
                continue
            lines = [f"# {section_id}", ""]
            if section_id == "findings" and table_id:
                lines.extend([f"下表列出來源的原始數據。[TABLE:{table_id} 回收成本]", ""])
            wanted = set(_section_claim_ids(section_id))
            for claim in (c for c in claims if c["claim_id"] in wanted):
                marker = " ".join(f"[CITE:{eid}]" for eid in claim["evidence_ids"])
                # The whole claim text, not a prefix of it: FS blocks a sentence that
                # omits a figure its claim asserts, and truncating here would
                # make this synthetic author commit exactly that.
                lines.extend([f"根據來源資料，{claim['claim_text']} {marker}。", ""])
                sentence_rows.append({
                    "sentence_id": f"s_{section_id}_{claim['claim_id']}",
                    "section_id": section_id,
                    "claim_ids": [claim["claim_id"]],
                    "evidence_ids": claim["evidence_ids"],
                    "citation_ids": claim["evidence_ids"],
                    "wording_strength": "hedged",
                    "draft_origin": "agent_draft",
                })
            (section_dir / f"{section_id}.md").write_text("\n".join(lines), encoding="utf-8")
        with open(run_dir / "sentence_map.jsonl", "w", encoding="utf-8") as handle:
            for row in sentence_rows:
                handle.write(json.dumps(row, ensure_ascii=False) + "\n")

        validated = validate_workflow(state.job_id)
        self.assertEqual((validated.qa or {}).get("qa_decision"), "pass")
        rendered = render_workflow(state.job_id)
        docx_path = (
            rendered.output.get("published_report_path")
            or rendered.output.get("final_docx_path")
            or rendered.output.get("rendered_docx_path")
        )
        self.assertTrue(docx_path and Path(docx_path).exists(), "nothing was rendered")
        document = Document(docx_path)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return rows, text, document.tables

    def test_the_delivered_document_carries_its_source_tables(self):
        """24 rows of source tables used to arrive as 0."""
        with tempfile.TemporaryDirectory() as tmpdir:
            _rows, text, tables = self._author_and_publish(tmpdir)
            self.assertGreaterEqual(len(tables), 1, "no Word table reached the document")
            cells = "\n".join(
                cell.text for table in tables for row in table.rows for cell in row.cells
            )
            self.assertIn("~80,000–85,000", cells, "the source's own numbers are missing")
            self.assertIn("來源：", text, "the table shipped without its provenance")

    def test_the_delivered_document_carries_its_sources(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            rows, text, _tables = self._author_and_publish(tmpdir)
            self.assertIn("[S1]", text, "no citation marker survived into the DOCX")
            # Either heading counts. The requirement is that the reader gets a
            # source list; which language its heading renders in is a separate
            # question, and conflating the two made a CI failure unreadable —
            # it could not say whether the section was missing or merely in
            # English, so the tail of the document goes into the message.
            self.assertTrue(
                "資料來源" in text or "Sources" in text,
                "the DOCX has no source list; document ends: " + repr(text[-400:]),
            )
            self.assertIn("recycling.md", text, "the source list names no file")
            # Traceability is checked the way a reader checks it: the entry
            # names a line span of a real file. Probing for the ledger's
            # internal id would only pass while that id was being printed into
            # the deliverable, which is itself the leak.
            traced = [
                row for row in rows
                if str(row.get("source_span") or "") and str(row["source_span"]) in text
            ]
            self.assertTrue(traced, "no source entry in the DOCX can be traced to a ledger row")
            leaked = [row for row in rows if row["evidence_id"] in text]
            self.assertFalse(leaked, "internal evidence ids reached the delivered document")

    def test_the_fallback_renderer_drops_the_sources_section(self):
        """A measured defect, not a passing behaviour.

        Without pandoc the python-docx fallback runs, and the generated
        Sources list does not survive it — References does. So a reader who
        installed the package without the render extra gets a document whose
        figures trace to nothing, which is the exact guarantee this section
        exists to provide.

        Recorded as failing-as-asserted rather than repaired here: the fix is
        in the fallback converter, and this round is a release. Delete this
        test when it is fixed; do not loosen the assertion above to match it.
        """
        import report_workflow.nodes.docx_render as docx_render

        with tempfile.TemporaryDirectory() as tmpdir:
            with patch.object(docx_render, "_find_pandoc", return_value=None):
                _rows, text, _tables = self._author_and_publish(tmpdir)
            self.assertIn("參考文獻", text, "references are dropped too now")
            self.assertNotIn(
                "資料來源",
                text,
                "the fallback renderer now keeps the Sources section — delete this test",
            )

    def test_no_workflow_marker_and_no_stranded_space_ship(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            _rows, text, _tables = self._author_and_publish(tmpdir)
            self.assertNotIn("[CITE:", text)
            self.assertNotIn("[TABLE:", text)
            self.assertNotIn(" 。", text)


class FeatureDecisionTests(unittest.TestCase):
    """An approved feature could not be turned on over MCP at all.

    start_report exposed no enable_* argument, so a recorded approval was
    refused with "the matching enable_* flag was not set" and the only way
    forward was to change the user's answer to `skip`.
    """

    def test_recorded_approval_is_the_flag(self):
        research, notebook = feature_flags_from_decisions(
            {"feature_decisions": {"web_research": "enable"}}, None, None
        )
        self.assertIs(research, True)
        self.assertIsNone(notebook)

    def test_recorded_decline_disables(self):
        research, _notebook = feature_flags_from_decisions(
            {"feature_decisions": {"web_research": "skip"}}, None, None
        )
        self.assertIs(research, False)

    def test_an_explicit_argument_still_wins(self):
        research, _notebook = feature_flags_from_decisions(
            {"feature_decisions": {"web_research": "enable"}}, False, None
        )
        self.assertIs(research, False)


if __name__ == "__main__":
    unittest.main()
