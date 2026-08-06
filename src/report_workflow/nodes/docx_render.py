"""DOCX_RENDER node - convert markdown to .docx.

Primary path: pandoc (robust, industry-standard Markdown?OCX converter).
Fallback path: python-docx regex-based converter (legacy).

The pandoc path uses a reference.docx template for consistent academic styling
(A4, Times New Roman 12pt, heading styles, 1.5x line spacing).
"""
import json
import logging
import re
import os
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..state import ReportState, WORKFLOW_RUNS_DIR
from ..errors import QAHardBlockError
from ..language import detect_document_language, localized_section_title
from ..parsers.office_math import cell_text, element_text
from .citation_bind import SOURCE_LIST_HEADING, SOURCE_LIST_HEADING_ZH
from .source_tables import replace_table_placeholders
from ..runtime_support import PLACEHOLDER_TEXT, load_jsonl
from ..policies import get_policy

logger = logging.getLogger(__name__)

# References-section heading names across supported document languages
# (blueprints ship "References", 參考文獻, or 參考資料). The regex and the
# literal tuple below are built from one list: they were separate copies, and a
# language added to one would have gone missing from the other.
_REFS_HEADING_WORDS = ("References", "Reference", "參考文獻", "參考資料")
_REFS_HEADING_NAMES = "(?:" + "|".join(_REFS_HEADING_WORDS) + ")"

# Path to the reference DOCX template (for pandoc --reference-doc)
_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "templates"
_REFERENCE_DOC = _TEMPLATE_DIR / "reference.docx"


# ------------------------------------------------------------------
# Pandoc-based converter (primary)
# ------------------------------------------------------------------

def _find_pandoc() -> str | None:
    """Find pandoc: on PATH, in a known install location, or bundled in a wheel."""
    pandoc_path = shutil.which("pandoc")
    if pandoc_path:
        return pandoc_path
    # Try common Windows install locations
    for candidate in [
        Path(r"C:\Program Files\Pandoc\pandoc.exe"),
        Path(r"C:\Users") / Path.home().name / "AppData" / "Local" / "Pandoc" / "pandoc.exe",
    ]:
        if candidate.exists():
            return str(candidate)
    return _bundled_pandoc()


def _bundled_pandoc() -> str | None:
    """The pandoc shipped inside `pypandoc-binary`, if the render extra is installed.

    Installing pandoc was the one manual step left in an otherwise one-command
    setup, and skipping it silently downgrades the document: the python-docx
    fallback renders without real Word tables or the template's layout. The
    wheel carries the binary, so `pip install "report-workflow[render]"`
    removes the step. A system pandoc still wins -- it is the one the user
    chose, and it is usually newer.
    """
    try:
        import pypandoc
    except ImportError:
        return None
    binary = "pandoc.exe" if os.name == "nt" else "pandoc"
    candidate = Path(pypandoc.__file__).resolve().parent / "files" / binary
    return str(candidate) if candidate.exists() else None


# ------------------------------------------------------------------
# Mermaid diagram conversion (optional)
# ------------------------------------------------------------------

_MERMAID_BLOCK_RE = re.compile(
    r'```mermaid\s*\n(.*?)\n\s*```',
    re.DOTALL,
)


def _find_mmdc() -> str | None:
    """Find the mermaid-cli (mmdc) executable on the system PATH."""
    mmdc_path = shutil.which("mmdc")
    if mmdc_path:
        return mmdc_path
    # Try common locations
    for candidate in [
        Path.home() / "AppData" / "Roaming" / "npm" / "mmdc.cmd",
        Path(r"C:\Program Files\nodejs\mmdc.cmd"),
        Path("/usr/local/bin/mmdc"),
        Path("/usr/bin/mmdc"),
    ]:
        if candidate.exists():
            return str(candidate)
    return None


def _convert_mermaid_blocks(md_content: str, output_dir: Path) -> tuple[str, int]:
    """Convert ```mermaid code fences to PNG images.

    For each mermaid block found:
    1. Write the mermaid source to a temp .mmd file
    2. Call mmdc to render it as PNG
    3. Replace the code fence with ![Figure N](path/to/png)

    Args:
        md_content: The markdown text to process.
        output_dir: Directory to write generated PNG files.

    Returns:
        (processed_markdown, count_of_converted_figures)
        If mmdc is not installed, returns the original text unchanged.
    """
    mmdc = _find_mmdc()
    if not mmdc:
        # Count mermaid blocks for logging
        count = len(_MERMAID_BLOCK_RE.findall(md_content))
        if count > 0:
            logger.warning(
                f"[DOCX_RENDER] Found {count} mermaid block(s) but mmdc not installed; "
                f"diagrams will not be rendered. Install: npm install -g @mermaid-js/mermaid-cli"
            )
        return md_content, 0

    figures_dir = output_dir / "figures"
    figures_dir.mkdir(parents=True, exist_ok=True)

    converted = 0
    figure_counter = 0

    def _replace_block(match: re.Match) -> str:
        nonlocal converted, figure_counter
        figure_counter += 1
        mermaid_source = match.group(1)

        mmd_path = figures_dir / f"mermaid_{figure_counter}.mmd"
        png_path = figures_dir / f"mermaid_figure_{figure_counter}.png"

        try:
            mmd_path.write_text(mermaid_source, encoding="utf-8")

            cmd = [
                mmdc,
                "-i", str(mmd_path),
                "-o", str(png_path),
                "-w", "800",
                "-b", "white",
                "--quiet",
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
                encoding="utf-8",
            )

            if result.returncode == 0 and png_path.exists():
                converted += 1
                logger.info(
                    f"[DOCX_RENDER] Mermaid block {figure_counter} -> {png_path.name}"
                )
                label = "圖" if detect_document_language(md_content) == "zh" else "Figure"
                return f"![{label} {figure_counter}]({png_path})"
            else:
                logger.warning(
                    f"[DOCX_RENDER] mmdc failed for block {figure_counter}: "
                    f"{result.stderr[:200] if result.stderr else 'unknown error'}"
                )
                return match.group(0)  # Keep original block

        except subprocess.TimeoutExpired:
            logger.warning(
                f"[DOCX_RENDER] mmdc timed out for block {figure_counter}"
            )
            return match.group(0)
        except Exception as exc:
            logger.warning(
                f"[DOCX_RENDER] mmdc error for block {figure_counter}: {exc}"
            )
            return match.group(0)

    processed = _MERMAID_BLOCK_RE.sub(_replace_block, md_content)
    if converted > 0:
        logger.info(
            f"[DOCX_RENDER] Converted {converted}/{figure_counter} mermaid block(s) to PNG"
        )
    return processed, converted


_IMAGE_LINK_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_FIGURE_PLACEHOLDER_RE = re.compile(r"\[FIGURE:\s*([^\]\s]+)(?:\s+([^\]]+))?\]", re.IGNORECASE)


def _load_figure_manifest(manifest_path: str) -> dict | None:
    """Load the generated figure manifest if present."""
    if not manifest_path or not Path(manifest_path).exists():
        return None
    try:
        with open(manifest_path, encoding="utf-8") as f:
            return json.load(f)
    except Exception as exc:
        logger.warning(f"[DOCX_RENDER] Could not load figure manifest: {exc}")
        return None


def _figure_alt_text(
    entry: dict,
    fallback_id: str,
    inline_caption: str = "",
    language: str = "en",
    display_number: str = "",
) -> str:
    label = "圖" if language == "zh" else "Figure"
    title = str(entry.get("title") or inline_caption or "").strip()
    figure_id = (
        display_number.strip()
        or str(entry.get("figure_id") or fallback_id).strip()
    )
    if title and figure_id and title.casefold() != figure_id.casefold():
        alt = f"{label} {figure_id}. {title}"
    elif figure_id:
        alt = f"{label} {figure_id}"
    else:
        alt = title or label
    return " ".join(alt.replace("[", "(").replace("]", ")").split())


def _replace_figure_placeholders(
    md_content: str,
    figure_manifest: dict | None,
    *,
    table_start_number: int = 0,
) -> tuple[str, int, list[str]]:
    """Replace [FIGURE:id] placeholders with markdown image links.

    FIGURE_BUILD already writes real image files and a manifest. The DOCX
    renderer's reliable image path is markdown image syntax, so normalize the
    generated placeholders into that existing path before pandoc sees the draft.
    """
    if not figure_manifest:
        return md_content, 0, []

    language = detect_document_language(md_content)
    entries_by_id: dict[str, dict] = {}
    for entry in figure_manifest.get("figures", []) or []:
        if not isinstance(entry, dict):
            continue
        figure_id = str(entry.get("figure_id") or "").strip()
        if figure_id:
            entries_by_id[figure_id.casefold()] = entry

    replaced = 0
    unresolved: list[str] = []
    # Figures and tables are numbered independently, as every style guide
    # expects. The caption used to print the author's figure_id, and ids are
    # unique across the whole plan, so a report with one chart and one table
    # rendered "圖 1." followed by "表 2." — 表 1 was unreachable. figure_id
    # stays the stable identity the manifest and gates match on.
    figure_number = 0
    # Source tables were already placed and numbered; continuing from there
    # keeps one sequence across the document, which is what a reader following
    # a cross-reference expects.
    table_number = table_start_number

    def replace(match: re.Match) -> str:
        nonlocal replaced, figure_number, table_number
        figure_id = match.group(1).strip()
        inline_caption = (match.group(2) or "").strip()
        entry = entries_by_id.get(figure_id.casefold())
        if not entry:
            unresolved.append(figure_id)
            return match.group(0)

        if str(entry.get("render_mode") or "") == "native_table":
            table_data = entry.get("data") or {}
            columns = [str(c) for c in (table_data.get("columns") or [])]
            rows = table_data.get("rows") or []
            if columns and rows:
                replaced += 1
                table_number += 1
                label = "表" if language == "zh" else "Table"
                title = str(entry.get("title") or inline_caption or "").strip()
                caption = f"{label} {table_number}. {title}".strip()

                def esc(value: object) -> str:
                    return str(value).replace("|", "\\|")

                header = "| " + " | ".join(esc(c) for c in columns) + " |"
                separator = "|" + "|".join(" --- " for _ in columns) + "|"
                body = "\n".join(
                    "| " + " | ".join(esc(v) for v in row) + " |" for row in rows
                )
                return f"{caption}\n\n{header}\n{separator}\n{body}"
            unresolved.append(figure_id)
            return match.group(0)

        image_path = Path(str(entry.get("path") or "").strip())
        if not image_path.exists():
            unresolved.append(figure_id)
            return match.group(0)

        replaced += 1
        figure_number += 1
        alt = _figure_alt_text(
            entry,
            figure_id,
            inline_caption,
            language=language,
            display_number=str(figure_number),
        )
        return f"![{alt}]({image_path.resolve().as_posix()})"

    return _FIGURE_PLACEHOLDER_RE.sub(replace, md_content), replaced, unresolved


def _absolutize_image_paths(
    md_content: str, base_dir: Path, draft_dir: Path | None = None
) -> str:
    """Rewrite local relative image links to absolute paths for pandoc.

    Pandoc's DOCX writer resolves image links relative to the process working
    directory in some Windows invocations, not always relative to the input
    markdown file. Absolute paths make figure embedding deterministic.

    A relative link means relative to the file that contains it, so the
    draft's own directory is tried first. Figures the pipeline builds carry
    absolute paths and are unaffected; a link written by hand in a section
    draft used to be resolved against the run directory alone, where it did
    not exist, and the figure then vanished from the document without a word.
    """

    def replace(match: re.Match) -> str:
        alt = match.group(1)
        raw_target = match.group(2).strip()
        if re.match(r"^[a-z]+://", raw_target, re.IGNORECASE):
            return match.group(0)
        target_no_title = raw_target.split(None, 1)[0].strip("<>")
        target_path = Path(target_no_title)
        if target_path.is_absolute():
            return match.group(0)
        candidates = [base_dir / target_path]
        if draft_dir is not None:
            candidates.insert(0, draft_dir / target_path)
        resolved = next(
            (c for c in candidates if c.exists()), candidates[-1]
        )
        return f"![{alt}]({resolved.resolve().as_posix()})"

    return _IMAGE_LINK_RE.sub(replace, md_content)


_TOC_TITLES = {"zh": "目錄", "en": "Table of Contents"}
_TOC_PLACEHOLDERS = {
    "zh": "（開啟後按 F9 或於列印時自動更新目錄）",
    "en": "(Press F9 after opening, or print, to populate the table of contents.)",
}


def _toc_openxml_block(language: str, page_break_before: bool) -> str:
    """Raw OOXML table-of-contents block for the pandoc input.

    pandoc's own --toc places the TOC ahead of everything, which puts it in
    front of the title page. Injecting the field manually keeps the title
    page first and localizes the heading. TOCHeading exists in the reference
    template, so the heading styles correctly and stays out of the TOC
    field's own listing.
    """
    title = _TOC_TITLES.get(language, _TOC_TITLES["en"])
    placeholder = _TOC_PLACEHOLDERS.get(language, _TOC_PLACEHOLDERS["en"])
    parts = []
    if page_break_before:
        parts.append('<w:p><w:r><w:br w:type="page"/></w:r></w:p>')
    parts.append(
        '<w:p><w:pPr><w:pStyle w:val="TOCHeading"/></w:pPr>'
        f"<w:r><w:t>{title}</w:t></w:r></w:p>"
    )
    parts.append(
        '<w:p><w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        f"<w:r><w:t>{placeholder}</w:t></w:r>"
        '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
    )
    parts.append('<w:p><w:r><w:br w:type="page"/></w:r></w:p>')
    return "```{=openxml}\n" + "\n".join(parts) + "\n```"


# A Chinese line that ends mid-paragraph: last character is a Han character
# or Chinese punctuation.
_CJK_LINE_END_RE = re.compile(r"[一-鿿㐀-䶿。，、；:：？！」』）】》%]$")
_CJK_BLOCK_PREFIXES = ("|", "#", ">", "-", "*", "+", "```", "[FIGURE", "[Source")


_LINK_TARGET_RE = re.compile(r"\]\([^)]*\)")
_CJK_GAP_RE = re.compile(
    r"([一-鿿㐀-䶿。，、；:：？！」』）】》])[ \t]+([一-鿿㐀-䶿「『（【《\[])"
)


def _close_cjk_gaps(text: str) -> str:
    return _CJK_GAP_RE.sub(r"\1\2", text)


def _normalize_cjk_typography(md: str) -> str:
    """Chinese sentences do not take a space between them.

    Each authored sentence is its own markdown line, and pandoc turns an
    intra-paragraph newline into a space — correct for English, wrong for
    Chinese, where it renders as "…增加。 5 N 時…". Join those lines
    directly and close the gap before a citation marker.
    """
    out: list[str] = []
    in_fence = False
    for line in md.splitlines():
        stripped = line.strip()
        if stripped.startswith("```"):
            in_fence = not in_fence
            out.append(line)
            continue
        joinable = (
            not in_fence
            and out
            and stripped
            and not stripped.startswith(_CJK_BLOCK_PREFIXES)
            and not out[-1].strip().startswith(_CJK_BLOCK_PREFIXES)
            and bool(_CJK_LINE_END_RE.search(out[-1].strip()))
        )
        if joinable:
            out[-1] = out[-1].rstrip() + stripped
        else:
            out.append(line)
    joined = "\n".join(out)
    # Close gaps left inside a line too: a stripped internal-source marker
    # leaves "轉動。 千分錶", and an authored marker leaves "4.8%。 [1]".
    # Only CJK-to-CJK (or CJK-to-citation) gaps close; a space between
    # Chinese and Latin ("撓度 1.52 mm") is real spacing and stays.
    #
    # Link and image targets are exempt. A run directory is named after the
    # user's prompt, so its path routinely contains a space between two CJK
    # words — closing that gap rewrote the figure's filename, pandoc found
    # nothing there, and the document rendered silently without the image.
    pieces: list[str] = []
    cursor = 0
    for match in _LINK_TARGET_RE.finditer(joined):
        pieces.append(_close_cjk_gaps(joined[cursor:match.start()]))
        pieces.append(match.group(0))
        cursor = match.end()
    pieces.append(_close_cjk_gaps(joined[cursor:]))
    return "".join(pieces)


def _xml_text_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _cover_openxml_block(cover_body: str) -> str:
    """Centered title-page paragraphs for a leading cover section.

    Raw openxml deliberately: the cover prose is plain sentences by the time
    it reaches the renderer (citations already bound), and markdown has no
    centering.
    """
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", cover_body) if p.strip()]
    parts = [
        '<w:p><w:pPr><w:jc w:val="center"/><w:spacing w:before="240" w:after="240"/></w:pPr>'
        f'<w:r><w:t xml:space="preserve">{_xml_text_escape(paragraph)}</w:t></w:r></w:p>'
        for paragraph in paragraphs
    ]
    return "```{=openxml}\n" + "\n".join(parts) + "\n```"


def _inject_toc(
    md_content: str,
    has_front_matter: bool,
    cover_title: str = "",
    title_leads: bool = False,
) -> str:
    """Insert the TOC block after the front matter or a leading cover section.

    Placement, in priority order: after the front-matter separator; after a
    leading cover section; else at the top of the document. A leading cover
    section is also promoted to a title page: its heading is dropped (a
    cover page does not label itself, and without a Heading 1 it stays out
    of the TOC field) and its paragraphs render centered. Only the pandoc
    input gets any of this: the python-docx fallback would render the
    raw-openxml fences as literal code blocks.
    """
    language = detect_document_language(md_content)
    if has_front_matter:
        separator = "\n\n---\n\n"
        idx = md_content.find(separator)
        if idx != -1:
            block = _toc_openxml_block(language, page_break_before=True)
            head = md_content[:idx]
            body = md_content[idx + len(separator):]
            return head + "\n\n" + block + "\n\n" + body
    if cover_title:
        headings = list(re.finditer(r"^# .*$", md_content, flags=re.MULTILINE))
        if headings and cover_title in headings[0].group(0):
            end = headings[1].start() if len(headings) >= 2 else len(md_content)
            cover_body = md_content[headings[0].end():end].strip()
            toc_block = _toc_openxml_block(language, page_break_before=True)
            head = md_content[:headings[0].start()]
            if cover_body:
                cover_block = _cover_openxml_block(cover_body)
                return head + cover_block + "\n\n" + toc_block + "\n\n" + md_content[end:]
            return head + toc_block + "\n\n" + md_content[end:]
    if title_leads:
        # Revised documents open with the base document's own title H1;
        # the TOC belongs after it, not on top of it.
        headings = list(re.finditer(r"^# .*$", md_content, flags=re.MULTILINE))
        if len(headings) >= 2:
            block = _toc_openxml_block(language, page_break_before=True)
            idx = headings[1].start()
            return md_content[:idx] + block + "\n\n" + md_content[idx:]
    block = _toc_openxml_block(language, page_break_before=False)
    return block + "\n\n" + md_content


def reference_docx_error(path: Path) -> str | None:
    """Return why a user-supplied reference .docx is unusable, or None if OK."""
    if not path.exists():
        return f"file not found: {path}"
    if path.suffix.lower() != ".docx":
        return f"not a .docx file: {path}"
    try:
        with zipfile.ZipFile(path) as z:
            if "word/styles.xml" not in z.namelist():
                return f"no word/styles.xml inside (not a Word document?): {path}"
    except zipfile.BadZipFile:
        return f"not a valid docx (zip) file: {path}"
    return None


def _reference_doc_body_carryover(path: Path) -> str:
    """What a supplied template contains that the render will not carry over.

    A course hands out a .docx with the cover page it wants — the department,
    the name and student-number blanks, the supervisor line — and pandoc's
    --reference-doc takes styles only. The template was accepted without a
    word, the fonts came through, and the cover the course actually grades was
    gone, replaced by this pipeline's own.

    Losing it may be the right trade: the styles are the part a template can
    give. Saying nothing about it is not.
    """
    try:
        lines = [
            " ".join(paragraph.text.split())
            for paragraph in Document(str(path)).paragraphs
            if paragraph.text.strip()
        ]
    except Exception:
        return ""
    return "; ".join(lines[:3])


def _resolve_reference_doc(spec: dict) -> Path:
    """Pick the styling template: user-supplied reference docx or the built-in.

    A user-supplied template that is missing or unreadable hard-blocks the
    render instead of silently falling back — the user asked for their
    formatting, so shipping the default would be the wrong document.
    """
    custom = str(spec.get("reference_docx_path") or "").strip()
    if not custom:
        return _REFERENCE_DOC
    path = Path(custom)
    error = reference_docx_error(path)
    if error:
        raise QAHardBlockError(f"Custom reference docx unusable: {error}")
    return path


def _pandoc_warnings(stderr: str) -> list[str]:
    """The first line of each pandoc warning, which is the actionable one.

    pandoc names what it could not do — "Could not fetch resource chart.png:
    replacing image with description", "Could not convert TeX math …, rendering
    as TeX" — and those are statements about the deliverable: a figure replaced
    by its alt text, a formula printed as raw TeX in a submitted report. They
    went to logger.info truncated at 300 characters, which is shorter than the
    TeX warning itself, so even the log was cut mid-sentence.

    Continuation lines are the parser's list of what it expected instead; that
    is for whoever fixes the markup, not for the report's QA.
    """
    warnings_found: list[str] = []
    for line in (stderr or "").splitlines():
        stripped = line.strip()
        if stripped.startswith("[WARNING]") and stripped not in warnings_found:
            warnings_found.append(stripped)
    return warnings_found


def _render_via_pandoc(
    md_path: str,
    output_path: str,
    reference_doc: Path | None = None,
    warnings_out: list[str] | None = None,
) -> bool:
    """Convert markdown to DOCX using pandoc.

    The table of contents is not pandoc's --toc: _inject_toc places a TOC
    field after the front matter in the markdown instead, so the title page
    stays first and the heading is localized.

    Returns True on success, False if pandoc is unavailable or fails.
    """
    pandoc = _find_pandoc()
    if not pandoc:
        logger.warning("[DOCX_RENDER] pandoc not found on PATH; falling back to python-docx")
        return False

    cmd = [pandoc, str(md_path), "-o", str(output_path)]

    # Use reference doc for styling if available
    ref_doc = reference_doc if reference_doc is not None else _REFERENCE_DOC
    if ref_doc.exists():
        cmd.extend(["--reference-doc", str(ref_doc)])

    # Standalone output
    cmd.append("--standalone")

    # Wrap long lines
    cmd.extend(["--wrap", "auto"])

    # Set metadata (prevents pandoc from using filename as title)
    cmd.extend(["-M", "title="])

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
            encoding="utf-8",
        )
        if result.returncode != 0:
            logger.error(
                f"[DOCX_RENDER] pandoc failed (exit {result.returncode}): {result.stderr[:500]}"
            )
            return False
        found = _pandoc_warnings(result.stderr)
        if found:
            logger.warning("[DOCX_RENDER] pandoc: " + " | ".join(found))
            if warnings_out is not None:
                warnings_out.extend(found)
        return True
    except FileNotFoundError:
        logger.warning("[DOCX_RENDER] pandoc executable not found")
        return False
    except subprocess.TimeoutExpired:
        logger.error("[DOCX_RENDER] pandoc timed out after 120s")
        return False
    except Exception as exc:
        logger.error(f"[DOCX_RENDER] pandoc error: {exc}")
        return False


# ------------------------------------------------------------------
# Post-render validation
# ------------------------------------------------------------------

_MARKUP_STRIP_RE = re.compile(
    r"```.*?```|!\[[^\]]*\]\([^)]*\)|\[[^\]]*\]\([^)]*\)|[#*_>`|-]", re.DOTALL
)


def _prose_length(text: str) -> int:
    """Characters of actual prose, markup and whitespace removed."""
    return len("".join(_MARKUP_STRIP_RE.sub(" ", text or "").split()))


def _docx_text_length(doc: Document) -> int:
    """Every character the document shows, tables included.

    ``doc.paragraphs`` does not reach inside tables, so a report whose
    results are a table counted as almost empty.

    ``paragraph.text`` is the runs only and an equation is not a run, so the
    theory section this pipeline itself renders from TeX counted as nothing on
    the way out while its source counted in full. Not seen failing — it takes a
    report that is more than half equations to cross the threshold — but the
    measurement was of something other than what the docstring claims.
    """
    total = sum(len(element_text(p._p)) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(cell_text(cell._tc))
    return total


def _validate_docx(docx_path: str, source_markdown: str | None = None) -> list[str]:
    """Validate a rendered DOCX file. Returns list of issues (empty = OK)."""
    issues = []
    path = Path(docx_path)

    if not path.exists():
        issues.append("DOCX file does not exist")
        return issues

    if path.stat().st_size < 1024:
        issues.append(f"DOCX file is suspiciously small ({path.stat().st_size} bytes)")

    try:
        doc = Document(str(path))
        # Check for content
        if len(doc.paragraphs) < 3:
            issues.append(f"DOCX has only {len(doc.paragraphs)} paragraphs; likely incomplete")

        # Check heading structure
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        if len(headings) == 0:
            issues.append("DOCX has no headings; structure may be broken")

        # Did the render lose the content? Measured against the markdown it
        # came from, not against a fixed character count. A 500-character
        # floor asks "is this long enough in English": Chinese says the same
        # thing in far fewer characters, so a complete report in Chinese was
        # told it was "likely incomplete" while its English translation
        # passed. The same threshold-tuned-in-English mistake as the block
        # floor that discarded short CJK headings.
        if source_markdown is not None:
            # A figure the draft asked for must not vanish quietly. Pandoc
            # emits the caption from the alt text whether or not it found the
            # file, so the document goes out with "圖 1." and a body that
            # says "由圖 1 可見" and no figure between them. The repair pass
            # already learns which files are missing and then discards the
            # finding.
            wanted = list(_IMAGE_LINK_RE.finditer(source_markdown))
            missing = [
                m.group(2).strip().strip("<>") for m in wanted
                if not Path(m.group(2).strip().strip("<>")).exists()
            ]
            if missing:
                issues.append(
                    f"{len(missing)} figure file(s) referenced by the draft do "
                    f"not exist: {', '.join(missing[:3])}"
                )
            elif wanted and len(doc.inline_shapes) < len(wanted):
                issues.append(
                    f"draft references {len(wanted)} figure(s) but the document "
                    f"embeds {len(doc.inline_shapes)}"
                )

            rendered = _docx_text_length(doc)
            source = _prose_length(source_markdown)
            if source and rendered < source * 0.5:
                issues.append(
                    f"DOCX carries {rendered} chars of the source's {source}; "
                    "content was lost in rendering"
                )

    except Exception as exc:
        issues.append(f"Cannot open DOCX for validation: {exc}")

    return issues


# ------------------------------------------------------------------
# Legacy python-docx fallback converter
# ------------------------------------------------------------------

# Box-drawing and block characters (U+2500-U+257F and related)
_ASCII_ART_CHARS_RE = re.compile(
    r"[\u2500-\u257f\u2501\u2574-\u257f\u2503\u250f\u2513\u251b\u2517\u2533\u253b\u2523\u252b"
    r"\u251b\u252b\u253b\u2523\u2502\u250c\u2510\u2518\u2514\u251c\u2525\u2528\u2534\u2538"
    r"\u2520\u252f\u253c\u2530\u253f\u2521\u2529\u2531\u2539\u2542\u254b]"
)
_GARBLED_BOX_ART_RE = re.compile(r"[\x80-\x9f\ue000-\uf8ff]")
_ASCII_ART_LANG_TAGS = {"ascii", "art", "figure", "diagram", "box", "graph", "text", ""}


def _is_ascii_art(fence_body: str) -> bool:
    """Return True if fence body looks like ASCII art."""
    if not fence_body:
        return False
    if _ASCII_ART_CHARS_RE.search(fence_body) or _GARBLED_BOX_ART_RE.search(fence_body):
        return True
    lines = fence_body.splitlines()
    if len(lines) >= 3:
        art_line_count = sum(
            1 for ln in lines
            if re.match(r"^[\s\-+=|<>#:\.\*\/\u2500-\u257f]+$", ln)
        )
        if art_line_count >= len(lines) * 0.5:
            return True
    return False


def _set_paragraph_title_style(doc: Document, title_text: str) -> None:
    """Add title text with the built-in 'Title' paragraph style."""
    para = doc.add_paragraph()
    para.style = doc.styles["Title"]
    run = para.add_run(title_text)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if run.font.size is None or run.font.size.pt < 16:
        run.font.size = Pt(18)


def _add_ascii_art_figure(doc: Document, fence_body: str, lang_tag: str) -> None:
    """Render ASCII art fence as a centered monospace paragraph."""
    caption_para = doc.add_paragraph()
    caption_run = caption_para.add_run("(Figure: see source appendix for diagram)")
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph()
    run = para.add_run(fence_body)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def _add_inline_formatted(para, text: str) -> None:
    """Add text to a paragraph, applying bold/italic/code inline patterns."""
    if not text:
        return
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def markdown_to_docx(
    md_text: str,
    output_path: str,
    figure_manifest: dict | None = None,
) -> None:
    """Convert markdown to DOCX using python-docx (legacy fallback).

    This is only used when pandoc is unavailable. The pandoc path
    handles tables, lists, code blocks, and inline formatting far
    more robustly than this regex-based approach.
    """
    doc = Document()

    figure_lookup: dict[str, str] = {}
    if figure_manifest:
        for entry in figure_manifest.get("figures", []):
            fid = entry.get("figure_id", "")
            fpath = entry.get("path", "")
            if fid and fpath:
                figure_lookup[fid] = fpath

    lines = md_text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # Code fences
        fence_match = re.match(r"^(```+|~~~+)\s*(\S*)?", line)
        if fence_match:
            fence_open = fence_match.group(1)
            lang_tag = (fence_match.group(2) or "").lower()
            fence_body_lines: list[str] = []
            i += 1
            while i < len(lines):
                closing = lines[i].strip()
                if re.match(rf"^{re.escape(fence_open)}$", closing):
                    i += 1
                    break
                fence_body_lines.append(lines[i])
                i += 1
            body_text = "\n".join(fence_body_lines).strip()
            if not body_text:
                continue

            is_art = _is_ascii_art(body_text) or lang_tag in _ASCII_ART_LANG_TAGS
            if is_art:
                _add_ascii_art_figure(doc, body_text, lang_tag)
            else:
                if lang_tag:
                    prefix = doc.add_paragraph()
                    prun = prefix.add_run(f"[code: {lang_tag}]")
                    prun.font.name = "Courier New"
                    prun.font.size = Pt(8)
                    prun.italic = True
                para = doc.add_paragraph()
                run = para.add_run(body_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            continue

        # Headings
        title_match = re.match(r"^#\s+\{\.Title\}\s+(.+)$", line)
        if title_match:
            _set_paragraph_title_style(doc, title_match.group(1))
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        # Bullet lists
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        # Numbered lists
        elif re.match(r"^\d+\.\s", line):
            match = re.match(r"^(\d+)\.\s(.*)", line)
            if match:
                doc.add_paragraph(match.group(2), style="List Number")
        # Tables
        elif " | " in line and i + 1 < len(lines) and "|" in lines[i + 1].strip():
            table_lines = []
            while i < len(lines) and (lines[i].strip().startswith("|") or " | " in lines[i]):
                table_lines.append(lines[i].strip())
                i += 1
            _SEP_RE = re.compile(r"^\|[\s\-\|\:\s]+\|$")
            data_lines = [ln for ln in table_lines if not _SEP_RE.match(ln)]
            if data_lines:
                num_cols = max(len(ln.split("|")) - 2 for ln in data_lines)
                num_cols = max(num_cols, 1)
                table = doc.add_table(rows=len(data_lines), cols=num_cols)
                for row_idx, tline in enumerate(data_lines):
                    cells = [c.strip() for c in tline.split("|")[1:-1]]
                    for col_idx, cell in enumerate(cells):
                        if col_idx < num_cols:
                            table.rows[row_idx].cells[col_idx].text = cell
            continue
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            _add_inline_formatted(para, line)

        i += 1

    doc.save(output_path)


def _add_hanging_indent_references(doc: Document, ref_md: str) -> None:
    """Append a formatted References section with hanging-indent paragraphs."""
    # Any level: the generated list and an authored one may differ, and the
    # heading must render at whatever level the markdown declares rather than
    # being forced to H2 — a demoted References nests under the section above
    # it in a Word table of contents.
    ref_re = re.compile(rf"^(#{{1,6}})\s+({_REFS_HEADING_NAMES})\s*$", re.MULTILINE)
    m = ref_re.search(ref_md)
    if not m:
        return

    refs_body = ref_md[m.end():].strip()
    entries = []
    current: list[str] = []
    for line in refs_body.split("\n"):
        stripped = line.strip()
        if stripped.startswith(("- ", "* ")):
            if current:
                entries.append("\n".join(current))
            current = [stripped[2:].strip()]
        elif re.match(r"^\[\d+\]\s+", stripped):
            if current:
                entries.append("\n".join(current))
            current = [stripped]
        elif stripped:
            current.append(stripped)
    if current:
        entries.append("\n".join(current))

    if not any(entry.strip() for entry in entries):
        # A heading with nothing under it reads as a broken document; a report
        # with no references simply has no References section.
        return

    doc.add_heading(m.group(2), level=len(m.group(1)))

    hanging_left = Inches(0.5)
    hanging_first = Inches(-0.5)

    for entry_text in entries:
        entry_text = entry_text.strip()
        if not entry_text:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = hanging_left
        para.paragraph_format.first_line_indent = hanging_first
        parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", entry_text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)


def _split_body_references(md_content: str, strict_refs: bool = True) -> tuple[str, str]:
    """Remove the inline References section from body markdown.

    Returns ``(md_without_references, body_refs_md)`` where ``body_refs_md``
    is a normalized bullet-list References block, or ``""`` when the section
    is missing or empty. An *empty* References section is removed outright —
    a document must never end with a dangling "References" heading that has
    nothing under it.
    """
    body_refs_md = ""
    # Match the heading at any level: upstream drafts carry "# References"
    # (H1) while normalized drafts carry "## References" (H2); both must be
    # captured or an empty section rides through to the rendered document.
    body_refs_match = re.search(
        rf"(?P<heading>^(?P<hashes>#{{1,6}})\s+{_REFS_HEADING_NAMES}"
        rf"[^\S\n]*(?:\n+|\Z))(?P<entries>.*?)(?=^#{{1,6}} |\Z)",
        md_content,
        re.MULTILINE | re.DOTALL,
    )
    if not body_refs_match:
        return md_content, body_refs_md

    entries_block = body_refs_match.group("entries")
    bullet_lines: list[str] = []
    for line in entries_block.splitlines():
        stripped = line.strip()
        if not stripped or stripped in _REFS_HEADING_WORDS:
            continue
        if stripped.startswith(("- ", "* ")):
            stripped = stripped[2:].strip()
        bullet_lines.append(f"- {stripped}")
    bullet_lines = _filter_body_reference_lines(bullet_lines, strict=strict_refs)
    if bullet_lines:
        # Keep the level the heading already had. Emitting a hardcoded H2
        # pushed 參考文獻 one level below its sibling sections, so Word nested
        # it under the last body section in the table of contents instead of
        # listing it alongside them.
        hashes = body_refs_match.group("hashes")
        body_refs_md = f"{hashes} References\n\n" + "\n\n".join(bullet_lines) + "\n"
    elif entries_block.strip():
        # The section carried authored content, but none of it is a citation —
        # commentary about the sources, say. Dropping it is right (References
        # lists references), but dropping it silently meant the author only
        # discovered the loss by reading the rendered document.
        logger.warning(
            "[DOCX_RENDER] References section removed: %d line(s) of authored "
            "content held no usable citation. References carries citations "
            "only; move source commentary into a body section.",
            len([ln for ln in entries_block.strip().splitlines() if ln.strip()]),
        )
    md_content = md_content[:body_refs_match.start()] + md_content[body_refs_match.end():]
    return md_content, body_refs_md


def _localize_reference_heading(ref_md: str, body_md: str, blueprint: dict) -> str:
    """Replace an English References heading with the blueprint's localized title.

    The citation chain writes ``## References`` as an internal marker; for a
    Chinese-language document the final rendered heading must follow the
    blueprint's ``title_zh`` (e.g. 參考文獻) like every other section.
    """
    if not ref_md.strip() or detect_document_language(body_md) != "zh":
        return ref_md
    sections = (blueprint or {}).get("sections", {}) or {}
    title = localized_section_title(sections.get("references"), "references", "zh")
    return re.sub(
        r"^(#{1,6})\s+References?\s*$",
        lambda m: f"{m.group(1)} {title}",
        ref_md,
        count=1,
        flags=re.MULTILINE,
    )


def _truncate_orphan_sections(md_text: str) -> str:
    """Strip orphan content after the canonical end of the main body."""
    sentinel_re = re.compile(r"^##?\s+End\s+of\s+Main\s+Report\s*$", re.MULTILINE | re.IGNORECASE)
    m = sentinel_re.search(md_text)
    if m:
        md_text = md_text[:m.start()]
    return md_text


# ------------------------------------------------------------------
# Pre-render sanity gate
# ------------------------------------------------------------------

_BRACKET_PLACEHOLDER_RE = re.compile(
    r"\[(?:Author Name|University|email@domain\.com|Your Name|Your University|INSERT .+?)\]",
    re.IGNORECASE,
)


def _pre_render_sanity_check(
    md_content: str,
    facts_freeze: dict | None = None,
    forbidden_fragments: list[str] | None = None,
) -> list[str]:
    """Run hard sanity checks on final markdown BEFORE rendering.

    Returns a list of blocking issues. Empty list = all clear.
    These checks catch problems that individual pipeline nodes miss
    because they operate on fragments, not the final assembled document.
    """
    issues: list[str] = []

    # 1. Duplicated headings
    headings = re.findall(r'^#+\s+(.+)$', md_content, re.MULTILINE)
    seen: dict[str, int] = {}
    for h in headings:
        h_clean = h.strip()
        seen[h_clean] = seen.get(h_clean, 0) + 1
    for h_text, count in seen.items():
        if count > 1:
            issues.append(f"Duplicated heading ({count}x): \"{h_text}\"")

    # 2. Duplicated References section
    ref_count = len(re.findall(r'^#{1,6}\s+References?\s*$', md_content, re.MULTILINE))
    if ref_count > 1:
        issues.append(f"Multiple References sections found ({ref_count})")

    # 3. Placeholder metadata
    placeholders = _BRACKET_PLACEHOLDER_RE.findall(md_content)
    if placeholders:
        issues.append(f"Placeholder metadata still present: {placeholders[:3]}")

    # 3b. Prompt/template metadata leakage
    leakage_patterns = [
        (r"\{\.Title\}", "Pandoc title marker leaked into publication text"),
        (r"\bRevise\s+the\s+base\s+document\b", "raw task instruction leaked into publication text"),
        (r"\bWrite\s+an?\s+academic\s+report\b", "raw task instruction leaked into publication text"),
        (r"\bIndependent\s+Researcher\b", "generic template affiliation leaked"),
        (r"\bauthor@example\.com\b", "generic template correspondence leaked"),
        (r"\bResearch\s+Author\b", "generic research author leaked"),
        (r"\bResearch\s+University\b", "generic research university leaked"),
        (r"\bresearch@university\.edu\b", "generic research correspondence leaked"),
        (r"^Author\s*$", "generic template author leaked"),
        (r"\(source\s*&\s*corpus\s*\(n\.d\.\)\)", "internal pseudo-citation leaked into publication text"),
    ]
    for pattern, label in leakage_patterns:
        if re.search(pattern, md_content, re.IGNORECASE | re.MULTILINE):
            issues.append(label)

    front_matter_head = "\n".join(md_content.splitlines()[:20])
    noisy_keyword_terms = ("Corpus", "Backtrader", "Pydantic", "Kelly", "Bayesian", "Ollama")
    keyword_line = re.search(r"Keywords:\s*(.+)", front_matter_head, re.IGNORECASE)
    if keyword_line:
        leaked_terms = [term for term in noisy_keyword_terms if re.search(rf"\b{re.escape(term)}\b", keyword_line.group(1), re.IGNORECASE)]
        if leaked_terms:
            issues.append("Front matter keywords contain implementation-noise terms: " + ", ".join(leaked_terms))

    bold_hit = re.search(r"^\s*\*\*\s+\S.*$", front_matter_head, re.MULTILINE)
    if bold_hit:
        issues.append(
            "Front matter contains leftover Markdown bold marker: "
            f'"{" ".join(bold_hit.group(0).split())[:120]}"'
        )

    corpus_hit = re.search(r"^.*source_corpus.*$", md_content, re.IGNORECASE | re.MULTILINE)
    if corpus_hit:
        line = " ".join(corpus_hit.group(0).split())[:120]
        issues.append(
            f'Internal source_corpus reference leaked into publication text: "{line}"'
        )

    # Link and image targets are structural: pandoc consumes them and the
    # reader never sees them. Figures are written under the run directory,
    # whose name is derived from the prompt, so scanning raw markdown
    # reported "prompt leaked into publication text" for any run that
    # renders an image — reliably so for a CJK prompt, which has no spaces
    # to trim and therefore lands in the directory name whole.
    prose_only = re.sub(r"\]\([^)]*\)", "]()", md_content)
    normalized_prose = " ".join(prose_only.split()).lower()
    for fragment in forbidden_fragments or []:
        clean = " ".join(fragment.split())
        if len(clean) < 20:
            continue
        if clean.lower() in normalized_prose:
            # Name the text, not just the rule. Without the fragment quoted
            # there is nothing to search the drafts for, and the author has to
            # read this function's source to find out what tripped it.
            excerpt = clean if len(clean) <= 80 else clean[:80] + "…"
            issues.append(
                f'Raw prompt fragment leaked into publication text: "{excerpt}"'
            )
            break

    # 4. Unresolved CITE markers
    cites = re.findall(r'\[CITE:[^\]]+\]', md_content)
    if cites:
        issues.append(f"Unresolved [CITE:] markers ({len(cites)}): {cites[:3]}")

    # 5. Internal markers that should not be in publication text
    internal_markers = re.findall(r'\[(?:Source|graphify|Note):[^\]]*\]', md_content)
    if internal_markers:
        issues.append(f"Internal markers in publication text: {internal_markers[:3]}")

    # 6. ASCII art code fences (box-drawing characters)
    code_fences = re.findall(r'```[\s\S]*?```', md_content)
    ascii_art_fences = [
        f for f in code_fences
        if _is_ascii_art(f)
    ]
    if ascii_art_fences:
        issues.append(
            f"ASCII art code fences ({len(ascii_art_fences)}) will render poorly in DOCX; "
            f"replace with mermaid diagrams or remove"
        )

    # 7. Orphan appendix fragments
    if re.search(r'traceability[_\s]appendix', md_content, re.IGNORECASE):
        issues.append("Traceability appendix content found in main document")

    # 8. End-of-report sentinel
    if re.search(r'End\s+of\s+Main\s+Report', md_content, re.IGNORECASE):
        issues.append("'End of Main Report' sentinel still present")

    # 9. Facts freeze verification
    if facts_freeze:
        for fact_key, expected_value in facts_freeze.items():
            if expected_value and expected_value not in md_content:
                issues.append(
                    f"Facts freeze violation: '{fact_key}' expected value "
                    f"'{expected_value}' not found in document"
                )

    return issues


def _filter_body_reference_lines(lines: list[str], strict: bool = True) -> list[str]:
    """Drop obviously internal or non-publication reference entries.

    ``strict`` additionally requires publication-shaped entries (DOI, arXiv,
    venue token, or italicized title). That is right for academic profiles,
    where internal-file citations are junk — but a technical document or
    business report legitimately cites internal documents (proposals,
    monthly reports, handbooks), so non-academic profiles pass
    ``strict=False`` and keep authored entries.
    """
    filtered = []
    internal_patterns = [
        r"\bsource_corpus\b",
        r"\bsource & corpus\b",
        r"\[Text file\]",
        r"\[Word document\]",
        r"\[Data file\]",
        r"\[Dataset\]",
        r"\bGRAPH_REPORT\b",
        r"\bmain_report\b",
        r"\bgraphify\b",
        r"https?://www\.backtrader\.com/?",
    ]
    for line in lines:
        lowered = line.lower()
        if any(re.search(pattern, lowered, re.IGNORECASE) for pattern in internal_patterns):
            continue
        if strict and not (
            re.search(r"doi[:\s]+10\.", line, re.IGNORECASE)
            or re.search(r"arxiv[:\s]+|\d{4}\.\d{4,5}", line, re.IGNORECASE)
            or re.search(r"\b(journal|proceedings|press|wiley|springer|elsevier|cambridge|oxford|mit press)\b", line, re.IGNORECASE)
            or re.search(r"\*[^*]+\*", line)
        ):
            continue
        filtered.append(line)
    return filtered


# ------------------------------------------------------------------
# Post-render table styling
# ------------------------------------------------------------------

def _style_tables_post_render(docx_path: str) -> None:
    """Apply professional table styling to all tables in the rendered DOCX.

    Adds grid borders and bold header row styling that pandoc's default
    output lacks.
    """
    try:
        doc = Document(docx_path)
        if not doc.tables:
            return

        for table in doc.tables:
            # Set table style to Table Grid when the reference document has it.
            # Some custom reference docs omit this style, so fall back to manual
            # borders instead of warning and leaving plain, borderless tables.
            try:
                table.style = doc.styles['Table Grid']
            except KeyError:
                _apply_manual_table_borders(table)

            # Bold the header row (first row)
            if table.rows:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                    # Light grey background for header cells
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'D9E2F3')  # Light blue-grey
                    shading.set(qn('w:val'), 'clear')
                    cell._tc.get_or_add_tcPr().append(shading)

        doc.save(docx_path)
        logger.info(f"[DOCX_RENDER] Table styling applied to {len(doc.tables)} table(s)")
    except Exception as exc:
        logger.warning(f"[DOCX_RENDER] Post-render table styling failed (non-fatal): {exc}")


def _apply_manual_table_borders(table) -> None:
    """Apply basic borders directly through WordprocessingML."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def _repair_missing_figures(docx_path: str, md_content: str) -> int:
    """Insert image files when pandoc failed to embed local markdown images."""
    image_matches = list(_IMAGE_LINK_RE.finditer(md_content))
    if not image_matches:
        return 0

    doc = Document(docx_path)
    if len(doc.inline_shapes) >= len(image_matches):
        return 0

    inserted = 0
    for match in image_matches:
        alt = match.group(1).strip()
        image_path = Path(match.group(2).strip().strip("<>"))
        if not image_path.exists():
            continue
        # Insert before the matching caption/alt paragraph when possible.
        target_para = None
        for para in doc.paragraphs:
            if alt and alt in para.text:
                target_para = para
                break
        if target_para is None:
            target_para = doc.add_paragraph()
        pic_para = target_para.insert_paragraph_before()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        run.add_picture(str(image_path), width=Inches(5.8))
        inserted += 1

    if inserted:
        doc.save(docx_path)
    return inserted


# ------------------------------------------------------------------
# Main node
# ------------------------------------------------------------------

def run_docx_render(state: ReportState) -> ReportState:
    """T15: DOCX_RENDER - convert markdown to .docx.

    Primary: pandoc with reference.docx template.
    Fallback: python-docx regex-based converter.
    Post-render: validates output structure.
    """
    qa_decision = state.qa.get("qa_decision")

    if qa_decision and qa_decision != "pass":
        raise QAHardBlockError(f"QA gate failed: {qa_decision}")

    # Select draft path per policy
    family = state.spec.get("report_profile", "academic_paper")
    policy = get_policy(family)

    if policy.citation.draft_prefer_marker_stripped:
        cited_md_path = state.drafts.get("publication_style_draft")
        if not cited_md_path or not Path(cited_md_path).exists():
            cited_md_path = state.drafts.get("merged_draft_cited_md")
        if not cited_md_path or not Path(cited_md_path).exists():
            cited_md_path = state.drafts.get("publication_draft_md")
    else:
        cited_md_path = state.drafts.get("publication_style_draft")
        if not cited_md_path or not Path(cited_md_path).exists():
            cited_md_path = state.drafts.get("merged_draft_cited_md")
        if not cited_md_path or not Path(cited_md_path).exists():
            cited_md_path = state.drafts.get("merged_draft_md")

    if not cited_md_path or not Path(cited_md_path).exists():
        raise QAHardBlockError("No merged draft found")

    with open(cited_md_path, encoding="utf-8") as f:
        md_content = f.read()

    if not md_content.strip():
        raise QAHardBlockError("Merged draft is empty")
    if PLACEHOLDER_TEXT in md_content:
        raise QAHardBlockError("Merged draft contains placeholder content")

    # Inject front matter at the top of the document
    front_matter_md = state.plan.get("front_matter_md", "")
    has_front_matter = bool(front_matter_md)
    if front_matter_md:
        md_content = front_matter_md + "\n\n---\n\n" + md_content

    # Remove orphan content after the canonical end of the main body
    md_content = _truncate_orphan_sections(md_content)

    # --- Mermaid diagram conversion (before sanity check) ---
    run_dir = WORKFLOW_RUNS_DIR / state.job_id
    md_content, mermaid_count = _convert_mermaid_blocks(md_content, run_dir)
    if mermaid_count > 0:
        state.output["mermaid_figures_converted"] = mermaid_count

    # Source tables first, so table numbering runs in reading order and the
    # figure pass continues the sequence rather than restarting it.
    evidence_ledger = load_jsonl(state.sources.get("evidence_ledger_path", ""))
    md_content, placed_tables, unresolved_tables = replace_table_placeholders(
        md_content, evidence_ledger
    )
    if placed_tables:
        state.output["source_tables_placed"] = placed_tables
    if unresolved_tables:
        state.runtime.setdefault("warnings", []).append(
            "Unresolved source table placeholder(s): "
            + ", ".join(sorted(set(unresolved_tables)))
        )

    figure_manifest = _load_figure_manifest(state.output.get("figure_manifest_path", ""))
    md_content, resolved_figures, unresolved_figures = _replace_figure_placeholders(
        md_content, figure_manifest, table_start_number=placed_tables
    )
    if resolved_figures:
        state.output["figure_placeholders_resolved"] = resolved_figures
    if unresolved_figures:
        state.runtime.setdefault("warnings", []).append(
            "Unresolved figure placeholder(s): " + ", ".join(sorted(set(unresolved_figures)))
        )

    # --- Pre-render sanity gate ---
    facts_freeze = state.plan.get("facts_freeze")
    user_prompt = state.spec.get("user_prompt", "")
    prompt_fragments = []
    if user_prompt:
        prompt_fragments.append(user_prompt)
        prompt_fragments.append(" ".join(user_prompt.split()[:12]))
    sanity_issues = _pre_render_sanity_check(md_content, facts_freeze, prompt_fragments)
    if sanity_issues:
        issue_list = "; ".join(sanity_issues)
        logger.error(f"[DOCX_RENDER] Pre-render sanity check FAILED: {issue_list}")
        raise QAHardBlockError(
            f"Pre-render sanity check failed ({len(sanity_issues)} issue(s)): {issue_list}"
        )

    # Append publication reference list
    pub_ref_list_path = state.citations.get("publication_reference_list_path", "")
    pub_ref_md = ""
    if pub_ref_list_path and Path(pub_ref_list_path).exists():
        with open(pub_ref_list_path, encoding="utf-8") as f:
            pub_ref_md = f.read()

    # Handle inline References section and always remove it from body markdown.
    # If generated publication refs exist, they take precedence. Otherwise reuse
    # curated body references after filtering internal/project-only entries.
    # Academic profiles keep the strict publication-shape filter; other
    # profiles legitimately cite internal documents, so authored entries
    # survive and do not depend on the citation chain having curated anything.
    strict_refs = state.spec.get("report_profile", "academic_paper") in (
        "academic_paper", "admissions_report", "admissions_project_report"
    )
    md_content, body_refs_md = _split_body_references(md_content, strict_refs=strict_refs)

    curated_count = int(state.citations.get("curated_reference_count", 0) or 0)
    if not pub_ref_md.strip() and body_refs_md.strip() and (curated_count > 0 or not strict_refs):
        pub_ref_md = body_refs_md

    # Append references to the end of the markdown for pandoc. The reference
    # artifact carries an English "## References" marker; localize it to the
    # blueprint title when the document language is Chinese.
    pub_ref_md = _localize_reference_heading(
        pub_ref_md, md_content, state.plan.get("blueprint") or {}
    )
    if pub_ref_md.strip():
        md_content = md_content.rstrip() + "\n\n" + pub_ref_md

    # The generated Sources list. Without it a document citing project sources
    # ships with markers pointing at nothing — which is what happened before:
    # the markers were deleted instead, and the reader could check no figure
    # in the document against anything.
    source_list_path = state.citations.get("publication_source_list_path", "")
    if source_list_path and Path(source_list_path).exists():
        source_list_md = Path(source_list_path).read_text(encoding="utf-8")
        if source_list_md.strip():
            if detect_document_language(md_content) == "zh":
                source_list_md = source_list_md.replace(
                    SOURCE_LIST_HEADING, SOURCE_LIST_HEADING_ZH, 1
                )
            md_content = md_content.rstrip() + "\n\n" + source_list_md

    run_dir = WORKFLOW_RUNS_DIR / state.job_id
    md_content = _absolutize_image_paths(
        md_content, run_dir, Path(cited_md_path).parent
    )
    if detect_document_language(md_content) == "zh":
        md_content = _normalize_cjk_typography(md_content)
    final_docx_path = run_dir / "rendered_report.docx"

    # Write the final markdown to a temp file for pandoc. The TOC field is
    # injected here (pandoc path only): --toc would place it before the
    # title page, and the fallback converter cannot consume raw openxml.
    toc_blueprint = state.plan.get("blueprint") or {}
    cover_title = ""
    if not has_front_matter and (toc_blueprint.get("section_order") or [])[:1] == ["cover"]:
        cover_title = localized_section_title(
            (toc_blueprint.get("sections") or {}).get("cover"),
            "cover",
            detect_document_language(md_content),
        )
    title_leads = (
        not has_front_matter
        and state.spec.get("task_intent") == "revise_existing"
    )
    pandoc_input_md = run_dir / "pandoc_input.md"
    with open(pandoc_input_md, "w", encoding="utf-8") as f:
        f.write(
            _inject_toc(
                md_content,
                has_front_matter,
                cover_title=cover_title,
                title_leads=title_leads,
            )
        )

    # --- Primary path: pandoc ---
    # Resolve outside the try: an unusable custom template must hard-block,
    # not get swallowed into the python-docx fallback.
    reference_doc = _resolve_reference_doc(state.spec)
    custom_template_requested = bool(str(state.spec.get("reference_docx_path") or "").strip())
    used_pandoc = False
    pandoc_warnings: list[str] = []
    try:
        used_pandoc = _render_via_pandoc(
            str(pandoc_input_md),
            str(final_docx_path),
            reference_doc=reference_doc,
            warnings_out=pandoc_warnings,
        )
    except Exception as exc:
        logger.warning(f"[DOCX_RENDER] pandoc path failed, falling back: {exc}")
        used_pandoc = False

    if not used_pandoc and custom_template_requested:
        raise QAHardBlockError(
            "Custom reference docx was requested but the pandoc render did not "
            "run; the python-docx fallback cannot apply a template, and "
            "shipping the default formatting instead would be the wrong document."
        )

    # --- Fallback path: python-docx ---
    if not used_pandoc:
        logger.info("[DOCX_RENDER] Using legacy python-docx converter")
        try:
            # Strip reference section from md_content for legacy converter
            # (it adds references separately with hanging indent)
            md_for_legacy = re.sub(
                rf"^#{{1,6}} {_REFS_HEADING_NAMES}\s*\n.*",
                "",
                md_content,
                flags=re.MULTILINE | re.DOTALL,
            )
            markdown_to_docx(md_for_legacy, str(final_docx_path), figure_manifest=figure_manifest)

            # Add APA references separately
            if pub_ref_md.strip():
                doc = Document(str(final_docx_path))
                _add_hanging_indent_references(doc, pub_ref_md)
                doc.save(str(final_docx_path))
        except Exception as exc:
            logger.exception("[DOCX_RENDER] legacy python-docx fallback failed")
            state.runtime["error"] = f"DOCX_RENDER failed: {type(exc).__name__}: {exc}"
            raise QAHardBlockError(f"DOCX render failed: {exc}") from exc

    # --- Post-render table styling ---
    if final_docx_path.exists():
        repaired_figures = _repair_missing_figures(str(final_docx_path), md_content)
        if repaired_figures:
            state.output["post_render_figures_repaired"] = repaired_figures
        _style_tables_post_render(str(final_docx_path))

    # --- Post-render validation ---
    validation_issues = _validate_docx(str(final_docx_path), md_content)
    # The renderer's own account of what it could not do belongs with the
    # checks that read the file afterwards, not in a log line nobody keeps.
    validation_issues.extend(f"pandoc: {item}" for item in pandoc_warnings)
    if custom_template_requested:
        carried = _reference_doc_body_carryover(reference_doc)
        if carried:
            validation_issues.append(
                "the reference .docx supplied its styles but not its own content; "
                f"this text is in the template and not in the report: {carried}"
            )
    if validation_issues:
        logger.warning(
            f"[DOCX_RENDER] Post-render validation issues: {'; '.join(validation_issues)}"
        )
        # Store issues but do not hard-block; the file may still be usable.
        state.runtime["docx_validation_issues"] = validation_issues

    renderer_used = "pandoc" if used_pandoc else "python-docx (fallback)"
    logger.info(
        f"[DOCX_RENDER] Done: renderer={renderer_used}, "
        f"output={final_docx_path}, "
        f"size={final_docx_path.stat().st_size if final_docx_path.exists() else 0} bytes"
    )

    # Surface renderer fallback as a warning the agent can see
    if not used_pandoc:
        state.runtime.setdefault("warnings", []).append(
            "?? DOCX was rendered using the python-docx fallback converter instead of pandoc. "
            "Output quality is degraded: tables may be broken, no table of contents, "
            "and complex formatting may be lost. Install pandoc for production-quality output: "
            "winget install JohnMacFarlane.Pandoc (Windows) / apt install pandoc (Linux)"
        )

    state.output["final_docx_path"] = str(final_docx_path)
    state.output["rendered_docx_path"] = str(final_docx_path)
    state.output["renderer_used"] = renderer_used
    # The template that was actually used, which is the user's whenever they
    # supplied one. Recording the packaged default here sent both template
    # reports to read a file the author never gave: the style map compared the
    # output against the built-in reference, and the field-fill report scanned
    # the built-in for placeholders, found none, and reported field_count 0 with
    # status pass — a clean bill of health for a template whose fields it had
    # not looked at.
    state.output["reference_docx_path"] = str(reference_doc) if reference_doc.exists() else ""
    state.output["reference_docx_applied"] = bool(used_pandoc and reference_doc.exists())
    return state
