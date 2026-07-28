"""POST_RENDER_VALIDATE - structural QA for rendered DOCX."""
from __future__ import annotations

import json
import re
from collections import Counter
from pathlib import Path
from typing import Any

from docx import Document

from ..errors import QAHardBlockError
from ..runtime_support import write_json_artifact
from ..state import ReportState


MAX_FRONT_MATTER_PARAGRAPHS = 8
MAX_HEADING_ENTRIES = 40
MAX_TABLE_ENTRIES = 25
MAX_TABLE_ROW_PREVIEW_CELLS = 8


def _docx_text_with_tables(doc: Document) -> str:
    parts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _outline_figure_count(state: ReportState) -> int:
    outline = state.plan.get("outline", {}) or {}
    count = 0
    for section in (outline.get("sections") or {}).values():
        fids = section.get("figure_ids", []) or []
        if isinstance(fids, list):
            count += len(fids)
    return count


def _outline_figure_ids(state: ReportState) -> set[str]:
    outline = state.plan.get("outline", {}) or {}
    ids: set[str] = set()
    for section in (outline.get("sections") or {}).values():
        fids = section.get("figure_ids", []) or []
        if isinstance(fids, list):
            ids.update(str(fid).lower() for fid in fids if str(fid).strip())
    return ids


def _manifest_figures(state: ReportState) -> list[dict]:
    manifest_path = state.output.get("figure_manifest_path", "")
    if not manifest_path or not Path(manifest_path).exists():
        return []
    try:
        with open(manifest_path, encoding="utf-8") as f:
            manifest = json.load(f)
        figures = manifest.get("figures", [])
        return [f for f in figures if isinstance(f, dict)]
    except Exception:
        return []


def _expected_figure_count(state: ReportState) -> int:
    figures = _manifest_figures(state)
    native_tables = sum(
        1 for f in figures if str(f.get("render_mode") or "") == "native_table"
    )
    image_count = len(figures) - native_tables
    # Outline may declare figures even when no manifest was produced (e.g.
    # matplotlib unavailable, malformed figure_plan.json). Treat outline as an
    # additional upper bound so silently-missing figures surface at render
    # time — minus the figures that legitimately render as native tables.
    return max(image_count, _outline_figure_count(state) - native_tables)


#: "Figure 3." opening a paragraph, or "圖 3." — the caption itself.
_FIGURE_CAPTION_RE = re.compile(
    r"^(?:(?:Figure|Fig\.?)\s+|(?:圖|图)\s*)(\d+|[a-z])[:.、．：]\s*",
    re.IGNORECASE,
)

#: The same figure named inside a sentence — "as Figure 3 shows", "由圖 3 可見".
#: Both of these were English-only, so the three checks built on them — a
#: reference with no caption, references with no embedded figure at all, and
#: references to figures the outline never declared — had never run against a
#: Chinese report. A word boundary does nothing next to CJK, hence the two
#: alternatives rather than one pattern.
_FIGURE_MENTION_RE = re.compile(
    r"(?:\b(?:Figure|Fig\.?)\s+|(?:圖|图)\s*)(\d+|[a-z])(?![0-9])",
    re.IGNORECASE,
)


#: A figure caption, in either language the renderer produces. The duplicate
#: check matched "Figure 1." only, so a Chinese report — the ordinary case
#: here — could carry two captions both reading "圖 1." and pass. The same
#: English-shaped rule that dropped short CJK headings and called a finished
#: Chinese report incomplete, in a third place.
_CAPTION_LABEL_RE = re.compile(
    r"^((?:Figure|Fig\.?|圖|图)\s*\d+[.、．:：]?\s+[^.。]+[.。]?)",
    re.IGNORECASE,
)


def _figure_build_error_hint(state: ReportState) -> str:
    """Why the figure never arrived, if the builder already recorded it.

    FIGURE_BUILD writes its failures into the manifest and carries on. That
    reason is the one thing an author needs and it was going unread.
    """
    manifest_path = state.output.get("figure_manifest_path", "")
    if not manifest_path or not Path(manifest_path).exists():
        return ""
    try:
        with open(manifest_path, encoding="utf-8") as f:
            manifest = json.load(f)
    except Exception:
        return ""
    errors = [str(e).strip() for e in (manifest.get("errors") or []) if str(e).strip()]
    if not errors:
        return ""
    return " — figure build reported: " + "; ".join(errors[:3])


def _figure_shortfall_hint(
    state: ReportState, docx_text: str, outline_figure_ids: set[str]
) -> str:
    """Explain why figures are missing, using ids the run already knows.

    A count on its own gives the author nothing to act on. The usual cause is
    an id mismatch: the draft cites a figure id the builder never assigned, so
    the placeholder is copied through unreplaced and the DOCX carries no image.
    Both halves of that mismatch are already in hand here.
    """
    built_ids = [
        str(figure.get("figure_id", "")).strip()
        for figure in _manifest_figures(state)
        if str(figure.get("figure_id", "")).strip()
    ]
    unresolved = sorted(set(re.findall(r"\[FIGURE:\s*([^\]\s]+)", docx_text, re.IGNORECASE)))

    parts: list[str] = []
    if unresolved:
        parts.append(
            "unresolved placeholder(s) " + ", ".join(f"[FIGURE:{fid}]" for fid in unresolved)
        )
    unknown = sorted(fid for fid in outline_figure_ids if fid not in {b.lower() for b in built_ids})
    if unknown:
        parts.append("outline figure_ids with no built figure: " + ", ".join(unknown))
    if not parts and built_ids:
        # Nothing is malformed — the figures were built and simply never
        # referenced. Saying so is the difference between an author looking for
        # a typo and an author placing a figure.
        parts.append(
            "no section references them; place [FIGURE:<id>] where the figure "
            "belongs, or drop the figure"
        )
    if not parts:
        return ""
    built = ", ".join(built_ids) if built_ids else "none"
    return f" — {'; '.join(parts)}. Built figure ids: {built}."


def _expected_table_count(state: ReportState) -> int:
    count = sum(
        1
        for f in _manifest_figures(state)
        if str(f.get("render_mode") or "") == "native_table"
    )
    draft_path = state.drafts.get("publication_style_draft") or state.drafts.get("merged_draft_cited_md")
    if draft_path and Path(draft_path).exists():
        markdown = Path(draft_path).read_text(encoding="utf-8")
        count += len(re.findall(r"^\|.+\|\s*$\n^\|?\s*:?-{3,}", markdown, re.MULTILINE))
    return count


def _clean_preview(text: str, limit: int = 160) -> str:
    cleaned = " ".join(str(text or "").split())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 3].rstrip() + "..."


def _style_name(paragraph: Any) -> str:
    try:
        return paragraph.style.name or "unknown"
    except Exception:
        return "unknown"


def _paragraph_style_counts(doc: Document) -> dict[str, int]:
    counts = Counter(_style_name(para) for para in doc.paragraphs)
    return dict(sorted(counts.items()))


def _heading_summary(doc: Document) -> list[dict[str, Any]]:
    headings: list[dict[str, Any]] = []
    for index, para in enumerate(doc.paragraphs):
        style_name = _style_name(para)
        if not style_name.startswith("Heading"):
            continue
        text = _clean_preview(para.text)
        if not text:
            continue
        headings.append({
            "paragraph_index": index,
            "style": style_name,
            "text": text,
        })
        if len(headings) >= MAX_HEADING_ENTRIES:
            break
    return headings


def _table_summaries(doc: Document) -> list[dict[str, Any]]:
    summaries: list[dict[str, Any]] = []
    for index, table in enumerate(doc.tables):
        row_count = len(table.rows)
        column_count = max((len(row.cells) for row in table.rows), default=0)
        non_empty_cells = 0
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    non_empty_cells += 1

        first_row_preview: list[str] = []
        if row_count:
            first_row_preview = [
                _clean_preview(cell.text, 80)
                for cell in table.rows[0].cells[:MAX_TABLE_ROW_PREVIEW_CELLS]
            ]

        summaries.append({
            "table_index": index,
            "rows": row_count,
            "columns": column_count,
            "non_empty_cells": non_empty_cells,
            "first_row_preview": first_row_preview,
        })
        if len(summaries) >= MAX_TABLE_ENTRIES:
            break
    return summaries


def _existing_report_path(state: ReportState, key: str) -> str:
    path = state.runtime.get(key) or state.output.get(key) or state.qa.get(key)
    if path and Path(path).exists():
        return str(path)
    return ""


def _build_layout_manifest(
    state: ReportState,
    doc: Document,
    docx_path: str,
    issues: list[str],
    expected_tables: int,
    expected_figures: int,
) -> dict[str, Any]:
    path = Path(docx_path)
    front_matter = [
        _clean_preview(para.text)
        for para in doc.paragraphs
        if para.text.strip()
    ][:MAX_FRONT_MATTER_PARAGRAPHS]

    return {
        "job_id": state.job_id,
        "status": "passed" if not issues else "failed",
        "report_profile": state.spec.get("report_profile", ""),
        "renderer_used": state.output.get("renderer_used", "unknown"),
        "docx": {
            "path": str(path),
            "file_size_bytes": path.stat().st_size if path.exists() else 0,
        },
        "counts": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
            "expected_tables": expected_tables,
            "expected_figures": expected_figures,
        },
        "paragraph_style_counts": _paragraph_style_counts(doc),
        "front_matter_preview": front_matter,
        "headings": _heading_summary(doc),
        "tables": _table_summaries(doc),
        "related_reports": {
            "post_render_repair_report": _existing_report_path(state, "post_render_repair_report_path"),
            "post_render_validate_report": _existing_report_path(state, "post_render_validate_report_path"),
            "visual_render_check_report": _existing_report_path(state, "visual_render_check_report_path"),
        },
        "issues": list(issues),
    }


def run_post_render_validate(state: ReportState) -> ReportState:
    """Hard-block rendered DOCX defects that require manual repair."""
    docx_path = state.output.get("final_docx_path") or state.output.get("rendered_docx_path")
    if not docx_path or not Path(docx_path).exists():
        raise QAHardBlockError("POST_RENDER_VALIDATE: rendered DOCX is missing")

    doc = Document(docx_path)
    text = _docx_text_with_tables(doc)
    issues: list[str] = []

    expected_figures = _expected_figure_count(state)
    expected_tables = _expected_table_count(state)
    outline_figure_ids = _outline_figure_ids(state)

    # A literal [FIGURE:id] in the deliverable is a defect on its own terms,
    # whatever the counts say. This used to be reachable only through the
    # count comparison below, and when every figure failed to build the
    # expected count fell to zero, matched the zero that were embedded, and
    # the check passed — while the placeholder text shipped in the document
    # the author hands in, under a "validation passed" message.
    leaked = sorted(set(re.findall(r"\[FIGURE:\s*([^\]\s]+)", text, re.IGNORECASE)))
    if leaked:
        issues.append(
            "unrendered placeholder(s) left in the document: "
            + ", ".join(f"[FIGURE:{fid}]" for fid in leaked)
            + _figure_build_error_hint(state)
        )

    if len(doc.inline_shapes) < expected_figures:
        issues.append(
            f"expected {expected_figures} embedded figure(s), found {len(doc.inline_shapes)}"
            + _figure_shortfall_hint(state, text, outline_figure_ids)
        )
    if len(doc.tables) < expected_tables:
        issues.append(f"expected {expected_tables} Word table(s), found {len(doc.tables)}")
    reference_heading_count = sum(
        1 for para in doc.paragraphs
        if para.style.name.startswith("Heading") and para.text.strip().lower() == "references"
    )
    if reference_heading_count > 1:
        issues.append(f"duplicate References headings found: {reference_heading_count}")

    figure_caption_labels: set[str] = set()
    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        match = _CAPTION_LABEL_RE.match(text)
        if not match:
            continue
        label = match.group(1).lower()
        if label in figure_caption_labels:
            issues.append(f"duplicate figure caption label: {match.group(1)}")
            break
        figure_caption_labels.add(label)

    leakage_patterns = [
        (r"\[CITE:", "unresolved CITE marker"),
        (r"\[Source:", "unresolved Source marker"),
        (r"\[FIGURE:", "unresolved FIGURE marker (figure not rendered)"),
        (r"\[graphify:", "unresolved graphify marker"),
        (r"source_corpus|claim_matrix|evidence_ledger", "internal artifact leakage"),
        (r"Revise the existing academic report|Use revised_report\.md", "prompt residue"),
        (r"author@example\.com|Independent Researcher|\{\.Title\}", "template metadata leakage"),
        (r"Research Author|Research University|research@university\.edu", "generic research metadata leakage"),
        (r"traceability appendix|End of Main Report|Appendix E|Appendices A and B", "appendix leakage"),
    ]
    for pattern, label in leakage_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            issues.append(label)

    front_paras = [para.text.strip() for para in doc.paragraphs if para.text.strip()][:8]
    front_text = "\n".join(front_paras)
    if "**" in front_text:
        issues.append("front matter contains leftover Markdown bold marker")
    if re.search(r"\[(?:Author Name|University|email@domain\.com|Your Name|INSERT .+?)\]", front_text, re.IGNORECASE):
        issues.append("front matter contains placeholder metadata")

    caption_ids: set[str] = set()
    mention_ids: set[str] = set()
    for para in doc.paragraphs:
        para_text = " ".join(para.text.split())
        caption_match = _FIGURE_CAPTION_RE.match(para_text)
        if caption_match:
            caption_ids.add(caption_match.group(1).lower())
            continue
        mention_ids.update(m.group(1).lower() for m in _FIGURE_MENTION_RE.finditer(para_text))

    if mention_ids:
        missing_captions = sorted(mention_ids - caption_ids)
        if missing_captions:
            issues.append("figure references without matching captions: " + ", ".join(missing_captions))
        if len(doc.inline_shapes) == 0:
            issues.append("figure references present but no embedded figures found")
        undeclared = sorted(mention_ids - outline_figure_ids) if outline_figure_ids else sorted(mention_ids)
        if undeclared:
            issues.append("figure references not declared in outline/manifest: " + ", ".join(undeclared))

    report = {
        "job_id": state.job_id,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "inline_shape_count": len(doc.inline_shapes),
        "expected_table_count": expected_tables,
        "expected_figure_count": expected_figures,
        "issues": issues,
        "status": "passed" if not issues else "failed",
    }
    state.runtime["post_render_validate_report_path"] = write_json_artifact(
        state, "post_render_validate_report.json", report
    )
    layout_manifest = _build_layout_manifest(
        state,
        doc,
        docx_path,
        issues,
        expected_tables,
        expected_figures,
    )
    state.runtime["post_render_layout_manifest_path"] = write_json_artifact(
        state, "post_render_layout_manifest.json", layout_manifest
    )
    state.output["post_render_layout_manifest_path"] = state.runtime["post_render_layout_manifest_path"]
    if issues:
        raise QAHardBlockError("POST_RENDER_VALIDATE: " + "; ".join(issues))
    return state
