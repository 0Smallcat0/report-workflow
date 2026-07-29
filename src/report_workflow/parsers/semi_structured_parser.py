"""Semi-structured parser for PDF, DOCX, TXT files."""
import re


def table_to_text(table: list[list[str]]) -> str:
    """Convert table cells into searchable text."""
    rows = []
    for row in table:
        rows.append(" | ".join(str(cell).strip() for cell in row if cell is not None))
    return "\n".join(rows)


def parse_pdf(file_path: str) -> dict:
    """Parse PDF using pdfplumber."""
    try:
        import pdfplumber
        blocks = []
        all_text = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    all_text.append(text)
                    blocks.append({
                        "block_id": f"page_{page_num}",
                        "block_type": "paragraph",
                        "content": text,
                        "page_number": page_num,
                        "table_data": None
                    })
                tables = page.extract_tables()
                for t_idx, table in enumerate(tables):
                    table_text = table_to_text(table)
                    blocks.append({
                        "block_id": f"page_{page_num}_table_{t_idx}",
                        "block_type": "table",
                        "content": table_text,
                        "page_number": page_num,
                        "table_data": table
                    })
                    if table_text:
                        all_text.append(table_text)
        return {
            "blocks": blocks,
            "raw_content": "\n\n".join(all_text),
            "success": True
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_docx(file_path: str) -> dict:
    """Parse DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(file_path)
        blocks = []
        all_text = []
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            all_text.append(text)
            block_type = "paragraph"
            if para.style.name.startswith("Heading"):
                block_type = "heading"
            blocks.append({
                "block_id": f"para_{i}",
                "block_type": block_type,
                "content": text,
                "page_number": None,
                "table_data": None
            })
        def _cell_text(raw: str) -> str:
            """One cell read as one cell.

            A header merged across two columns comes back from python-docx as
            both underlying cells, each holding the label twice with a newline
            between. Kept verbatim, the header row broke into three lines of
            different widths, the data row no longer lined up with it, and the
            row 機台 A1 / 5 / 7 reached the ledger as
            {"機台": "A1", "解析度(μm)\\n解析度(μm)": "7"} — the 5 μm reading
            dropped at ingestion and a column keyed on text nobody wrote.

            Repeated identical segments are the library's echo of one merged
            label, so one is kept; anything else joins with a space, because a
            cell holding two different lines really does hold both.
            """
            segments = [part.strip() for part in str(raw).splitlines() if part.strip()]
            if not segments:
                return ""
            if len(set(segments)) == 1:
                return segments[0]
            return " ".join(segments)

        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([_cell_text(cell.text) for cell in row.cells])
            table_text = table_to_text(rows)
            blocks.append({
                "block_id": f"table_{t_idx}",
                "block_type": "table",
                "content": table_text,
                "page_number": None,
                "table_data": rows
            })
            if table_text:
                all_text.append(table_text)
        return {
            "blocks": blocks,
            "raw_content": "\n\n".join(all_text),
            "success": True
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_txt(file_path: str) -> dict:
    """Parse TXT file by splitting into heading/paragraph/list/code blocks.

    Adds line_start, line_end, content_hash, source_file_path, and quote metadata
    to every block so evidence can be traced back to exact source locations.
    """
    import hashlib

    try:
        from .source_text import read_source_lines

        lines = read_source_lines(file_path)

        blocks = []
        all_text_parts = []
        block_counter = 0
        i = 0

        while i < len(lines):
            raw_line = lines[i]
            stripped = raw_line.strip()

            # ---------- heading ----------
            if stripped.startswith("#"):
                # Consume all leading heading lines
                heading_texts = []
                while i < len(lines) and lines[i].strip().startswith("#"):
                    heading_texts.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(heading_texts)
                line_start = i - len(heading_texts) + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"txt_{block_counter}",
                    "block_type": "heading",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content,
                })
                continue

            # ---------- code block ([code] ... [/code] or ``` ```) ----------
            if stripped.startswith("[code]") or stripped.startswith("```"):
                code_lines = []
                code_start = i
                # Consume until closing delimiter
                i += 1
                while i < len(lines):
                    end_marker = lines[i].strip()
                    if end_marker.startswith("[/code]") or end_marker.startswith("```"):
                        i += 1
                        break
                    code_lines.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(code_lines)
                line_start = code_start + 1
                line_end = i - 1
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"txt_{block_counter}",
                    "block_type": "code_block",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                all_text_parts.append(content)
                continue

            # ---------- delimited table ----------
            table_rows, table_end = _delimited_table_rows(lines, i)
            if table_rows:
                table_lines = [lines[n].rstrip() for n in range(i, table_end)]
                content = "\n".join(table_lines)
                all_text_parts.extend(table_lines)
                line_start = i + 1
                line_end = table_end
                i = table_end
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "table",
                    "content": content,
                    "page_number": None,
                    "table_data": table_rows,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- list item ----------
            if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
                list_items = []
                list_start = i
                while i < len(lines):
                    line_stripped = lines[i].strip()
                    is_bullet = (
                        line_stripped.startswith("- ")
                        or line_stripped.startswith("* ")
                        or line_stripped.startswith("+ ")
                    )
                    # A wrapped bullet continues on indented lines; splitting
                    # them into separate blocks fragments one logical entry
                    # (e.g. one literature citation) across evidence units.
                    is_continuation = (
                        bool(list_items)
                        and bool(line_stripped)
                        and lines[i][:1] in (" ", "\t")
                        and not is_bullet
                    )
                    if not (is_bullet or is_continuation):
                        break
                    list_items.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(list_items)
                line_start = list_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"txt_{block_counter}",
                    "block_type": "list_item",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- numbered list ----------
            import re as _re
            numbered = _re.compile(r"^\d+[.)]\s").match
            if numbered(stripped):
                num_items = []
                num_start = i
                while i < len(lines) and numbered(lines[i].strip()):
                    num_items.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(num_items)
                line_start = num_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"txt_{block_counter}",
                    "block_type": "list_item",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- blank line ----------
            if not stripped:
                i += 1
                continue

            # ---------- paragraph (collect non-blank lines until blank or heading) ----------
            para_lines = []
            para_start = i
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#"):
                # A pasted table usually follows its lead-in sentence with no
                # blank line between them, so the paragraph has to yield to it
                # or the whole table is swallowed as prose.
                if para_lines and _delimited_table_rows(lines, i)[0]:
                    break
                para_lines.append(lines[i].rstrip())
                all_text_parts.append(lines[i].rstrip())
                i += 1
            if para_lines:
                content = "\n".join(para_lines)
                line_start = para_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"txt_{block_counter}",
                    "block_type": "paragraph",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })

        return {
            "blocks": blocks,
            "raw_content": "\n".join(all_text_parts),
            "success": True,
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


_PIPE_SEPARATOR_ROW_RE = re.compile(r"^\s*\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")
_MIN_TABLE_ROWS = 2
_MIN_TABLE_COLUMNS = 2


def _split_table_row(line: str) -> list[str] | None:
    """Cells of one delimited row, or None when the line is not one.

    Two notations reach us. Pasting a selection out of a spreadsheet yields
    tab-separated lines — the most common way a measurement table enters a
    notes file — and hand-written markdown yields pipe rows. Both are tables;
    only the delimiter differs.
    """
    stripped = line.strip()
    if not stripped:
        return None
    if "\t" in line:
        cells = [cell.strip() for cell in line.rstrip("\n").split("\t")]
    elif stripped.count("|") >= _MIN_TABLE_COLUMNS - 1 and _PIPE_SEPARATOR_ROW_RE.match(stripped) is None:
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    else:
        return None
    if len(cells) < _MIN_TABLE_COLUMNS:
        return None
    if not any(cells):
        return None
    return cells


def _delimited_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Consume a delimited table starting at ``start``.

    Returns ``(rows, end_index)``, or ``([], start)`` when there is no table
    here. A run must hold at least two rows of the same column count before it
    counts, so an ordinary sentence that happens to contain a tab or a vertical
    bar stays prose.
    """
    first = _split_table_row(lines[start])
    if first is None:
        return [], start

    width = len(first)
    rows = [first]
    index = start + 1
    while index < len(lines):
        if _PIPE_SEPARATOR_ROW_RE.match(lines[index].strip()):
            # Markdown's header underline carries no data.
            index += 1
            continue
        cells = _split_table_row(lines[index])
        if cells is None or len(cells) != width:
            break
        rows.append(cells)
        index += 1

    if len(rows) < _MIN_TABLE_ROWS:
        return [], start
    return rows, index


def parse_markdown(file_path: str) -> dict:
    """Parse Markdown file by splitting into heading/paragraph/list/code blocks.

    Adds line_start, line_end, content_hash, source_file_path, and quote metadata
    to every block so evidence can be traced back to exact source locations.
    Markdown headings start with # (like parse_txt for TXT files).
    """
    import hashlib

    try:
        from .source_text import read_source_lines

        lines = read_source_lines(file_path)

        blocks = []
        all_text_parts = []
        block_counter = 0
        i = 0

        while i < len(lines):
            raw_line = lines[i]
            stripped = raw_line.strip()

            # ---------- heading ----------
            if stripped.startswith("#"):
                # Consume all leading heading lines
                heading_texts = []
                while i < len(lines) and lines[i].strip().startswith("#"):
                    heading_texts.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(heading_texts)
                line_start = i - len(heading_texts) + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "heading",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content,
                })
                continue

            # ---------- code block (``` ``` or [code] [/code]) ----------
            if stripped.startswith("[code]") or stripped.startswith("```"):
                code_lines = []
                code_start = i
                i += 1
                while i < len(lines):
                    end_marker = lines[i].strip()
                    if end_marker.startswith("[/code]") or end_marker.startswith("```"):
                        i += 1
                        break
                    code_lines.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(code_lines)
                line_start = code_start + 1
                line_end = i - 1
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "code_block",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                all_text_parts.append(content)
                continue

            # ---------- delimited table ----------
            table_rows, table_end = _delimited_table_rows(lines, i)
            if table_rows:
                table_lines = [lines[n].rstrip() for n in range(i, table_end)]
                content = "\n".join(table_lines)
                all_text_parts.extend(table_lines)
                line_start = i + 1
                line_end = table_end
                i = table_end
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "table",
                    "content": content,
                    "page_number": None,
                    "table_data": table_rows,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- list item ----------
            if stripped.startswith("- ") or stripped.startswith("* ") or stripped.startswith("+ "):
                list_items = []
                list_start = i
                while i < len(lines):
                    line_stripped = lines[i].strip()
                    is_bullet = (
                        line_stripped.startswith("- ")
                        or line_stripped.startswith("* ")
                        or line_stripped.startswith("+ ")
                    )
                    # A wrapped bullet continues on indented lines; splitting
                    # them into separate blocks fragments one logical entry
                    # (e.g. one literature citation) across evidence units.
                    is_continuation = (
                        bool(list_items)
                        and bool(line_stripped)
                        and lines[i][:1] in (" ", "\t")
                        and not is_bullet
                    )
                    if not (is_bullet or is_continuation):
                        break
                    list_items.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(list_items)
                line_start = list_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "list_item",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- numbered list ----------
            import re as _re
            numbered = _re.compile(r"^\d+[.)]\s").match
            if numbered(stripped):
                num_items = []
                num_start = i
                while i < len(lines) and numbered(lines[i].strip()):
                    num_items.append(lines[i].rstrip())
                    all_text_parts.append(lines[i].rstrip())
                    i += 1
                content = "\n".join(num_items)
                line_start = num_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "list_item",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })
                continue

            # ---------- blank line ----------
            if not stripped:
                i += 1
                continue

            # ---------- paragraph ----------
            para_lines = []
            para_start = i
            while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#"):
                # A pasted table usually follows its lead-in sentence with no
                # blank line between them, so the paragraph has to yield to it
                # or the whole table is swallowed as prose.
                if para_lines and _delimited_table_rows(lines, i)[0]:
                    break
                para_lines.append(lines[i].rstrip())
                all_text_parts.append(lines[i].rstrip())
                i += 1
            if para_lines:
                content = "\n".join(para_lines)
                line_start = para_start + 1
                line_end = i
                block_counter += 1
                content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]
                blocks.append({
                    "block_id": f"md_{block_counter}",
                    "block_type": "paragraph",
                    "content": content,
                    "page_number": None,
                    "table_data": None,
                    "source_file_path": file_path,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_hash": content_hash,
                    "quote": content[:200] + ("..." if len(content) > 200 else ""),
                })

        return {
            "blocks": blocks,
            "raw_content": "\n".join(all_text_parts),
            "success": True,
        }
    except Exception as e:
        return {"blocks": [], "error": str(e), "success": False}


def parse_semi_structured(file_path: str, file_type: str) -> dict:
    """Parse semi-structured file and return content blocks."""
    if file_type == "pdf":
        result = parse_pdf(file_path)
    elif file_type == "docx":
        result = parse_docx(file_path)
    elif file_type in ("txt", "md"):
        # md is handled by parse_markdown (same block structure as txt)
        result = parse_markdown(file_path)
    else:
        return {"blocks": [], "error": f"Unsupported file type: {file_type}", "success": False}

    # Enrich blocks with Chinese heuristic annotations when CJK content detected
    if result.get("success") and result.get("blocks"):
        raw = result.get("raw_content", "")
        if _has_cjk_content(raw):
            annotate_chinese_heuristics(result["blocks"])

    return result


# ------------------------------------------------------------------
# Chinese engineering document heuristics (ported from report-from-notebooklm)
# ------------------------------------------------------------------

import re as _re_module

from ..language import CJK_RE as _CJK_RE
_FORMULA_HINT_RE = _re_module.compile(r"(=|\b(?:Re|Cd|C_D|CL|C_L|Cp|C_p)\b|ρ|ν|μ|sqrt|√)")
_QUESTION_RE = _re_module.compile(r"^\d+[.、)]")
_SECTION_HEADING_RE = _re_module.compile(r"^(?:[一二三四五六七八九十]+[、.]|[IVX]+[.])")
_NUMBER_RE = _re_module.compile(r"[-+]?\d+(?:\.\d+)?(?:[eE][-+]?\d+)?")
_CONSTANT_KEYWORDS = (
    "直徑", "弦長", "厚度", "密度", "黏度", "流速",
    "壓力", "攻角", "面積", "雷諾數",
    "Pa", "mm", "cm", "m/s", "kg", "gf", "Hz",
)
_TABLE_HINT_KEYWORDS = ("表", "Table", "data", "資料")


def _has_cjk_content(text: str) -> bool:
    """Return True if text contains significant CJK characters (>5%)."""
    if not text:
        return False
    cjk_count = len(_CJK_RE.findall(text[:2000]))
    return cjk_count > len(text[:2000]) * 0.05


def annotate_chinese_heuristics(blocks: list[dict]) -> None:
    """Enrich parsed blocks with Chinese engineering document annotations.

    Adds optional hint fields to each block:
      - chinese_hints.formula_hint: True if line contains formula patterns
      - chinese_hints.constant_hint: True if line contains constant/unit keywords
      - chinese_hints.question_hint: True if line looks like a numbered question
      - chinese_hints.table_hint: True if line references a table
      - chinese_hints.section_heading: True if line matches Chinese section heading pattern
    """
    for block in blocks:
        content = block.get("content", "")
        if not content:
            continue

        hints: dict = {}
        lines = content.splitlines()

        # Formula detection
        if _FORMULA_HINT_RE.search(content):
            hints["formula_hint"] = True

        # Constant/unit detection
        if any(kw in content for kw in _CONSTANT_KEYWORDS):
            hints["constant_hint"] = True
            # Also check if numbers are present alongside constants
            if _NUMBER_RE.search(content):
                hints["numeric_constant"] = True

        # Question detection (numbered questions in discussion sections)
        for line in lines:
            stripped = line.strip()
            if _QUESTION_RE.match(stripped):
                hints["question_hint"] = True
                break

        # Table reference detection
        if any(kw in content for kw in _TABLE_HINT_KEYWORDS):
            hints["table_hint"] = True

        # Chinese section heading detection
        for line in lines:
            stripped = line.strip()
            if _SECTION_HEADING_RE.match(stripped):
                hints["section_heading"] = True
                break

        if hints:
            block["chinese_hints"] = hints

