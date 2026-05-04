"""Template and fixed-field packaging reports."""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from ..state import ReportState
from .common import existing_path, write_json, write_text


def _style_font_details(style) -> dict:
    font = getattr(style, "font", None)
    if not font:
        return {}
    size = getattr(font, "size", None)
    return {
        "font_name": font.name or "",
        "font_size_pt": round(size.pt, 2) if size else None,
        "bold": bool(font.bold) if font.bold is not None else None,
        "italic": bool(font.italic) if font.italic is not None else None,
    }


def _document_style_summary(path: str | None) -> dict:
    if not path:
        return {"path": "", "exists": False, "load_error": ""}
    docx_path = Path(path)
    if not docx_path.exists():
        return {"path": str(docx_path), "exists": False, "load_error": ""}

    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        return {"path": str(docx_path), "exists": True, "load_error": str(exc)}

    paragraph_style_counts: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else "unknown"
        paragraph_style_counts[style_name] = paragraph_style_counts.get(style_name, 0) + 1

    table_style_counts: dict[str, int] = {}
    for table in doc.tables:
        style_name = table.style.name if table.style else "unknown"
        table_style_counts[style_name] = table_style_counts.get(style_name, 0) + 1

    paragraph_style_names = sorted(
        style.name for style in doc.styles
        if style.type == WD_STYLE_TYPE.PARAGRAPH
    )
    table_style_names = sorted(
        style.name for style in doc.styles
        if style.type == WD_STYLE_TYPE.TABLE
    )

    key_style_names = [
        "Normal",
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Body Text",
        "Caption",
    ]
    styles_by_name = {style.name: style for style in doc.styles}
    key_styles = {
        name: _style_font_details(style)
        for name in key_style_names
        if (style := styles_by_name.get(name))
    }

    section_summary = []
    for section in doc.sections:
        section_summary.append({
            "page_width_inches": round(section.page_width.inches, 2),
            "page_height_inches": round(section.page_height.inches, 2),
            "top_margin_inches": round(section.top_margin.inches, 2),
            "bottom_margin_inches": round(section.bottom_margin.inches, 2),
            "left_margin_inches": round(section.left_margin.inches, 2),
            "right_margin_inches": round(section.right_margin.inches, 2),
        })

    return {
        "path": str(docx_path),
        "exists": True,
        "load_error": "",
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "paragraph_style_counts": paragraph_style_counts,
        "table_style_counts": table_style_counts,
        "paragraph_style_names": paragraph_style_names,
        "table_style_names": table_style_names,
        "key_styles": key_styles,
        "sections": section_summary,
    }


def _normalize_field_key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", str(value).casefold()).strip("_")


def _field_label(key: str) -> str:
    return str(key).replace("_", " ").strip().title()


def _docx_text(path: str | None) -> str:
    if not path:
        return ""
    docx_path = Path(path)
    if not docx_path.exists():
        return ""
    try:
        doc = Document(str(docx_path))
    except Exception:
        return ""

    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_template_placeholders(text: str) -> list[dict]:
    placeholders: dict[str, dict] = {}
    patterns = [
        r"\{\{\s*([^{}]{2,60}?)\s*\}\}",
        r"<<\s*([^<>]{2,60}?)\s*>>",
        r"\[\s*([A-Za-z][A-Za-z0-9 _.-]{1,60}?)\s*\]",
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            raw = match.group(1).strip()
            if raw.lower().startswith(("cite:", "source:", "figure:")):
                continue
            key = _normalize_field_key(raw)
            if key:
                placeholders[key] = {
                    "key": key,
                    "label": _field_label(raw),
                    "placeholder": match.group(0),
                }
    return sorted(placeholders.values(), key=lambda item: item["key"])


def _front_matter_field_values(front_matter: dict) -> dict[str, dict]:
    fields: dict[str, dict] = {}
    standard = {
        "title": "Title",
        "short_title": "Running Title",
        "author_block": "Author",
        "affiliation_block": "Affiliation",
        "correspondence": "Correspondence",
        "acknowledgements": "Acknowledgements",
        "funding": "Funding",
        "conflict_note": "Conflict Of Interest",
    }
    for key, label in standard.items():
        value = front_matter.get(key)
        if value:
            fields[key] = {
                "key": key,
                "label": label,
                "value": str(value).strip(),
                "source": "front_matter",
            }

    keywords = front_matter.get("keywords")
    if keywords:
        fields["keywords"] = {
            "key": "keywords",
            "label": "Keywords",
            "value": ", ".join(str(keyword).strip() for keyword in keywords if str(keyword).strip()),
            "source": "front_matter",
        }

    template_fields = front_matter.get("template_fields") or {}
    if isinstance(template_fields, dict):
        for raw_key, raw_value in template_fields.items():
            key = _normalize_field_key(raw_key)
            value = str(raw_value).strip()
            if key and value:
                fields[key] = {
                    "key": key,
                    "label": _field_label(raw_key),
                    "value": value,
                    "source": "template_fields",
                }
    return fields


def _build_template_style_map_md(style_map: dict) -> str:
    rendered = style_map["rendered_docx"]
    reference = style_map["reference_docx"]
    comparison = style_map["style_comparison"]

    lines = [
        "# Template Style Map",
        "",
        f"- Status: {style_map['status']}",
        f"- Reference template mode: {style_map['reference_template_mode'] or 'unknown'}",
        f"- Renderer: {style_map['renderer_used'] or 'unknown'}",
        f"- Reference DOCX applied: {style_map['reference_docx_applied']}",
        f"- Reference DOCX: {reference['path'] or 'missing'}",
        f"- Rendered DOCX: {rendered['path'] or 'missing'}",
        "",
        "## Rendered Style Usage",
        "",
    ]
    for name, count in sorted(rendered.get("paragraph_style_counts", {}).items()):
        lines.append(f"- {name}: {count}")

    if rendered.get("table_style_counts"):
        lines.extend(["", "## Rendered Table Styles", ""])
        for name, count in sorted(rendered["table_style_counts"].items()):
            lines.append(f"- {name}: {count}")

    if reference.get("key_styles"):
        lines.extend(["", "## Reference Key Styles", ""])
        for name, details in reference["key_styles"].items():
            font_name = details.get("font_name") or "inherited"
            font_size = details.get("font_size_pt") or "inherited"
            lines.append(f"- {name}: {font_name}, {font_size} pt")

    lines.extend([
        "",
        "## Style Match",
        "",
        f"- Rendered paragraph styles defined in reference: {len(comparison['rendered_styles_defined_in_reference'])}",
        f"- Rendered paragraph styles not defined in reference: {len(comparison['rendered_styles_not_defined_in_reference'])}",
    ])
    if style_map["warnings"]:
        lines.extend(["", "## Warnings", ""])
        lines.extend(f"- {warning}" for warning in style_map["warnings"])

    lines.append("")
    return "\n".join(lines)


def build_template_style_map(state: ReportState, run_dir: Path) -> dict[str, str]:
    rendered_docx_path = state.output.get("final_docx_path") or state.output.get("rendered_docx_path")
    reference_docx_path = state.output.get("reference_docx_path", "")
    rendered = _document_style_summary(rendered_docx_path)
    reference = _document_style_summary(reference_docx_path)

    rendered_styles = set(rendered.get("paragraph_style_counts", {}))
    reference_styles = set(reference.get("paragraph_style_names", []))
    warnings = []
    if not rendered.get("exists"):
        warnings.append("Rendered DOCX is missing; style usage could not be inspected.")
    if rendered.get("load_error"):
        warnings.append(f"Rendered DOCX could not be read: {rendered['load_error']}")
    if not reference.get("exists"):
        warnings.append("Reference DOCX is missing; template style map is incomplete.")
    if reference.get("load_error"):
        warnings.append(f"Reference DOCX could not be read: {reference['load_error']}")
    if not state.output.get("reference_docx_applied"):
        warnings.append("Reference DOCX was not applied by the renderer; output may use fallback styling.")
    if state.output.get("renderer_used") != "pandoc":
        warnings.append("Renderer was not pandoc; DOCX style fidelity may be degraded.")

    style_map = {
        "job_id": state.job_id,
        "created_at": datetime.now().isoformat(),
        "status": "review" if warnings else "pass",
        "report_profile": state.spec.get("report_profile", ""),
        "reference_template_mode": state.spec.get("reference_template_mode", ""),
        "renderer_used": state.output.get("renderer_used", ""),
        "reference_docx_applied": bool(state.output.get("reference_docx_applied")),
        "reference_docx": reference,
        "rendered_docx": rendered,
        "style_comparison": {
            "rendered_styles_defined_in_reference": sorted(rendered_styles & reference_styles),
            "rendered_styles_not_defined_in_reference": sorted(rendered_styles - reference_styles),
            "reference_paragraph_style_count": len(reference_styles),
            "rendered_paragraph_style_count": len(rendered_styles),
        },
        "warnings": warnings,
    }

    json_path = run_dir / "template_style_map.json"
    md_path = run_dir / "template_style_map.md"
    write_json(json_path, style_map)
    write_text(md_path, _build_template_style_map_md(style_map))
    state.output["template_style_map_path"] = str(json_path)
    state.output["template_style_map_md_path"] = str(md_path)
    return {"json": str(json_path), "markdown": str(md_path)}


def _build_template_field_fill_report_md(report: dict) -> str:
    lines = [
        "# Template Field Fill Report",
        "",
        f"- Status: {report['status']}",
        f"- Reference template mode: {report['reference_template_mode'] or 'unknown'}",
        f"- Rendered DOCX: {report['rendered_docx_path'] or 'missing'}",
        f"- Reference DOCX: {report['reference_docx_path'] or 'missing'}",
        "",
        "## Fields",
        "",
    ]
    for field in report["fields"]:
        lines.append(
            f"- {field['label']}: {field['status']} "
            f"(source: {field['source']}, rendered matches: {field['rendered_match_count']})"
        )
    if report["warnings"]:
        lines.extend(["", "## Warnings", ""])
        lines.extend(f"- {warning}" for warning in report["warnings"])
    lines.append("")
    return "\n".join(lines)


def build_template_field_fill_report(state: ReportState, run_dir: Path) -> dict[str, str]:
    front_matter = state.plan.get("front_matter") or {}
    field_values = _front_matter_field_values(front_matter)
    rendered_docx_path = state.output.get("final_docx_path") or state.output.get("rendered_docx_path")
    reference_docx_path = state.output.get("reference_docx_path", "")
    rendered_text = _docx_text(rendered_docx_path)
    reference_text = _docx_text(reference_docx_path)
    placeholders = _extract_template_placeholders(reference_text)

    for placeholder in placeholders:
        field_values.setdefault(
            placeholder["key"],
            {
                "key": placeholder["key"],
                "label": placeholder["label"],
                "value": "",
                "source": "reference_placeholder",
            },
        )

    normalized_rendered = " ".join(rendered_text.split()).casefold()
    fields = []
    warnings = []
    for key in sorted(field_values):
        field = dict(field_values[key])
        value = str(field.get("value", "")).strip()
        rendered_match_count = 0
        if value:
            rendered_match_count = normalized_rendered.count(" ".join(value.split()).casefold())
        placeholder = next((item for item in placeholders if item["key"] == key), None)
        if not value:
            status = "missing_value"
            warnings.append(f"{field['label']} has a template placeholder but no structured value.")
        elif rendered_match_count > 0:
            status = "filled"
        else:
            status = "value_available_not_found"
            warnings.append(f"{field['label']} has a value but was not found in the rendered DOCX.")
        field.update({
            "placeholder": placeholder.get("placeholder", "") if placeholder else "",
            "required_by_reference": bool(placeholder),
            "rendered_match_count": rendered_match_count,
            "status": status,
        })
        fields.append(field)

    if state.spec.get("reference_template_mode") == "fixed_template" and not fields:
        warnings.append("fixed_template mode has no structured front matter or template fields to fill.")
    if rendered_docx_path and not rendered_text:
        warnings.append("Rendered DOCX text could not be inspected for field values.")

    report = {
        "job_id": state.job_id,
        "created_at": datetime.now().isoformat(),
        "status": "review" if warnings else "pass",
        "report_profile": state.spec.get("report_profile", ""),
        "reference_template_mode": state.spec.get("reference_template_mode", ""),
        "rendered_docx_path": existing_path(rendered_docx_path),
        "reference_docx_path": existing_path(reference_docx_path),
        "placeholder_count": len(placeholders),
        "field_count": len(fields),
        "filled_count": sum(1 for field in fields if field["status"] == "filled"),
        "missing_value_count": sum(1 for field in fields if field["status"] == "missing_value"),
        "not_found_count": sum(1 for field in fields if field["status"] == "value_available_not_found"),
        "fields": fields,
        "warnings": warnings,
    }

    json_path = run_dir / "template_field_fill_report.json"
    md_path = run_dir / "template_field_fill_report.md"
    write_json(json_path, report)
    write_text(md_path, _build_template_field_fill_report_md(report))
    state.output["template_field_fill_report_path"] = str(json_path)
    state.output["template_field_fill_report_md_path"] = str(md_path)
    return {"json": str(json_path), "markdown": str(md_path)}
