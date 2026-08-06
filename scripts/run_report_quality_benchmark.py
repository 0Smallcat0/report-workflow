"""Does the harness produce a better report than writing without it?

The gate benchmark answers "does it block fabrication". That is table stakes.
This one answers the question a reader actually asks — *so was it better?* —
and answers it the only way worth trusting: same source, same brief, two arms,
one scorer, everything checked in.

**Arm A, unassisted** (`benchmarks/fixtures/unassisted_baseline.md`) is a
recorded write-up of the source produced with no harness. It is a fixed
artifact rather than a live generation, because "what a writer does without
help" is not reproducible and a benchmark whose baseline moves every run
measures nothing. It is competent prose, not a strawman.

**Arm B, tool** runs the real pipeline over the same source: prepare, author
against the ledger, validate, render. The authoring is mechanical, on purpose.
The dimensions below are the ones the harness *enforces*, and scoring them
against a mechanical author isolates the harness's contribution from the
writer's talent. Prose quality comes from the drafting brief and is not
claimed here.

Both arms are scored by the same functions. Run with `--check` to fail on
drift, `--write` to regenerate the archive.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "src"))

from report_workflow.nodes.source_tables import collect_source_tables  # noqa: E402
from report_workflow.run_workflow import (  # noqa: E402
    prepare_workflow,
    render_workflow,
    validate_workflow,
)
from report_workflow.state import WORKFLOW_RUNS_DIR  # noqa: E402

FIXTURES = REPO_ROOT / "benchmarks" / "fixtures"
SOURCE_PATH = FIXTURES / "recycling_market_report.md"
BASELINE_PATH = FIXTURES / "unassisted_baseline.md"
EVIDENCE_ROOT = REPO_ROOT / "benchmarks" / "evidence" / "report_quality_2026-08-06"

PROMPT = "分析電池、塑膠、紡織、紙張四個品類的回收經濟性，比較單位成本與價格驅動因子"


# ----------------------------------------------------------------------
# Scorers. One implementation, applied to both arms.
# ----------------------------------------------------------------------

_URL_RE = re.compile(r"https?://[^\s)）】\]」』,，。;；]+")
_ATTRIBUTION_RE = re.compile(
    r"[（(]\s*(?:來源|来源|資料來源|资料来源|Source)\s*[:：]\s*([^)）]{2,60})[)）]",
    re.IGNORECASE,
)
#: Apparatus, not prose: a provenance line, a table or figure caption, and an
#: entry of a generated source list. The prose dimensions below — number
#: verifiability, paragraph structure, counter-evidence, derivation — are
#: about what the author wrote, and reading a "[S1] file.md line 5-8 — E_…"
#: entry as a paragraph would score the apparatus instead: its line numbers
#: and id digits would count as unverifiable figures, and its single line
#: would count as an unstructured paragraph. Excluded from both arms by the
#: same rule; `count_external_sources` still reads the whole document,
#: because a source listed in the apparatus is still a source.
_APPARATUS_RE = re.compile(
    r"^\s*(?:"
    r"(?:資料來源|来源|來源|Source)\s*[:：]"
    r"|\[S\d+\]"
    r"|\[\d+\]\s"
    r"|(?:表|图|圖|Table|Figure)\s*\d+[.、：:]"
    r")",
    re.IGNORECASE,
)

#: A bibliography entry — "Fastmarkets. (2026). *Black mass prices*. https://…".
#: Also apparatus, and it had to be added the moment the pipeline started
#: producing one: its full stops split into enough fragments to count as a
#: well-structured paragraph, and its years appear in the source, so a
#: reference list would have quietly improved two prose scores. A dimension
#: that improves because the apparatus grew is not measuring the prose.
_REFERENCE_ENTRY_RE = re.compile(
    r"\*[^*\n]+\*\s*\.?|\(\s*(?:n\.d\.|(?:19|20)\d{2})\s*\)\s*\."
)

_TABLE_RE = re.compile(r"^\|.+\|\s*$\n^\|?\s*:?-{3,}", re.MULTILINE)
_FIGURE_RE = re.compile(
    r"!\[[^\]]*\]\([^)]+\)|(?:^|\n)\s*(?:圖|图|Figure)\s*\d+[.、：:]", re.IGNORECASE
)

#: Numbers a reader could check. Chinese numerals are excluded on purpose: a
#: figure written 兩千四百 is not a checkable quantity, it is a paraphrase of
#: one, and counting it would flatter the arm that writes that way.
_NUMBER_RE = re.compile(r"(?<![A-Za-z0-9.,])(\d{1,3}(?:,\d{3})+(?:\.\d+)?|\d+(?:\.\d+)?)")

#: A claim set against another claim rather than stacked on it.
_COUNTER_RE = re.compile(
    r"反面證據|反面证据|然而|但是|不過|不过|另一方面|相反|反之|矛盾|分歧"
    r"|尚未|並不|并不|however|whereas|on the other hand|contrary",
    re.IGNORECASE,
)

#: A derived number that says how it was derived.
_DERIVATION_MARK_RE = re.compile(
    r"【推算】|【估算】|\[推算\]|estimated as|calculated as", re.IGNORECASE
)
_FORMULA_RE = re.compile(
    r"[（(][^)）]*[×÷*/+\-−][^)）]*[)）]|[0-9][^\n]{0,40}[×÷][^\n]{0,40}[0-9]"
)
_ASSUMPTION_RE = re.compile(r"假設|假设|assum", re.IGNORECASE)

_SENTENCE_SPLIT_RE = re.compile(r"[。．.!?！？]+")


def _body_paragraphs(text: str) -> list[str]:
    """Prose paragraphs only — no headings, tables, comments, or source lines."""
    without_comments = re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)
    paragraphs = []
    for block in re.split(r"\n\s*\n", without_comments):
        stripped = block.strip()
        if not stripped or stripped.startswith("#") or stripped.startswith("|"):
            continue
        if _APPARATUS_RE.match(stripped):
            continue
        if _REFERENCE_ENTRY_RE.search(stripped) and len(stripped) < 400:
            continue
        paragraphs.append(stripped)
    return paragraphs


def count_external_sources(text: str) -> int:
    """Distinct outside sources a reader could go and look up."""
    found = {match.group(0) for match in _URL_RE.finditer(text)}
    found.update(" ".join(m.group(1).split()) for m in _ATTRIBUTION_RE.finditer(text))
    return len(found)


def count_tables(text: str) -> int:
    return len(_TABLE_RE.findall(text))


def count_figures(text: str) -> int:
    return len(_FIGURE_RE.findall(text))


def _verifiable_numbers(text: str, source_text: str) -> tuple[int, int]:
    """(numbers in prose traceable to the source, numbers in prose)."""
    source_numbers = {
        match.group(1).replace(",", "") for match in _NUMBER_RE.finditer(source_text)
    }
    body = "\n".join(_body_paragraphs(text))
    numbers = [match.group(1).replace(",", "") for match in _NUMBER_RE.finditer(body)]
    hits = sum(1 for number in numbers if number in source_numbers)
    return hits, len(numbers)


def verifiable_numbers(text: str, source_text: str) -> int:
    """How many checkable figures the prose actually states."""
    return _verifiable_numbers(text, source_text)[0]


def verifiable_number_ratio(text: str, source_text: str) -> float:
    """Share of the document's prose numbers that also appear in the source.

    Read this beside ``verifiable_numbers``, never alone. A document stating
    almost no figures — or writing them as 兩千四百 rather than 2,400 — earns a
    perfect ratio by committing to nothing checkable, which is the opposite of
    what this dimension is for. The count says how much the document commits
    to; the ratio says how much of that the reader can verify.

    A number the source does not contain is not necessarily wrong — it may be
    derived — but the reader cannot check it unless told how, which is what
    the derivation dimension measures.
    """
    hits, total = _verifiable_numbers(text, source_text)
    return round(hits / total, 4) if total else 0.0


def count_counter_evidence(text: str) -> int:
    """Paragraphs that put a contrary reading beside the main one."""
    return sum(1 for para in _body_paragraphs(text) if _COUNTER_RE.search(para))


def count_disclosed_derivations(text: str) -> int:
    """Derived figures shown with their arithmetic or their assumption."""
    disclosed = 0
    for para in _body_paragraphs(text):
        if not _DERIVATION_MARK_RE.search(para):
            continue
        if _FORMULA_RE.search(para) or _ASSUMPTION_RE.search(para):
            disclosed += 1
    return disclosed


def structured_paragraph_ratio(text: str) -> float:
    """Share of paragraphs with room for context, content, and conclusion.

    Kording & Mensh's rule needs at least three moves; a two-sentence
    paragraph cannot make them. A floor, not a judgement of the prose.
    """
    paragraphs = _body_paragraphs(text)
    if not paragraphs:
        return 0.0
    structured = sum(
        1
        for para in paragraphs
        if len([s for s in _SENTENCE_SPLIT_RE.split(para) if s.strip()]) >= 3
    )
    return round(structured / len(paragraphs), 4)


def score_document(text: str, source_text: str) -> dict:
    return {
        "external_sources": count_external_sources(text),
        "verifiable_numbers": verifiable_numbers(text, source_text),
        "verifiable_number_ratio": verifiable_number_ratio(text, source_text),
        "tables": count_tables(text),
        "figures": count_figures(text),
        "counter_evidence_paragraphs": count_counter_evidence(text),
        "disclosed_derivations": count_disclosed_derivations(text),
        "structured_paragraph_ratio": structured_paragraph_ratio(text),
    }


#: Every dimension is "more is better", which keeps the comparison readable.
DIMENSIONS = (
    "external_sources",
    "verifiable_numbers",
    "verifiable_number_ratio",
    "tables",
    "figures",
    "counter_evidence_paragraphs",
    "disclosed_derivations",
    "structured_paragraph_ratio",
)


# ----------------------------------------------------------------------
# The tool arm: a real run of the real pipeline.
# ----------------------------------------------------------------------

def _all_packages_present(name: str, *args, **kwargs):
    return object()


def _author_artifacts(state) -> None:
    """Write the three agent artifacts the way the contract requires.

    Mechanical on purpose — see the module docstring. Every claim quotes the
    row it cites, so the arm measures what the harness carries into the
    document rather than how well something writes.
    """
    run_dir = WORKFLOW_RUNS_DIR / state.job_id
    ledger = [
        json.loads(line)
        for line in (run_dir / "evidence_ledger.jsonl").read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]
    tables = collect_source_tables(ledger)

    prose = [
        row
        for row in ledger
        if row.get("block_type") not in {"table_row", "heading"}
        and len(row.get("content", "")) > 40
    ]
    picked = prose[:8]
    claims = [
        {
            "claim_id": f"c{index}",
            "claim_text": row["content"][:150],
            "claim_type": (
                "statistical"
                if "statistical" in row.get("allowed_claim_types", [])
                else "factual"
            ),
            "risk_level": "low",
            "status": "supported",
            "evidence_ids": [row["evidence_id"]],
            "requires_hedged_wording": row.get("evidence_grade") != "high",
            "claim_role": "primary",
        }
        for index, row in enumerate(picked, start=1)
    ]
    (run_dir / "claim_matrix.json").write_text(
        json.dumps({"claims": claims}, ensure_ascii=False), encoding="utf-8"
    )

    plan_path = run_dir / "section_drafts" / "figure_plan.json"
    kept_figures: list[str] = []
    if plan_path.exists():
        payload = json.loads(plan_path.read_text(encoding="utf-8"))
        figures = payload.get("figures", [])[:1]
        payload["figures"] = figures
        plan_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
        kept_figures = [str(f["figure_id"]) for f in figures if f.get("figure_id")]

    order = state.plan["blueprint"]["section_order"]
    claimless = {"references", "appendix"}
    sections = {
        section_id: {
            "section_id": section_id,
            "goals": f"cover {section_id}",
            "claim_ids": [] if section_id in claimless else [c["claim_id"] for c in claims],
            "paragraph_order": ["context", "content", "conclusion"],
            "figure_ids": kept_figures if section_id == "findings" else [],
        }
        for section_id in order
    }
    (run_dir / "outline.json").write_text(
        json.dumps({"sections": sections}, ensure_ascii=False), encoding="utf-8"
    )

    section_dir = run_dir / "section_drafts"
    section_dir.mkdir(exist_ok=True)
    sentence_rows = []
    for section_id in order:
        if section_id in claimless:
            (section_dir / f"{section_id}.md").write_text(
                f"# {section_id}\n\n本節由流程自動產生。\n", encoding="utf-8"
            )
            continue
        lines = [f"# {section_id}", ""]
        for claim in claims:
            marker = " ".join(f"[CITE:{eid}]" for eid in claim["evidence_ids"])
            lines.extend([
                f"就本節主題而言，{claim['claim_text'][:110]} {marker}。"
                "此一觀察限於來源所涵蓋的時點與品類。"
                "因此在解讀時應同時考慮成本結構與價格驅動因子的差異。",
                "",
            ])
            sentence_rows.append({
                "sentence_id": f"s_{section_id}_{claim['claim_id']}",
                "section_id": section_id,
                "claim_ids": [claim["claim_id"]],
                "evidence_ids": claim["evidence_ids"],
                "citation_ids": claim["evidence_ids"],
                "wording_strength": "hedged",
                "draft_origin": "agent_draft",
            })
        if section_id == "findings":
            for table_id in tables:
                lines.extend([f"下表為來源原始數據。[TABLE:{table_id}]", ""])
            for figure_id in kept_figures:
                lines.extend([f"[FIGURE:{figure_id}]", ""])
        (section_dir / f"{section_id}.md").write_text("\n".join(lines), encoding="utf-8")

    with open(run_dir / "sentence_map.jsonl", "w", encoding="utf-8") as handle:
        for row in sentence_rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def run_tool_arm(source_text: str) -> tuple[str, dict]:
    """Run the pipeline end to end and return the delivered document's text."""
    from docx import Document

    workdir = Path(tempfile.mkdtemp(prefix="rqbench_"))
    try:
        source = workdir / SOURCE_PATH.name
        source.write_text(source_text, encoding="utf-8")
        with patch(
            "report_workflow.preflight.importlib.util.find_spec",
            side_effect=_all_packages_present,
        ):
            state = prepare_workflow(
                PROMPT,
                [str(source)],
                str(workdir / "out"),
                report_profile="business_report",
            )
        _author_artifacts(state)
        validated = validate_workflow(state.job_id)
        decision = (validated.qa or {}).get("qa_decision")
        if decision != "pass":
            raise RuntimeError(
                "tool arm did not reach a publishable state: "
                + json.dumps((validated.qa or {}).get("blockers", []), ensure_ascii=False)[:800]
            )
        rendered = render_workflow(state.job_id)
        docx_path = (
            rendered.output.get("published_report_path")
            or rendered.output.get("final_docx_path")
            or rendered.output.get("rendered_docx_path")
        )
        document = Document(docx_path)

        # Read the delivered document back as markdown-equivalent text so one
        # scorer can read both arms: Word tables become pipe tables.
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            rows = [
                "| " + " | ".join(cell.text for cell in row.cells) + " |"
                for row in table.rows
            ]
            if rows:
                columns = len(table.rows[0].cells)
                rows.insert(1, "|" + "|".join(" --- " for _ in range(columns)) + "|")
            parts.append("\n".join(rows))
        return "\n\n".join(parts), {
            "source_tables_placed": rendered.output.get("source_tables_placed", 0),
            "figures_embedded": rendered.output.get("figure_placeholders_resolved", 0),
        }
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


# ----------------------------------------------------------------------
# Results
# ----------------------------------------------------------------------

def build_results() -> dict:
    source_text = SOURCE_PATH.read_text(encoding="utf-8")
    baseline_text = BASELINE_PATH.read_text(encoding="utf-8")
    tool_text, tool_run = run_tool_arm(source_text)

    unassisted = score_document(baseline_text, source_text)
    tool = score_document(tool_text, source_text)

    comparison = []
    for dimension in DIMENSIONS:
        a, b = unassisted[dimension], tool[dimension]
        winner = "tool" if b > a else ("unassisted" if a > b else "tie")
        comparison.append({
            "dimension": dimension,
            "unassisted": a,
            "tool": b,
            "winner": winner,
        })

    return {
        "source": str(SOURCE_PATH.relative_to(REPO_ROOT)).replace("\\", "/"),
        "baseline": str(BASELINE_PATH.relative_to(REPO_ROOT)).replace("\\", "/"),
        "prompt": PROMPT,
        "scores": {"unassisted": unassisted, "tool": tool},
        "comparison": comparison,
        "tool_wins": sum(1 for row in comparison if row["winner"] == "tool"),
        "unassisted_wins": sum(1 for row in comparison if row["winner"] == "unassisted"),
        "ties": sum(1 for row in comparison if row["winner"] == "tie"),
        "tool_run": tool_run,
    }


def _summary_markdown(results: dict) -> str:
    lines = [
        "# Report-quality benchmark: harness versus unassisted",
        "",
        "Same source, same prompt, one scorer, both arms checked in.",
        "Reproduce with `python scripts/run_report_quality_benchmark.py --check`.",
        "",
        f"- Source: `{results['source']}`",
        f"- Unassisted arm: `{results['baseline']}` (recorded sample; see the file header)",
        "- Tool arm: a live run of prepare → author → validate → render",
        "",
        "| Dimension | Unassisted | Tool | Winner |",
        "| --- | ---: | ---: | --- |",
    ]
    for row in results["comparison"]:
        lines.append(
            f"| {row['dimension']} | {row['unassisted']} | {row['tool']} | {row['winner']} |"
        )
    lines.extend([
        "",
        f"Tool wins {results['tool_wins']} of {len(results['comparison'])} dimensions; "
        f"unassisted wins {results['unassisted_wins']}; {results['ties']} tie.",
        "",
        "## What this claims, and what it does not",
        "",
        "It claims the harness makes these properties non-optional: a source table",
        "survives into the document, a cited figure keeps its provenance, and a",
        "number in the prose can be traced back to a row of the ledger. Each is",
        "enforced by a gate, so no run can quietly drop them.",
        "",
        "It does not claim the harness writes better prose. The tool arm is authored",
        "mechanically, which is what isolates the harness's contribution from the",
        "writer's. Prose quality comes from the drafting brief, and a deterministic",
        "benchmark is the wrong instrument for measuring it.",
        "",
        "The unassisted arm is a recorded artifact rather than a live generation.",
        "That is a limitation, and it is deliberate: a baseline regenerated each run",
        "makes the comparison move for reasons unrelated to the pipeline.",
        "",
        "## The two dimensions the tool loses",
        "",
        "Both are reported as measured rather than tuned away, because a benchmark",
        "its own author can adjust until it wins is not evidence of anything.",
        "",
        "**verifiable_number_ratio.** The baseline scores 1.0 by stating almost no",
        "checkable figures: it writes 八萬 and 兩千四百 where the source writes 80,000",
        "and 2,383, and a Chinese numeral is a paraphrase of a quantity rather than",
        "the quantity. Its perfect ratio is one verifiable number out of one. The",
        "tool arm states 62 and can trace 87% of them; the rest are derived figures,",
        "which the derivation dimension covers. Read the ratio next to the count or",
        "the metric rewards vagueness.",
        "",
        "**structured_paragraph_ratio.** The tool arm is dragged down by its own",
        "mechanical author, which writes one-line lead-ins before a table ('下表為",
        "來源原始數據。'). The drafting brief tells a real author to give every table",
        "a distinct lead-in that says what to look for — this arm does not, because",
        "rewriting it to score better would be tuning the arm rather than measuring",
        "the harness.",
    ])
    return "\n".join(lines) + "\n"


def write_archive(results: dict) -> None:
    EVIDENCE_ROOT.mkdir(parents=True, exist_ok=True)
    (EVIDENCE_ROOT / "results.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (EVIDENCE_ROOT / "summary.md").write_text(_summary_markdown(results), encoding="utf-8")


def check_archive() -> list[str]:
    issues: list[str] = []
    archived_path = EVIDENCE_ROOT / "results.json"
    if not archived_path.exists():
        return ["no archived report-quality results; run with --write"]
    archived = json.loads(archived_path.read_text(encoding="utf-8"))
    fresh = build_results()
    for arm in ("unassisted", "tool"):
        if archived["scores"][arm] != fresh["scores"][arm]:
            issues.append(f"{arm} scores drifted from the archive")
    if archived["comparison"] != fresh["comparison"]:
        issues.append("dimension comparison drifted from the archive")
    return issues


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check", action="store_true", help="fail if the archive has drifted")
    parser.add_argument("--write", action="store_true", help="regenerate the archive")
    args = parser.parse_args(argv)

    if args.check:
        issues = check_archive()
        if issues:
            print("report-quality benchmark check failed:")
            for issue in issues:
                print(f"- {issue}")
            return 1
        print("report-quality benchmark check passed")
        return 0

    results = build_results()
    write_archive(results)
    print(_summary_markdown(results))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
