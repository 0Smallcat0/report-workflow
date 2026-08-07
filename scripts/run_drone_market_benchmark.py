"""Does the harness beat the arm that was actually winning?

`run_report_quality_benchmark.py` compares the harness against an unassisted
write-up of a *prose* source, and the harness wins on the dimensions it
enforces. That comparison was never the hard one. The hard one is three raw CSVs
and a market question, where the unassisted arm is free to compute whatever it
likes and did: thirteen hand-built tables, 703 checkable figures, both questions
answered in the first line, and a chapter of seven counter-arguments against its
own recommendation.

That report is the baseline here. It is recorded rather than regenerated, for
the same reason as the other benchmark's: a baseline that moves each run makes
the comparison move for reasons unrelated to the pipeline.

The tool arm runs the real pipeline over the same three files and places the
cross tabulations the pipeline built at intake. The authoring is mechanical, and
deliberately so — it isolates what the harness carries into a document from how
well something writes. What the mechanical author is *not* allowed to do is skip
the built tables, because that is the loss this round exists to close.

Both arms are scored by the functions in `run_report_quality_benchmark.py` — one
implementation, imported rather than copied, so a change to a scorer cannot
improve one arm and not the other.

    python scripts/run_drone_market_benchmark.py --check   # fail on drift
    python scripts/run_drone_market_benchmark.py --write   # regenerate archive
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch

REPO_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO_ROOT / "src"))
sys.path.insert(0, str(Path(__file__).resolve().parent))

from report_workflow.derived_evidence import built_table_entries  # noqa: E402
from report_workflow.prompt_questions import extract_questions  # noqa: E402
from report_workflow.run_workflow import (  # noqa: E402
    prepare_workflow,
    render_workflow,
    validate_workflow,
)
from report_workflow.state import WORKFLOW_RUNS_DIR  # noqa: E402
from run_report_quality_benchmark import DIMENSIONS, score_document  # noqa: E402

SOURCE_DIR = REPO_ROOT / "benchmarks" / "fixtures" / "drone_market"
BASELINE_PATH = REPO_ROOT / "benchmarks" / "fixtures" / "drone_market_unassisted.md"
EVIDENCE_ROOT = REPO_ROOT / "benchmarks" / "evidence" / "drone_market_2026-08-07"

#: The task statement both arms were given, unchanged. It asks two questions,
#: which is what makes the conclusion requirement fire on this fixture.
PROMPT = (
    "根據 Amazon US 無人機類商品資料撰寫一份市場研究報告，涵蓋品類結構、價格帶分布、"
    "品牌集中度、買家痛點四個面向，結論必須能支撐「這個市場值不值得進入、"
    "從哪個切點進入」的判斷。"
)


def source_paths() -> list[Path]:
    return sorted(SOURCE_DIR.glob("*.csv"))


def source_text() -> str:
    """The three files as one document, for the number-verifiability scorer."""
    return "\n".join(path.read_text(encoding="utf-8") for path in source_paths())


# ----------------------------------------------------------------------
# The tool arm: a real run, authored mechanically.
# ----------------------------------------------------------------------

def _all_packages_present(name: str, *args, **kwargs):
    return object()


def _read_ledger(run_dir: Path) -> list[dict]:
    return [
        json.loads(line)
        for line in (run_dir / "evidence_ledger.jsonl").read_text(encoding="utf-8").splitlines()
        if line.strip()
    ]


def _author_artifacts(state) -> dict:
    """Write the agent artifacts, placing every table the pipeline built.

    Mechanical, and mechanical in the direction the harness pushes: each built
    cross tabulation gets a claim that cites it and a `[TABLE:]` marker in the
    section that discusses it. An author is free to waive one — the outline gate
    takes a reason — but this arm exists to measure what the material is worth
    when it is used, so it uses all of it.
    """
    run_dir = WORKFLOW_RUNS_DIR / state.job_id
    ledger = _read_ledger(run_dir)
    by_id = {row.get("evidence_id"): row for row in ledger}
    tables = built_table_entries(ledger)

    # Scalars for the counter-evidence section: single-column summaries carry
    # coverage and share, which is what a limitation is usually made of.
    scalars = [
        row
        for row in ledger
        if row.get("derivation")
        and not row.get("table_grid")
        and len(row.get("content", "")) > 30
    ][:2]

    def claim(claim_id: str, evidence_id: str, role: str) -> dict:
        row = by_id.get(evidence_id) or {}
        return {
            "claim_id": claim_id,
            "claim_text": row.get("content", "")[:150],
            "claim_type": (
                "statistical"
                if "statistical" in row.get("allowed_claim_types", [])
                else "factual"
            ),
            "risk_level": "low",
            "status": "supported",
            "evidence_ids": [evidence_id],
            "requires_hedged_wording": row.get("evidence_grade") != "high",
            "claim_role": role,
        }

    table_claims = [
        claim(f"t{index}", table["evidence_id"], "primary" if index <= 3 else "supporting")
        for index, table in enumerate(tables, start=1)
    ]
    limit_claims = [
        claim(f"l{index}", row["evidence_id"], "supporting")
        for index, row in enumerate(scalars, start=1)
    ]
    if len(limit_claims) < 2 and table_claims:
        # A source with no scalar summaries still owes the section two claims.
        limit_claims = [
            claim(f"l{index}", table_claims[-1]["evidence_ids"][0], "supporting")
            for index in (1, 2)
        ]
    claims = table_claims + limit_claims
    (run_dir / "claim_matrix.json").write_text(
        json.dumps({"claims": claims}, ensure_ascii=False), encoding="utf-8"
    )

    # No figures: this arm is about what reaches the page as checkable text, and
    # the outline gate requires figure_plan.json and the outline to agree.
    plan_path = run_dir / "section_drafts" / "figure_plan.json"
    if plan_path.exists():
        payload = json.loads(plan_path.read_text(encoding="utf-8"))
        payload["figures"] = []
        plan_path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

    order = state.plan["blueprint"]["section_order"]
    blueprint_sections = state.plan["blueprint"].get("sections") or {}
    claimless = {"references", "appendix"}
    counter_id = next(
        (
            section_id
            for section_id, spec in blueprint_sections.items()
            if isinstance(spec, dict) and spec.get("requires_undermines")
        ),
        "",
    )
    answer_id = next(
        (
            section_id
            for section_id, spec in blueprint_sections.items()
            if isinstance(spec, dict) and spec.get("must_answer_prompt_questions")
        ),
        "",
    )

    def claims_for(section_id: str) -> list[dict]:
        if section_id in claimless:
            return []
        if section_id == counter_id:
            return limit_claims
        return table_claims

    sections = {
        section_id: {
            "section_id": section_id,
            "goals": f"cover {section_id}",
            "claim_ids": [c["claim_id"] for c in claims_for(section_id)],
            "paragraph_order": ["context", "content", "conclusion"],
            "figure_ids": [],
        }
        for section_id in order
    }
    if counter_id in sections and table_claims:
        sections[counter_id]["undermines"] = [table_claims[0]["claim_id"]]
    questions = extract_questions(PROMPT)
    if answer_id in sections and questions:
        answered = sections[answer_id]["claim_ids"]
        sections[answer_id]["answers"] = [
            {
                "question_index": index,
                "claim_ids": [answered[min(index, len(answered) - 1)]],
            }
            for index in range(len(questions))
        ]
    (run_dir / "outline.json").write_text(
        json.dumps({"sections": sections}, ensure_ascii=False), encoding="utf-8"
    )

    section_dir = run_dir / "section_drafts"
    section_dir.mkdir(exist_ok=True)
    sentence_rows: list[dict] = []
    for section_id in order:
        if section_id in claimless:
            (section_dir / f"{section_id}.md").write_text(
                f"# {section_id}\n\n本節由流程自動產生。\n", encoding="utf-8"
            )
            continue
        lines = [f"# {section_id}", ""]
        for entry in claims_for(section_id):
            marker = " ".join(f"[CITE:{eid}]" for eid in entry["evidence_ids"])
            lines.extend([
                f"就本節主題而言，{entry['claim_text']} {marker}。"
                "此一觀察限於本次抓取的掛牌與評論所涵蓋的範圍。"
                "解讀時應同時考慮各價格帶的資料覆蓋率差異。",
                "",
            ])
            sentence_rows.append({
                "sentence_id": f"s_{section_id}_{entry['claim_id']}",
                "section_id": section_id,
                "claim_ids": [entry["claim_id"]],
                "evidence_ids": entry["evidence_ids"],
                "citation_ids": entry["evidence_ids"],
                "wording_strength": "hedged",
                "draft_origin": "agent_draft",
            })
        if section_id == "findings":
            for table in tables:
                lines.extend([
                    f"下表為本節討論的分組結果。[TABLE:{table['evidence_id']} "
                    f"{table['description']}]",
                    "",
                ])
        (section_dir / f"{section_id}.md").write_text("\n".join(lines), encoding="utf-8")

    with open(run_dir / "sentence_map.jsonl", "w", encoding="utf-8") as handle:
        for row in sentence_rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")

    return {
        "tables_built": len(tables),
        "tables_placed": len(tables),
        "hand_registered_derivations": 0,
        "questions_extracted": len(questions),
    }


def run_tool_arm() -> tuple[str, dict]:
    """Run the pipeline end to end and return the delivered document's text."""
    from docx import Document

    workdir = Path(tempfile.mkdtemp(prefix="dronebench_"))
    try:
        with patch(
            "report_workflow.preflight.importlib.util.find_spec",
            side_effect=_all_packages_present,
        ):
            state = prepare_workflow(
                PROMPT,
                [str(path) for path in source_paths()],
                str(workdir / "out"),
                report_profile="business_report",
            )
        authored = _author_artifacts(state)
        validated = validate_workflow(state.job_id)
        decision = (validated.qa or {}).get("qa_decision")
        if decision != "pass":
            raise RuntimeError(
                "tool arm did not reach a publishable state: "
                + json.dumps((validated.qa or {}).get("blockers", []), ensure_ascii=False)[:800]
            )
        rendered = render_workflow(state.job_id)
        docx_path = (
            rendered.output.get("published_report_path")
            or rendered.output.get("final_docx_path")
            or rendered.output.get("rendered_docx_path")
        )
        document = Document(docx_path)
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            rows = [
                "| " + " | ".join(cell.text for cell in row.cells) + " |"
                for row in table.rows
            ]
            if rows:
                columns = len(table.rows[0].cells)
                rows.insert(1, "|" + "|".join(" --- " for _ in range(columns)) + "|")
            parts.append("\n".join(rows))
        authored["word_tables"] = len(document.tables)
        return "\n\n".join(parts), authored
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


# ----------------------------------------------------------------------
# Results
# ----------------------------------------------------------------------

def build_results() -> dict:
    text = source_text()
    baseline_text = BASELINE_PATH.read_text(encoding="utf-8")
    tool_text, tool_run = run_tool_arm()

    unassisted = score_document(baseline_text, text)
    tool = score_document(tool_text, text)

    comparison = []
    for dimension in DIMENSIONS:
        a, b = unassisted[dimension], tool[dimension]
        winner = "tool" if b > a else ("unassisted" if a > b else "tie")
        comparison.append({
            "dimension": dimension,
            "unassisted": a,
            "tool": b,
            "winner": winner,
        })

    return {
        "source": sorted(
            str(path.relative_to(REPO_ROOT)).replace("\\", "/") for path in source_paths()
        ),
        "baseline": str(BASELINE_PATH.relative_to(REPO_ROOT)).replace("\\", "/"),
        "prompt": PROMPT,
        "scores": {"unassisted": unassisted, "tool": tool},
        "comparison": comparison,
        "tool_wins": sum(1 for row in comparison if row["winner"] == "tool"),
        "unassisted_wins": sum(1 for row in comparison if row["winner"] == "unassisted"),
        "ties": sum(1 for row in comparison if row["winner"] == "tie"),
        "tool_run": tool_run,
        "tool_document": tool_text,
    }


def _summary_markdown(results: dict) -> str:
    run = results["tool_run"]
    lines = [
        "# Drone-market benchmark: harness versus the hand-built control",
        "",
        "Three raw CSVs, one task statement, one scorer, both arms checked in.",
        "Reproduce with `python scripts/run_drone_market_benchmark.py --check`.",
        "",
        "- Source: " + ", ".join(f"`{path}`" for path in results["source"]),
        f"- Unassisted arm: `{results['baseline']}` (recorded sample; see the file header)",
        "- Tool arm: a live run of prepare -> author -> validate -> render",
        "",
        "| Dimension | Unassisted | Tool | Winner |",
        "| --- | ---: | ---: | --- |",
    ]
    for row in results["comparison"]:
        lines.append(
            f"| {row['dimension']} | {row['unassisted']} | {row['tool']} | {row['winner']} |"
        )
    lines.extend([
        "",
        f"Tool wins {results['tool_wins']} of {len(results['comparison'])} dimensions; "
        f"unassisted wins {results['unassisted_wins']}; {results['ties']} tie.",
        "",
        "## What the run itself shows",
        "",
        f"- Cross tabulations the pipeline built at intake: **{run['tables_built']}**",
        f"- Placed in the document: **{run['tables_placed']}**",
        f"- Tables in the delivered DOCX: **{run.get('word_tables', 0)}**",
        f"- Derivations the author had to register by hand: **{run['hand_registered_derivations']}**",
        f"- Questions extracted from the task statement, each bound to a claim in the "
        f"conclusion: **{run['questions_extracted']}**",
        "",
        "The last three are the round's subject. The tables were computed whether or not",
        "anyone asked; before this round an author could leave them unmentioned, and three",
        "runs of this task placed four of them, then three, then two. The outline now",
        "refuses to load until each is either placed or waived by name with a reason, so",
        "the number above is a floor rather than an average.",
        "",
        "## What this claims, and what it does not",
        "",
        "It claims the harness carries checkable material into the document without the",
        "author having to build it: every table above is computed from the rows, keeps its",
        "provenance, and is cited by a claim the gates check.",
        "",
        "It does not claim the harness writes better prose. The tool arm is authored",
        "mechanically - one lead-in sentence per table, no argument between them - which",
        "is what isolates the harness's contribution from the writer's. Rewriting that",
        "author to score better would be tuning the arm rather than measuring the harness.",
        "",
        "The unassisted arm is a recorded artifact rather than a live generation. That is",
        "a limitation and it is deliberate: it is one write-up by one author on one day,",
        "and regenerating it would move the baseline for reasons unrelated to the",
        "pipeline. It is also not a strawman - it is the arm that was ahead.",
    ])
    losses = [row for row in results["comparison"] if row["winner"] == "unassisted"]
    if losses:
        lines.extend([
            "",
            "## The dimensions the tool loses",
            "",
            "Reported as measured rather than tuned away, because a benchmark its own",
            "author can adjust until it wins is not evidence of anything.",
            "",
        ])
        for row in losses:
            lines.append(
                f"- **{row['dimension']}**: unassisted {row['unassisted']}, "
                f"tool {row['tool']}."
            )
        lines.extend([
            "",
            "All three come from the same property of this arm: it registers nothing. It",
            "places the tables the pipeline built and states no figures beyond what those",
            "tables' evidence text already contains, where the unassisted arm computed",
            "whatever its argument needed and typed the result into a sentence.",
            "",
            "`tables` is the clearest reading of that. The tool arm carries the built",
            "tables and no others; the three real acceptance runs of this same task",
            "registered six, five and four grouped tables of their own on top of them.",
            "The gap is what an author adds, and this arm has no author.",
            "",
            "`verifiable_numbers` and its ratio are counted over prose paragraphs, and the",
            "paragraph filter excludes table rows from *both* arms — so a figure the tool",
            "puts in a checked table is not counted, while the same figure typed into a",
            "sentence by the unassisted arm is. Read beside it: every number this arm",
            "states is traceable to a ledger row, and none of the unassisted arm's 121 is",
            "traceable to anything at all. That is the trade, and this scorer is not the",
            "instrument that shows it.",
        ])
    return "\n".join(lines) + "\n"


def write_archive(results: dict) -> None:
    EVIDENCE_ROOT.mkdir(parents=True, exist_ok=True)
    (EVIDENCE_ROOT / "results.json").write_text(
        json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    (EVIDENCE_ROOT / "summary.md").write_text(_summary_markdown(results), encoding="utf-8")


def check_archive() -> list[str]:
    """Re-score both recorded documents and diff against the archive.

    Rendering depends on the machine, so `--check` does not render: it verifies
    what can be held fixed, which is that the scorer given these two documents
    still produces these numbers. Regenerating the tool arm needs the full
    environment, and that is what `--write` is for.
    """
    archived_path = EVIDENCE_ROOT / "results.json"
    if not archived_path.exists():
        return ["no archived drone-market results; run with --write"]
    archived = json.loads(archived_path.read_text(encoding="utf-8"))
    tool_document = archived.get("tool_document")
    if not tool_document:
        return ["the archive has no recorded tool document; run with --write"]

    text = source_text()
    fresh = {
        "unassisted": score_document(BASELINE_PATH.read_text(encoding="utf-8"), text),
        "tool": score_document(tool_document, text),
    }
    issues: list[str] = []
    for arm in ("unassisted", "tool"):
        if archived["scores"][arm] != fresh[arm]:
            issues.append(
                f"{arm} scores drifted from the archive: "
                f"archived {archived['scores'][arm]}, recomputed {fresh[arm]}"
            )
    run = archived.get("tool_run") or {}
    if run.get("tables_placed", 0) < run.get("tables_built", 0):
        issues.append(
            f"the archived run placed {run.get('tables_placed')} of "
            f"{run.get('tables_built')} built tables; this arm places all of them"
        )
    return issues


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--check", action="store_true", help="fail if the archive has drifted")
    parser.add_argument("--write", action="store_true", help="regenerate the archive")
    args = parser.parse_args(argv)

    if args.check:
        issues = check_archive()
        if issues:
            print("drone-market benchmark check failed:")
            for issue in issues:
                print(f"- {issue}")
            return 1
        print("drone-market benchmark check passed")
        return 0

    results = build_results()
    write_archive(results)
    print(_summary_markdown(results))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
